import os
import hashlib
import html
import json
import re
import secrets
import shutil
import smtplib
import sqlite3
import tempfile
import threading
import time
import zipfile
from datetime import date, datetime, timedelta, timezone
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from email.message import EmailMessage
from functools import wraps
from io import BytesIO
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qsl, urlencode, urlsplit, urlunsplit
from urllib.request import Request, urlopen
from zoneinfo import ZoneInfo, ZoneInfoNotFoundError

from flask import (
    Flask,
    abort,
    flash,
    g,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    session,
    url_for,
)
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageOps
from pillow_heif import register_heif_opener
from image_processing import compress_image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from reportlab.platypus import LongTable, Paragraph, SimpleDocTemplate, Spacer, TableStyle


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
DB_PATH = os.path.join(DATA_DIR, "invoices.db")
ATTACHMENTS_DIR = os.path.join(DATA_DIR, "attachments")
CONTRACT_ATTACHMENTS_DIR = os.path.join(DATA_DIR, "contract-attachments")
REPORT_ATTACHMENTS_DIR = os.path.join(DATA_DIR, "service-report-attachments")
EXPENSE_ATTACHMENTS_DIR = os.path.join(DATA_DIR, "expense-attachments")
CUSTOMER_REIMBURSEMENT_DIR = os.path.join(DATA_DIR, "customer-reimbursements")
COMPANY_ATTACHMENT_DIR = os.path.join(DATA_DIR, "company-attachments")
USER_ATTACHMENT_DIR = os.path.join(DATA_DIR, "user-attachments")
SHARED_PHOTOS_DIR = os.environ.get("SHARED_PHOTOS_DIR", "/app/shared-photos")
ALLOWED_ATTACHMENT_EXTENSIONS = {"pdf", "png", "jpg", "jpeg", "webp", "gif", "doc", "docx", "xls", "xlsx"}
ALLOWED_IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "webp", "gif", "heic", "heif"}
REPORT_PHOTO_FOLDERS = {
    "arrival": "现场到达时间照片",
    "departure": "离开现场时间照片",
    "self_check": "自检照片",
    "site": "现场服务照片",
    "mileage_proof": "里程佐证",
}

register_heif_opener()
ALLOWED_ATTACHMENT_LABEL = "Word、Excel、PDF、PNG、JPG、JPEG、WEBP、GIF"

DEFAULT_COMPANY_PROFILE = {
    "name": "Prasinos Power LLC",
    "address": "6131 Fenske Lane, Needville, TX 77461, US",
    "email": "info@prasinospower.com",
    "phone": "+1 910 910 9191",
    "registration_number": "805498356",
    "ein": "99-2440677",
    "tax_note": (
        "No Texas sales tax is charged on this invoice because the customer is located outside "
        "the United States. Customer is responsible for any applicable local taxes, withholding, "
        "foreign exchange costs, intermediary bank fees, or wire transfer fees unless otherwise agreed in writing."
    ),
}

DEFAULT_PAYMENT_INSTRUCTIONS = {
    "method": "International bank wire transfer",
    "beneficiary": "Prasinos Power LLC",
    "bank_name": "Chase Bank",
    "account_number": "2909930519",
    "routing_number": "111000614",
    "swift_bic": "CHASUS33XXX",
}

DEFAULT_INVOICE_TERMS = (
    "Payment is due by the due date stated on this invoice. Bank fees, intermediary fees, "
    "and foreign exchange charges are the responsibility of the payer unless otherwise agreed in writing."
)

DEFAULT_SMTP_SETTINGS = {
    "host": "smtp.gmail.com",
    "port": "587",
    "user": "delanochen@gmail.com",
    "password": "",
    "from": "delanochen@gmail.com",
    "tls": "true",
}
DEFAULT_TIMEZONE = "America/Chicago"
NOMINATIM_URL = os.environ.get("NOMINATIM_URL", "https://nominatim.openstreetmap.org/search")
CENSUS_GEOCODER_URL = os.environ.get(
    "CENSUS_GEOCODER_URL",
    "https://geocoding.geo.census.gov/geocoder/locations/onelineaddress",
)
GOOGLE_MAPS_BROWSER_API_KEY_ENV = os.environ.get("GOOGLE_MAPS_BROWSER_API_KEY", "").strip()
GOOGLE_GEOCODING_API_KEY_ENV = os.environ.get("GOOGLE_GEOCODING_API_KEY", "").strip()
GOOGLE_GEOCODING_URL = os.environ.get("GOOGLE_GEOCODING_URL", "https://maps.googleapis.com/maps/api/geocode/json")
NOMINATIM_USER_AGENT = os.environ.get(
    "NOMINATIM_USER_AGENT",
    "PrasinosPowerInvoiceTool/1.0 (info@prasinospower.com)",
)
NOMINATIM_COUNTRY_CODES = os.environ.get("NOMINATIM_COUNTRY_CODES", "us,ca").strip()
GEOCODING_ENABLED = os.environ.get("GEOCODING_ENABLED", "true").strip().lower() in {"1", "true", "yes", "on"}
GEOCODER_VERSION = "5"
GOOGLE_PHOTOS_ALLOWED_HOSTS = {"photos.app.goo.gl", "photos.google.com", "www.photos.google.com"}
GOOGLE_PHOTOS_MAX_IMPORT = 1000
GOOGLE_PHOTOS_MAX_IMAGE_BYTES = 30 * 1024 * 1024
_geocode_lock = threading.Lock()
_last_geocode_request_at = 0.0
US_STATE_CODES = {
    "AL", "AK", "AZ", "AR", "CA", "CO", "CT", "DE", "FL", "GA", "HI", "ID", "IL", "IN", "IA",
    "KS", "KY", "LA", "ME", "MD", "MA", "MI", "MN", "MS", "MO", "MT", "NE", "NV", "NH", "NJ",
    "NM", "NY", "NC", "ND", "OH", "OK", "OR", "PA", "RI", "SC", "SD", "TN", "TX", "UT", "VT",
    "VA", "WA", "WV", "WI", "WY", "DC",
}

STATUS_LABELS = {
    "draft": "保存未提交",
    "submitted": "待经理审核",
    "returned": "已退回",
    "completed": "已完成",
    "void": "作废",
}

EXPENSE_STATUS_LABELS = {
    "draft": "保存未提交",
    "submitted": "待经理审核",
    "returned": "已退回",
    "approved": "已通过",
}

EXPENSE_PAYOUT_LABELS = {
    "pending": "待报销",
    "paid": "已报销",
}

CUSTOMER_REIMBURSEMENT_STATUS_LABELS = {
    "draft": "保存未提交",
    "submitted": "待经理审核",
    "returned": "已退回",
    "approved": "已通过",
}

ROLE_OPTIONS = {
    "admin": "管理员",
    "manager": "经理",
    "finance": "财务",
    "employee": "员工",
    "external_manager": "外部管理员",
    "external_employee": "外部员工",
}

MENU_PERMISSION_GROUPS = [
    {
        "label": "主菜单",
        "items": [
            {"key": "dashboard", "label": "概览", "roles": {"admin", "manager", "finance"}},
            {"key": "contracts", "label": "合同", "roles": {"admin", "manager", "finance", "external_manager"}},
            {"key": "service_orders", "label": "工单", "roles": set(ROLE_OPTIONS)},
            {"key": "service_order_map", "label": "站点地图", "roles": {"admin", "manager", "finance", "employee", "external_manager"}},
            {"key": "expense_processing", "label": "报销处理", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "new_invoice", "label": "新建发票", "roles": {"manager", "finance"}},
            {"key": "messages", "label": "消息", "roles": set(ROLE_OPTIONS)},
        ],
    },
    {
        "label": "报表",
        "items": [
            {"key": "invoices", "label": "发票查询", "roles": {"admin", "manager", "finance", "external_manager"}},
            {"key": "invoice_query", "label": "发票明细查询", "roles": {"admin", "manager", "finance", "external_manager"}},
            {"key": "customer_reimbursement_query", "label": "工单结算报表", "roles": {"admin", "manager", "finance", "external_manager"}},
            {"key": "buyer_query", "label": "站点查询", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "service_order_query", "label": "工单查询", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "service_report_query", "label": "日报查询", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "expense_query", "label": "报销明细查询", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "audit_log_report", "label": "操作日志", "roles": {"admin", "manager"}},
        ],
    },
    {
        "label": "薪酬管理",
        "items": [
            {"key": "labor_hours_report", "label": "工时统计", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "payroll_report", "label": "薪酬统计", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "payroll_calendar", "label": "薪酬日历", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "employee_grades", "label": "员工等级", "roles": {"admin", "manager", "finance"}},
            {"key": "payroll_subsidies", "label": "补贴参数", "roles": {"admin", "manager", "finance"}},
        ],
    },
    {
        "label": "基础数据",
        "items": [
            {"key": "clients", "label": "客户", "roles": {"admin", "manager", "finance", "external_manager"}},
            {"key": "owners", "label": "业主", "roles": {"admin", "manager", "finance"}},
            {"key": "buyers", "label": "站点", "roles": {"admin", "manager", "finance", "external_manager"}},
            {"key": "work_order_types", "label": "工单类型", "roles": {"admin", "manager"}},
            {"key": "projects", "label": "项目", "roles": {"admin", "manager"}},
            {"key": "countries", "label": "国家", "roles": {"admin", "manager"}},
            {"key": "company_info", "label": "公司信息", "roles": {"admin", "manager", "finance", "employee"}},
            {"key": "system_settings", "label": "系统设置", "roles": {"admin"}},
            {"key": "database_console", "label": "数据库工具", "roles": {"admin"}},
            {"key": "users", "label": "用户/我的资料", "roles": set(ROLE_OPTIONS)},
        ],
    },
]

DEFAULT_MENU_ROLES = {
    item["key"]: set(item["roles"])
    for group in MENU_PERMISSION_GROUPS
    for item in group["items"]
}

ACTION_LABELS = {
    "view": "查看",
    "create": "新增",
    "edit": "编辑",
    "delete": "删除",
    "approve": "审核",
    "export": "导出",
    "send": "发送",
    "pay": "核销",
    "execute": "执行",
}

ROLE_ACTION_PERMISSION_GROUPS = [
    {
        "label": "主业务",
        "items": [
            {"key": "contracts", "label": "合同", "actions": {"view": {"admin", "manager", "finance", "external_manager"}, "create": {"admin", "manager", "finance"}, "edit": {"admin", "manager", "finance"}, "delete": {"admin", "manager", "finance"}}},
            {"key": "service_orders", "label": "工单", "actions": {"view": set(ROLE_OPTIONS), "create": {"manager", "finance", "employee"}, "edit": {"admin", "manager", "finance", "employee", "external_manager", "external_employee"}, "delete": {"admin", "manager", "finance"}}},
            {"key": "service_reports", "label": "工作日报", "actions": {"view": set(ROLE_OPTIONS), "create": {"admin", "manager", "finance", "employee", "external_employee"}, "edit": {"admin", "manager", "finance", "employee", "external_employee"}, "delete": {"admin", "manager"}, "export": {"admin", "manager", "finance", "employee", "external_employee"}}},
            {"key": "invoices", "label": "发票", "actions": {"view": {"admin", "manager", "finance", "external_manager"}, "create": {"manager", "finance"}, "edit": {"admin", "manager", "finance"}, "delete": {"admin", "manager", "finance"}, "export": {"admin", "manager", "finance", "external_manager"}, "send": {"manager", "finance"}, "pay": {"admin", "manager", "finance"}}},
            {"key": "expenses", "label": "员工报销", "actions": {"view": {"admin", "manager", "finance", "employee"}, "create": {"manager", "finance", "employee"}, "edit": {"admin", "manager", "finance", "employee"}, "delete": {"admin", "manager", "finance", "employee"}, "approve": {"manager", "finance"}}},
            {"key": "customer_reimbursements", "label": "工单结算", "actions": {"view": {"admin", "manager", "finance", "external_manager"}, "create": {"admin", "manager", "finance"}, "edit": {"admin", "manager", "finance"}, "delete": {"admin", "manager", "finance"}, "approve": {"admin", "manager"}, "export": {"admin", "manager", "finance", "external_manager"}, "send": {"manager", "finance"}}},
        ],
    },
    {
        "label": "基础数据",
        "items": [
            {"key": "clients", "label": "客户", "actions": {"view": {"admin", "manager", "finance", "external_manager"}, "create": {"admin", "manager", "finance"}, "edit": {"admin", "manager", "finance"}, "delete": {"admin", "manager", "finance"}}},
            {"key": "owners", "label": "业主", "actions": {"view": {"admin", "manager", "finance"}, "create": {"admin", "manager", "finance"}, "edit": {"admin", "manager", "finance"}, "delete": {"admin", "manager", "finance"}}},
            {"key": "buyers", "label": "站点", "actions": {"view": {"admin", "manager", "finance", "external_manager"}, "create": {"admin", "manager", "finance", "external_manager"}, "edit": {"admin", "manager", "finance", "external_manager"}, "delete": {"admin", "manager", "finance"}}},
            {"key": "work_order_types", "label": "工单类型", "actions": {"view": {"admin", "manager"}, "create": {"admin", "manager"}, "edit": {"admin", "manager"}, "delete": {"admin", "manager"}}},
            {"key": "projects", "label": "项目", "actions": {"view": {"admin", "manager"}, "create": {"admin", "manager"}, "edit": {"admin", "manager"}, "delete": {"admin", "manager"}}},
            {"key": "countries", "label": "国家", "actions": {"view": {"admin", "manager"}, "create": {"admin", "manager"}, "edit": {"admin", "manager"}}},
            {"key": "users", "label": "用户", "actions": {"view": set(ROLE_OPTIONS), "create": {"admin", "external_manager"}, "edit": {"admin", "external_manager"}, "delete": {"admin"}, "approve": {"admin", "manager"}}},
        ],
    },
    {
        "label": "薪酬与系统",
        "items": [
            {"key": "labor_hours_report", "label": "工时统计", "actions": {"view": {"admin", "manager", "finance", "employee"}, "export": {"admin", "manager", "finance"}}},
            {"key": "payroll_report", "label": "薪酬统计", "actions": {"view": {"admin", "manager", "finance", "employee"}, "export": {"admin", "manager", "finance"}}},
            {"key": "payroll_calendar", "label": "薪酬日历", "actions": {"view": {"admin", "manager", "finance", "employee"}, "export": {"admin", "manager", "finance"}}},
            {"key": "employee_grades", "label": "员工等级", "actions": {"view": {"admin", "manager", "finance"}, "create": {"admin", "manager", "finance"}, "edit": {"admin", "manager", "finance"}}},
            {"key": "payroll_subsidies", "label": "补贴参数", "actions": {"view": {"admin", "manager", "finance"}, "edit": {"admin", "manager", "finance"}}},
            {"key": "company_info", "label": "公司信息", "actions": {"view": {"admin", "manager", "finance", "employee"}, "edit": {"admin"}}},
            {"key": "system_settings", "label": "系统设置", "actions": {"view": {"admin"}, "edit": {"admin"}}},
            {"key": "database_console", "label": "数据库工具", "actions": {"view": {"admin"}, "execute": {"admin"}}},
        ],
    },
]

DEFAULT_ACTION_ROLES = {
    (item["key"], action): set(roles)
    for group in ROLE_ACTION_PERMISSION_GROUPS
    for item in group["items"]
    for action, roles in item["actions"].items()
}

def permission_tree_groups():
    menu_labels = {
        item["key"]: item["label"]
        for group in MENU_PERMISSION_GROUPS
        for item in group["items"]
    }
    seen_keys = set()
    groups = []
    for group in ROLE_ACTION_PERMISSION_GROUPS:
        items = []
        for item in group["items"]:
            seen_keys.add(item["key"])
            items.append(
                {
                    "key": item["key"],
                    "label": item["label"],
                    "menu_key": item["key"] if item["key"] in DEFAULT_MENU_ROLES else "",
                    "actions": list(item["actions"].keys()),
                }
            )
        groups.append({"label": group["label"], "items": items})
    menu_only_items = [
        {"key": key, "label": label, "menu_key": key, "actions": []}
        for key, label in menu_labels.items()
        if key not in seen_keys
    ]
    if menu_only_items:
        groups.insert(0, {"label": "菜单入口", "items": menu_only_items})
    return groups

SUPPORTED_LANGUAGES = {
    "zh-CN": {"label": "简体中文", "flag": "🇨🇳"},
    "en": {"label": "English", "flag": "🇺🇸"},
    "nl": {"label": "Nederlands", "flag": "🇳🇱"},
    "de": {"label": "Deutsch", "flag": "🇩🇪"},
    "es": {"label": "Español", "flag": "🇪🇸"},
}
DEFAULT_LANGUAGE = "zh-CN"

PROJECT_TYPE_LABELS = {
    "invoice": "发票项目",
    "expense": "员工报销",
    "customer_expense": "工单结算",
}

CONTRACT_TYPE_LABELS = {
    "framework": "框架合同",
    "project": "项目合同",
}

CONTRACT_STATUS_LABELS = {
    "draft": "草稿",
    "review": "审核中",
    "signed": "已签署",
    "active": "执行中",
    "completed": "已完成",
    "terminated": "已终止",
}

CUSTOMER_REIMBURSEMENT_PROJECTS = {
    "标准工时": 70,
    "交通工时": 35,
    "加班工时": 105,
    "节假日工时": 140,
    "里程费": 1,
}

CUSTOMER_REIMBURSEMENT_INVOICE_PROJECTS = {
    "Technical Services": "labor_total",
    "Travel Expenses Reimbursement": "travel_total",
    "Mileage Reimbursement": "mileage_total",
}

DEFAULT_OWNER_NAMES = [
    "未知",
    "Qcells",
    "Stella",
    "Stem",
    "SunGrid",
    "Adapture Renewables",
    "Aypa",
    "spearmint",
    "plus power",
]

DEFAULT_SITE_OWNER_NAMES = {
    "edinburg": "Stella",
    "revolution": "spearmint",
    "ebony": "plus power",
    "anemoi": "plus power",
}


def normalized_project_name(value):
    value = str(value or "").translate({ord(char): None for char in "\u200b\u200c\u200d\ufeff"})
    return " ".join(value.strip().split())


def project_name_key(value):
    return normalized_project_name(value).casefold()


app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", secrets.token_hex(32))


def now():
    return datetime.now(app_timezone()).replace(microsecond=0).isoformat()


def db():
    if "db" not in g:
        os.makedirs(DATA_DIR, exist_ok=True)
        os.makedirs(ATTACHMENTS_DIR, exist_ok=True)
        os.makedirs(CONTRACT_ATTACHMENTS_DIR, exist_ok=True)
        os.makedirs(REPORT_ATTACHMENTS_DIR, exist_ok=True)
        os.makedirs(EXPENSE_ATTACHMENTS_DIR, exist_ok=True)
        os.makedirs(CUSTOMER_REIMBURSEMENT_DIR, exist_ok=True)
        os.makedirs(COMPANY_ATTACHMENT_DIR, exist_ok=True)
        os.makedirs(USER_ATTACHMENT_DIR, exist_ok=True)
        g.db = sqlite3.connect(DB_PATH)
        g.db.row_factory = sqlite3.Row
    return g.db


def merge_duplicate_projects(connection):
    grouped = {}
    for row in connection.execute("select * from projects order by is_active desc, id asc").fetchall():
        normalized_name = normalized_project_name(row["name"])
        key = project_name_key(normalized_name)
        if row["name"] != normalized_name or row["name_key"] != key:
            connection.execute(
                "update projects set name = ?, name_key = ? where id = ?",
                (normalized_name, key, row["id"]),
            )
        grouped.setdefault((row["project_type"], key), []).append(row)
    for rows in grouped.values():
        if len(rows) < 2:
            continue
        canonical = rows[0]
        canonical_name = normalized_project_name(canonical["name"])
        duplicate_ids = [row["id"] for row in rows[1:]]
        placeholders = ",".join("?" for _ in duplicate_ids)
        connection.execute(
            f"update invoice_items set project_id = ? where project_id in ({placeholders})",
            [canonical["id"], *duplicate_ids],
        )
        connection.execute(
            f"update expense_items set project_id = ?, project = ? where project_id in ({placeholders})",
            [canonical["id"], canonical_name, *duplicate_ids],
        )
        connection.execute(
            f"update expenses set project_id = ?, project = ? where project_id in ({placeholders})",
            [canonical["id"], canonical_name, *duplicate_ids],
        )
        connection.execute(f"delete from projects where id in ({placeholders})", duplicate_ids)


def project_name_exists(project_type, name, excluded_project_id=None):
    normalized_name_key = project_name_key(name)
    for row in db().execute("select id, name, name_key from projects where project_type = ?", (project_type,)).fetchall():
        if excluded_project_id and row["id"] == excluded_project_id:
            continue
        if (row["name_key"] or project_name_key(row["name"])) == normalized_name_key:
            return True
    return False


@app.teardown_appcontext
def close_db(error=None):
    connection = g.pop("db", None)
    if connection is not None:
        connection.close()


def ensure_column(connection, table, column, definition):
    columns = {row["name"] for row in connection.execute(f"pragma table_info({table})").fetchall()}
    if column not in columns:
        connection.execute(f"alter table {table} add column {column} {definition}")


def init_db():
    with app.app_context():
        connection = db()
        connection.executescript(
            """
            create table if not exists users (
                id integer primary key autoincrement,
                name text not null,
                email text not null unique,
                password_hash text not null,
                role text not null default 'user',
                address text not null default '',
                employee_grade_id integer,
                client_id integer,
                created_at text not null,
                foreign key(employee_grade_id) references employee_grades(id),
                foreign key(client_id) references clients(id)
            );

            create table if not exists employee_grades (
                id integer primary key autoincrement,
                grade_name text not null unique,
                base_salary real not null default 0,
                standard_hourly_rate real not null default 0,
                transport_hourly_rate real not null default 0,
                overtime_hourly_rate real not null default 0,
                holiday_hourly_rate real not null default 0,
                is_active integer not null default 1,
                created_at text not null
            );

            create table if not exists clients (
                id integer primary key autoincrement,
                client_number text not null unique,
                name text not null,
                short_name text not null,
                contact_name text,
                email text,
                address text,
                country text not null default 'China',
                created_at text not null
            );

            create table if not exists owners (
                id integer primary key autoincrement,
                owner_number text not null unique,
                name text not null,
                created_at text not null
            );

            create table if not exists projects (
                id integer primary key autoincrement,
                name text not null,
                name_key text,
                project_type text not null default 'invoice',
                default_amount real not null default 0,
                unit_price real not null default 0,
                tax_rate real not null default 0,
                is_active integer not null default 1,
                created_at text not null
            );

            create table if not exists contracts (
                id integer primary key autoincrement,
                contract_number text not null unique,
                client_id integer not null,
                contract_type text not null,
                title text not null,
                status text not null default 'draft',
                signed_date text,
                start_date text,
                end_date text,
                currency text not null default 'USD',
                amount real,
                payment_terms text,
                rate_card text,
                project_name text,
                notes text,
                created_by integer not null,
                created_at text not null,
                updated_at text not null,
                foreign key(client_id) references clients(id),
                foreign key(created_by) references users(id)
            );

            create table if not exists contract_attachments (
                id integer primary key autoincrement,
                contract_id integer not null,
                original_filename text not null,
                stored_filename text not null,
                content_type text,
                uploaded_by integer not null,
                uploaded_at text not null,
                foreign key(contract_id) references contracts(id) on delete cascade,
                foreign key(uploaded_by) references users(id)
            );

            create table if not exists invoices (
                id integer primary key autoincrement,
                invoice_number text not null unique,
                client_id integer not null,
                issue_date text not null,
                due_date text not null,
                currency text not null default 'USD',
                notes text,
                status text not null default 'submitted',
                return_reason text,
                sent_at text,
                paid_at text,
                payment_amount real,
                payment_note text,
                created_by integer not null,
                created_at text not null,
                foreign key(client_id) references clients(id),
                foreign key(created_by) references users(id)
            );

            create table if not exists invoice_items (
                id integer primary key autoincrement,
                invoice_id integer not null,
                project_id integer not null,
                description text not null,
                amount real not null,
                tax_rate real not null default 0,
                foreign key(invoice_id) references invoices(id) on delete cascade,
                foreign key(project_id) references projects(id)
            );

            create table if not exists invoice_attachments (
                id integer primary key autoincrement,
                invoice_id integer not null,
                original_filename text not null,
                stored_filename text not null,
                content_type text,
                uploaded_by integer not null,
                uploaded_at text not null,
                foreign key(invoice_id) references invoices(id) on delete cascade,
                foreign key(uploaded_by) references users(id)
            );

            create table if not exists invoice_save_tokens (
                token text primary key,
                invoice_id integer,
                created_at text not null,
                foreign key(invoice_id) references invoices(id) on delete cascade
            );

            create table if not exists messages (
                id integer primary key autoincrement,
                user_id integer not null,
                title text not null,
                body text not null,
                link text,
                is_read integer not null default 0,
                created_at text not null,
                foreign key(user_id) references users(id)
            );

            create table if not exists settings (
                key text primary key,
                value text not null
            );

            create table if not exists role_menu_permissions (
                role text not null,
                menu_key text not null,
                is_enabled integer not null default 1,
                updated_by integer,
                updated_at text not null,
                primary key(role, menu_key),
                foreign key(updated_by) references users(id) on delete set null
            );

            create table if not exists role_action_permissions (
                role text not null,
                resource_key text not null,
                action_key text not null,
                is_enabled integer not null default 1,
                updated_by integer,
                updated_at text not null,
                primary key(role, resource_key, action_key),
                foreign key(updated_by) references users(id) on delete set null
            );

            create table if not exists countries (
                code text primary key,
                region_code text not null,
                is_active integer not null default 1,
                sort_order integer not null default 0,
                created_at text not null
            );

            create table if not exists country_translations (
                country_code text not null,
                language_code text not null,
                name text not null,
                region_name text not null,
                primary key(country_code, language_code),
                foreign key(country_code) references countries(code) on delete cascade
            );

            create table if not exists company_attachments (
                id integer primary key autoincrement,
                original_filename text not null,
                stored_filename text not null,
                content_type text,
                uploaded_by integer not null,
                uploaded_at text not null,
                foreign key(uploaded_by) references users(id)
            );

            create table if not exists user_attachments (
                id integer primary key autoincrement,
                user_id integer not null,
                original_filename text not null,
                stored_filename text not null,
                content_type text,
                uploaded_by integer not null,
                uploaded_at text not null,
                foreign key(user_id) references users(id) on delete cascade,
                foreign key(uploaded_by) references users(id)
            );

            create table if not exists user_service_orders (
                user_id integer not null,
                service_order_id integer not null,
                assigned_by integer not null,
                assigned_at text not null,
                primary key(user_id, service_order_id),
                foreign key(user_id) references users(id) on delete cascade,
                foreign key(service_order_id) references service_orders(id) on delete cascade,
                foreign key(assigned_by) references users(id)
            );

            create table if not exists buyers (
                id integer primary key autoincrement,
                buyer_number text not null unique,
                country text,
                country_code text not null default 'US',
                name text not null,
                owner text,
                contact_name text,
                contact_details text,
                email text,
                site_size text,
                detailed_address text,
                equipment_manufacturer text,
                latitude real,
                longitude real,
                geocode_address text,
                geocode_status text not null default 'pending',
                geocode_attempted_at text,
                geocode_version text,
                created_at text not null
            );

            create table if not exists work_order_types (
                id integer primary key autoincrement,
                code text not null unique,
                name text not null,
                description text,
                is_active integer not null default 1,
                created_at text not null
            );

            create table if not exists service_orders (
                id integer primary key autoincrement,
                order_number text not null unique,
                client_id integer,
                client_name text not null,
                site_address text not null,
                client_order_number text not null,
                status text not null default 'open',
                created_by integer not null,
                created_at text not null,
                foreign key(client_id) references clients(id),
                foreign key(created_by) references users(id)
            );

            create table if not exists service_reports (
                id integer primary key autoincrement,
                service_order_id integer not null,
                report_date text not null,
                actual_work_date text not null,
                total_service_hours real not null default 0,
                travel_hours real not null default 0,
                public_transport_hours real not null default 0,
                driving_miles real not null default 0,
                departure_address text,
                site_address text,
                total_time text,
                cabinet_number text,
                arrival_time text,
                departure_time text,
                service_description text,
                created_by integer not null,
                created_at text not null,
                updated_at text not null,
                foreign key(service_order_id) references service_orders(id) on delete cascade,
                foreign key(created_by) references users(id)
            );

            create table if not exists service_report_workers (
                id integer primary key autoincrement,
                report_id integer not null,
                user_id integer not null,
                driving_miles real not null default 0,
                foreign key(report_id) references service_reports(id) on delete cascade,
                foreign key(user_id) references users(id)
            );

            create table if not exists audit_logs (
                id integer primary key autoincrement,
                user_id integer,
                user_name text not null,
                action text not null,
                entity_type text not null,
                entity_id integer,
                entity_label text not null,
                summary text,
                created_at text not null,
                foreign key(user_id) references users(id) on delete set null
            );

            create table if not exists service_report_saved_parts (
                id integer primary key autoincrement,
                report_id integer not null,
                part_number text,
                part_name text,
                quantity text,
                status text,
                sort_order integer not null default 0,
                foreign key(report_id) references service_reports(id) on delete cascade
            );

            create table if not exists service_report_replaced_parts (
                id integer primary key autoincrement,
                report_id integer not null,
                part_number text,
                part_name text,
                old_serial_number text,
                new_serial_number text,
                quantity text,
                sort_order integer not null default 0,
                foreign key(report_id) references service_reports(id) on delete cascade
            );

            create table if not exists service_report_attachments (
                id integer primary key autoincrement,
                report_id integer not null,
                category text not null,
                original_filename text not null,
                stored_filename text not null,
                content_type text,
                uploaded_by integer not null,
                uploaded_at text not null,
                foreign key(report_id) references service_reports(id) on delete cascade,
                foreign key(uploaded_by) references users(id)
            );

            create table if not exists service_report_save_tokens (
                token text primary key,
                report_id integer,
                created_at text not null,
                foreign key(report_id) references service_reports(id) on delete cascade
            );

            create table if not exists expenses (
                id integer primary key autoincrement,
                service_order_id integer not null,
                expense_number text not null unique,
                project_id integer,
                project text not null,
                expense_date text not null,
                amount real not null default 0,
                currency text not null default 'USD',
                description text,
                status text not null default 'draft',
                return_reason text,
                reviewed_by integer,
                reviewed_at text,
                created_by integer not null,
                created_at text not null,
                updated_at text not null,
                foreign key(service_order_id) references service_orders(id) on delete cascade,
                foreign key(project_id) references projects(id),
                foreign key(created_by) references users(id),
                foreign key(reviewed_by) references users(id)
            );

            create table if not exists expense_items (
                id integer primary key autoincrement,
                expense_id integer not null,
                project_id integer not null,
                project text not null,
                amount real not null default 0,
                description text,
                sort_order integer not null default 0,
                foreign key(expense_id) references expenses(id) on delete cascade,
                foreign key(project_id) references projects(id)
            );

            create table if not exists expense_save_tokens (
                token text primary key,
                expense_id integer,
                created_at text not null,
                foreign key(expense_id) references expenses(id) on delete cascade
            );

            create table if not exists expense_attachments (
                id integer primary key autoincrement,
                expense_id integer not null,
                original_filename text not null,
                stored_filename text not null,
                content_type text,
                uploaded_by integer not null,
                uploaded_at text not null,
                foreign key(expense_id) references expenses(id) on delete cascade,
                foreign key(uploaded_by) references users(id)
            );

            create table if not exists customer_reimbursements (
                id integer primary key autoincrement,
                service_order_id integer not null,
                file_name text not null,
                stored_filename text not null,
                status text not null default 'draft',
                return_reason text,
                reviewed_by integer,
                reviewed_at text,
                labor_total real not null default 0,
                lodging_total real not null default 0,
                travel_total real not null default 0,
                mileage_total real not null default 0,
                total_amount real not null default 0,
                invoice_id integer,
                created_by integer not null,
                created_at text not null,
                foreign key(service_order_id) references service_orders(id) on delete cascade,
                foreign key(invoice_id) references invoices(id) on delete set null,
                foreign key(created_by) references users(id),
                foreign key(reviewed_by) references users(id)
            );

            create table if not exists customer_reimbursement_items (
                id integer primary key autoincrement,
                customer_reimbursement_id integer not null,
                worker_name text not null,
                project_date text not null,
                standard_hours real not null default 0,
                transport_hours real not null default 0,
                overtime_hours real not null default 0,
                holiday_hours real not null default 0,
                standard_rate real not null default 0,
                transport_rate real not null default 0,
                overtime_rate real not null default 0,
                holiday_rate real not null default 0,
                labor_total real not null default 0,
                lodging real not null default 0,
                airfare real not null default 0,
                baggage real not null default 0,
                rental_car real not null default 0,
                fuel real not null default 0,
                parking real not null default 0,
                taxi real not null default 0,
                miles real not null default 0,
                mileage_rate real not null default 0,
                mileage_total real not null default 0,
                other real not null default 0,
                total real not null default 0,
                sort_order integer not null default 0,
                foreign key(customer_reimbursement_id) references customer_reimbursements(id) on delete cascade
            );

            create table if not exists customer_reimbursement_attachments (
                id integer primary key autoincrement,
                customer_reimbursement_id integer not null,
                original_filename text not null,
                stored_filename text not null,
                content_type text,
                uploaded_by integer not null,
                uploaded_at text not null,
                foreign key(customer_reimbursement_id) references customer_reimbursements(id) on delete cascade,
                foreign key(uploaded_by) references users(id)
            );
            """
        )
        ensure_column(connection, "invoices", "service_order_id", "integer")
        ensure_column(connection, "service_orders", "contract_id", "integer")
        ensure_column(connection, "users", "is_active", "integer not null default 1")
        ensure_column(connection, "users", "region_code", "text not null default 'americas'")
        ensure_column(connection, "users", "country_code", "text not null default 'US'")
        ensure_column(connection, "users", "employee_grade_id", "integer")
        ensure_column(connection, "users", "address", "text not null default ''")
        ensure_column(connection, "service_reports", "actual_work_date", "text")
        ensure_column(connection, "service_report_workers", "driving_miles", "real not null default 0")
        connection.execute(
            "update service_reports set actual_work_date = report_date where trim(coalesce(actual_work_date, '')) = ''"
        )
        connection.execute(
            """
            update service_reports
            set actual_work_date = '2026-06-24'
            where report_date = '2026-07-01'
              and (trim(coalesce(actual_work_date, '')) = '' or actual_work_date = report_date)
            """
        )
        connection.execute(
            """
            update service_reports
            set actual_work_date = '2026-06-25'
            where report_date = '2026-07-02'
              and (trim(coalesce(actual_work_date, '')) = '' or actual_work_date = report_date)
            """
        )
        ensure_column(connection, "projects", "project_type", "text not null default 'invoice'")
        ensure_column(connection, "projects", "name_key", "text")
        ensure_column(connection, "projects", "unit_price", "real not null default 0")
        ensure_column(connection, "customer_reimbursements", "status", "text not null default 'draft'")
        ensure_column(connection, "customer_reimbursements", "return_reason", "text")
        ensure_column(connection, "customer_reimbursements", "reviewed_by", "integer")
        ensure_column(connection, "customer_reimbursements", "reviewed_at", "text")
        ensure_column(connection, "expenses", "project_id", "integer")
        ensure_column(connection, "expenses", "payout_status", "text not null default 'pending'")
        ensure_column(connection, "expenses", "reimbursed_by", "integer")
        ensure_column(connection, "expenses", "reimbursed_at", "text")
        ensure_column(connection, "service_orders", "latitude", "real")
        ensure_column(connection, "service_orders", "longitude", "real")
        ensure_column(connection, "service_orders", "geocode_address", "text")
        ensure_column(connection, "service_orders", "geocode_status", "text not null default 'pending'")
        ensure_column(connection, "service_orders", "geocode_attempted_at", "text")
        ensure_column(connection, "service_orders", "geocode_version", "text")
        ensure_column(connection, "service_orders", "buyer_id", "integer")
        ensure_column(connection, "service_orders", "client_id", "integer")
        ensure_column(connection, "service_orders", "buyer_contact_name", "text")
        ensure_column(connection, "service_orders", "buyer_contact_details", "text")
        ensure_column(connection, "service_orders", "start_date", "text")
        ensure_column(connection, "service_orders", "work_order_type_id", "integer")
        ensure_column(connection, "service_orders", "region_code", "text not null default 'americas'")
        ensure_column(connection, "service_orders", "country_code", "text not null default 'US'")
        ensure_column(connection, "buyers", "client_id", "integer")
        ensure_column(connection, "buyers", "owner", "text")
        ensure_column(connection, "buyers", "owner_id", "integer")
        ensure_column(connection, "buyers", "email", "text")
        ensure_column(connection, "buyers", "site_size", "text")
        ensure_column(connection, "buyers", "country_code", "text not null default 'US'")
        site_country_migration = connection.execute(
            "select value from settings where key = ?",
            ("buyers_country_code_v1",),
        ).fetchone()
        if not site_country_migration:
            connection.execute("update buyers set country_code = 'US', country = 'US'")
            connection.execute(
                "update buyers set country_code = 'CA', country = 'CA' where buyer_number = 'BUY00017'"
            )
            connection.execute(
                "insert or replace into settings (key, value) values (?, ?)",
                ("buyers_country_code_v1", now()),
            )
        try:
            merge_duplicate_projects(connection)
        except sqlite3.Error:
            app.logger.exception("Failed to merge duplicate projects during startup.")
        connection.execute("drop index if exists idx_projects_type_name_unique")
        existing_owner_names = [
            row["owner"]
            for row in connection.execute(
                """
                select distinct trim(owner) as owner
                from buyers
                where trim(coalesce(owner, '')) != ''
                  and owner_id is null
                order by trim(owner)
                """
            ).fetchall()
        ]
        next_owner_sequence = connection.execute("select coalesce(max(id), 0) + 1 as value from owners").fetchone()["value"]
        for owner_name in DEFAULT_OWNER_NAMES:
            if connection.execute(
                "select 1 from owners where lower(trim(name)) = lower(trim(?))",
                (owner_name,),
            ).fetchone():
                continue
            while True:
                owner_number = f"OWN{next_owner_sequence:05d}"
                next_owner_sequence += 1
                if not connection.execute("select id from owners where owner_number = ?", (owner_number,)).fetchone():
                    break
            connection.execute(
                "insert into owners (owner_number, name, created_at) values (?, ?, ?)",
                (owner_number, owner_name, now()),
            )
        for owner_name in existing_owner_names:
            existing_owner = connection.execute(
                "select id from owners where lower(trim(name)) = lower(trim(?))",
                (owner_name,),
            ).fetchone()
            if existing_owner:
                owner_id = existing_owner["id"]
            else:
                while True:
                    owner_number = f"OWN{next_owner_sequence:05d}"
                    next_owner_sequence += 1
                    if not connection.execute("select id from owners where owner_number = ?", (owner_number,)).fetchone():
                        break
                cursor = connection.execute(
                    "insert into owners (owner_number, name, created_at) values (?, ?, ?)",
                    (owner_number, owner_name, now()),
                )
                owner_id = cursor.lastrowid
            connection.execute(
                "update buyers set owner_id = ? where owner_id is null and lower(trim(owner)) = lower(trim(?))",
                (owner_id, owner_name),
            )
        for site_name, owner_name in DEFAULT_SITE_OWNER_NAMES.items():
            owner = connection.execute(
                "select id, name from owners where lower(trim(name)) = lower(trim(?))",
                (owner_name,),
            ).fetchone()
            if owner:
                connection.execute(
                    """
                    update buyers
                    set owner_id = ?, owner = ?
                    where lower(trim(name)) = lower(trim(?))
                    """,
                    (owner["id"], owner["name"], site_name),
                )
        unknown_owner_row = connection.execute("select id, name from owners where name = ?", ("未知",)).fetchone()
        if unknown_owner_row:
            connection.execute(
                """
                update buyers
                set owner_id = ?, owner = ?
                where owner_id is null
                   or trim(coalesce(owner, '')) = ''
                """,
                (unknown_owner_row["id"], unknown_owner_row["name"]),
            )
        connection.execute(
            "update users set region_code = 'americas', country_code = 'US' "
            "where trim(coalesce(region_code, '')) = '' or trim(coalesce(country_code, '')) = ''"
        )
        connection.execute(
            "update service_orders set region_code = 'americas', country_code = 'US' "
            "where trim(coalesce(region_code, '')) = '' or trim(coalesce(country_code, '')) = ''"
        )
        connection.execute(
            """
            update buyers
            set client_id = (select id from clients where client_number = '00001' limit 1)
            where client_id is null
              and exists (select 1 from clients where client_number = '00001')
            """
        )
        connection.execute(
            """
            update service_orders
            set client_id = (
                select clients.id
                from clients
                where lower(trim(clients.name)) = lower(trim(service_orders.client_name))
                   or lower(trim(clients.short_name)) = lower(trim(service_orders.client_name))
                order by clients.id
                limit 1
            )
            where client_id is null
            """
        )
        connection.execute(
            """
            update service_orders
            set client_id = (
                select id from clients where client_number = '00001' limit 1
            )
            where client_id is null
              and exists (select 1 from clients where client_number = '00001')
            """
        )
        existing_buyer_names = {
            row["name"].strip().casefold(): row["id"]
            for row in connection.execute("select id, name from buyers").fetchall()
            if row["name"] and row["name"].strip()
        }
        next_buyer_sequence = connection.execute("select coalesce(max(id), 0) + 1 as value from buyers").fetchone()["value"]
        legacy_orders = connection.execute(
            """
            select id, client_name, site_address
            from service_orders
            where buyer_id is null and trim(coalesce(client_name, '')) != ''
            order by id
            """
        ).fetchall()
        for legacy_order in legacy_orders:
            buyer_name = legacy_order["client_name"].strip()
            buyer_id = existing_buyer_names.get(buyer_name.casefold())
            if not buyer_id:
                while True:
                    buyer_number = f"BUY{next_buyer_sequence:05d}"
                    next_buyer_sequence += 1
                    if not connection.execute(
                        "select id from buyers where buyer_number = ?",
                        (buyer_number,),
                    ).fetchone():
                        break
                cursor = connection.execute(
                    """
                    insert into buyers (
                        buyer_number, country, country_code, name, detailed_address, created_at
                    ) values (?, 'US', 'US', ?, ?, ?)
                    """,
                    (buyer_number, buyer_name, legacy_order["site_address"], now()),
                )
                buyer_id = cursor.lastrowid
                existing_buyer_names[buyer_name.casefold()] = buyer_id
            connection.execute(
                "update service_orders set buyer_id = ? where id = ?",
                (buyer_id, legacy_order["id"]),
            )
        connection.execute(
            """
            insert into expense_items (expense_id, project_id, project, amount, description, sort_order)
            select expenses.id, expenses.project_id, expenses.project, expenses.amount, expenses.description, 0
            from expenses
            where expenses.project_id is not null
              and not exists (
                  select 1 from expense_items where expense_items.expense_id = expenses.id
              )
            """
        )
        connection.execute(
            """
            delete from messages
            where title = '发票已确认完成'
              and exists (
                  select 1 from messages as reviewed
                  where reviewed.user_id = messages.user_id
                    and reviewed.link = messages.link
                    and reviewed.title = '发票已审核完成'
              )
            """
        )
        connection.execute(
            """
            update messages
            set title = '发票已审核完成',
                body = replace(body, '已由经理确认完成', '已审核完成')
            where title = '发票已确认完成'
            """
        )
        connection.execute(
            """
            update customer_reimbursements
            set status = 'approved'
            where invoice_id is not null and status != 'approved'
            """
        )
        connection.execute(
            """
            update invoices
            set status = 'completed', return_reason = null
            where status in ('submitted', 'returned')
            """
        )
        connection.execute("update users set role = 'external_manager' where role = 'external'")
        connection.execute(
            """
            update users
            set name = '陈永明'
            where name like '陈永%' and instr(name, '�') > 0
            """
        )
        connection.execute(
            """
            update customer_reimbursement_items
            set worker_name = '陈永明'
            where worker_name like '陈永%' and instr(worker_name, '�') > 0
            """
        )
        seed_customer_reimbursement_projects(connection)
        seed_countries(connection)
        seed_settings(connection)
        seed_role_permissions(connection)
        admin_email = os.environ.get("ADMIN_EMAIL", "admin@example.com").strip().lower()
        admin_password = os.environ.get("ADMIN_PASSWORD", "change-me-now")
        if not connection.execute("select id from users where email = ?", (admin_email,)).fetchone():
            connection.execute(
                """
                insert into users (name, email, password_hash, role, created_at)
                values (?, ?, ?, 'admin', ?)
                """,
                ("Admin", admin_email, generate_password_hash(admin_password), now()),
            )
        connection.commit()


def seed_countries(connection):
    countries = [
        (
            "US",
            "americas",
            10,
            {
                "zh-CN": ("美国", "美洲"),
                "en": ("United States", "Americas"),
                "nl": ("Verenigde Staten", "Amerika"),
                "de": ("Vereinigte Staaten", "Amerika"),
                "es": ("Estados Unidos", "América"),
            },
        ),
        (
            "CN",
            "asia",
            20,
            {
                "zh-CN": ("中国", "亚洲"),
                "en": ("China", "Asia"),
                "nl": ("China", "Azië"),
                "de": ("China", "Asien"),
                "es": ("China", "Asia"),
            },
        ),
        (
            "CA",
            "americas",
            25,
            {
                "zh-CN": ("加拿大", "美洲"),
                "en": ("Canada", "Americas"),
                "nl": ("Canada", "Amerika"),
                "de": ("Kanada", "Amerika"),
                "es": ("Canadá", "América"),
            },
        ),
        (
            "NL",
            "europe",
            30,
            {
                "zh-CN": ("荷兰", "欧洲"),
                "en": ("Netherlands", "Europe"),
                "nl": ("Nederland", "Europa"),
                "de": ("Niederlande", "Europa"),
                "es": ("Países Bajos", "Europa"),
            },
        ),
    ]
    for code, region_code, sort_order, translations in countries:
        connection.execute(
            """
            insert or ignore into countries (code, region_code, is_active, sort_order, created_at)
            values (?, ?, 1, ?, ?)
            """,
            (code, region_code, sort_order, now()),
        )
        for language_code, (name, region_name) in translations.items():
            connection.execute(
                """
                insert or ignore into country_translations
                    (country_code, language_code, name, region_name)
                values (?, ?, ?, ?)
                """,
                (code, language_code, name, region_name),
            )


def current_language():
    language = session.get("language", DEFAULT_LANGUAGE)
    return language if language in SUPPORTED_LANGUAGES else DEFAULT_LANGUAGE


def clear_session_preserving_language():
    language = current_language()
    session.clear()
    session["language"] = language


def country_rows(include_inactive=False):
    active_clause = "" if include_inactive else "where countries.is_active = 1"
    language = current_language()
    return db().execute(
        f"""
        select countries.*,
               coalesce(local.name, chinese.name, english.name, countries.code) as name,
               coalesce(local.region_name, chinese.region_name, english.region_name, countries.region_code) as region_name
        from countries
        left join country_translations local
          on local.country_code = countries.code and local.language_code = ?
        left join country_translations chinese
          on chinese.country_code = countries.code and chinese.language_code = 'zh-CN'
        left join country_translations english
          on english.country_code = countries.code and english.language_code = 'en'
        {active_clause}
        order by countries.sort_order, name, countries.code
        """,
        (language,),
    ).fetchall()


def country_by_code(code, include_inactive=False):
    clause = "" if include_inactive else "and is_active = 1"
    return db().execute(
        f"select * from countries where code = ? {clause}",
        ((code or "").strip().upper(),),
    ).fetchone()


def country_translations(country_code):
    return db().execute(
        """
        select * from country_translations
        where country_code = ?
        order by case language_code when 'zh-CN' then 0 when 'en' then 1 when 'nl' then 2 else 3 end,
                 language_code
        """,
        (country_code,),
    ).fetchall()


def country_from_form(default_code="US"):
    country = country_by_code(request.form.get("country_code") or default_code, include_inactive=True)
    if not country:
        country = country_by_code(default_code, include_inactive=True)
    return country


def report_location_filters(table_alias="service_orders"):
    region_code = request.args.get("region_code", "").strip().lower()
    country_code = request.args.get("country_code", "").strip().upper()
    countries = country_rows(include_inactive=True)
    valid_country_codes = {country["code"] for country in countries}
    valid_region_codes = {country["region_code"] for country in countries}
    if country_code not in valid_country_codes:
        country_code = ""
    if region_code not in valid_region_codes:
        region_code = ""
    clauses = []
    params = []
    if region_code:
        clauses.append(f"{table_alias}.region_code = ?")
        params.append(region_code)
    if country_code:
        clauses.append(f"{table_alias}.country_code = ?")
        params.append(country_code)
    regions = []
    seen_regions = set()
    for country in countries:
        if country["region_code"] in seen_regions:
            continue
        seen_regions.add(country["region_code"])
        regions.append({"code": country["region_code"], "name": country["region_name"]})
    return region_code, country_code, countries, regions, clauses, params


def seed_settings(connection):
    defaults = {}
    for key, value in DEFAULT_COMPANY_PROFILE.items():
        defaults[f"company_{key}"] = value
    for key, value in DEFAULT_PAYMENT_INSTRUCTIONS.items():
        defaults[f"payment_{key}"] = value
    for key, value in DEFAULT_SMTP_SETTINGS.items():
        defaults[f"smtp_{key}"] = value
    defaults["invoice_terms"] = DEFAULT_INVOICE_TERMS
    defaults["google_maps_browser_api_key"] = GOOGLE_MAPS_BROWSER_API_KEY_ENV
    defaults["google_geocoding_api_key"] = GOOGLE_GEOCODING_API_KEY_ENV
    defaults["payroll_cycle_start"] = "2026-07-06"
    defaults["payroll_car_allowance_method"] = "daily"
    defaults["payroll_car_daily_amount"] = "60"
    defaults["payroll_car_mileage_rate"] = "0.5"
    defaults["payroll_meal_allowance_method"] = "daily"
    defaults["payroll_meal_daily_amount"] = "50"
    defaults["payroll_lodging_limit"] = "0"
    defaults["payroll_historical_paid_date"] = date.today().isoformat()
    for key, value in defaults.items():
        connection.execute(
            "insert or ignore into settings (key, value) values (?, ?)",
            (key, value),
        )


def seed_role_permissions(connection):
    timestamp = now()
    for menu_key, roles in DEFAULT_MENU_ROLES.items():
        for role in ROLE_OPTIONS:
            connection.execute(
                """
                insert or ignore into role_menu_permissions (
                    role, menu_key, is_enabled, updated_by, updated_at
                ) values (?, ?, ?, null, ?)
                """,
                (role, menu_key, 1 if role in roles else 0, timestamp),
            )
    for (resource_key, action_key), roles in DEFAULT_ACTION_ROLES.items():
        for role in ROLE_OPTIONS:
            connection.execute(
                """
                insert or ignore into role_action_permissions (
                    role, resource_key, action_key, is_enabled, updated_by, updated_at
                ) values (?, ?, ?, ?, null, ?)
                """,
                (role, resource_key, action_key, 1 if role in roles else 0, timestamp),
            )


def get_setting(key, default=""):
    row = db().execute("select value from settings where key = ?", (key,)).fetchone()
    return row["value"] if row else default


def set_setting(key, value):
    db().execute(
        """
        insert into settings (key, value) values (?, ?)
        on conflict(key) do update set value = excluded.value
        """,
        (key, value),
    )


def setting_float(key, default=0):
    try:
        return float(get_setting(key, str(default)) or default)
    except (TypeError, ValueError):
        return float(default)


def payroll_subsidy_settings():
    car_method = get_setting("payroll_car_allowance_method", "daily")
    if car_method not in {"daily", "mileage"}:
        car_method = "daily"
    meal_method = get_setting("payroll_meal_allowance_method", "daily")
    if meal_method not in {"daily", "provided"}:
        meal_method = "daily"
    return {
        "cycle_start": get_setting("payroll_cycle_start", "2026-07-06"),
        "car_allowance_method": car_method,
        "car_daily_amount": setting_float("payroll_car_daily_amount", 60),
        "car_mileage_rate": setting_float("payroll_car_mileage_rate", 0.5),
        "meal_allowance_method": meal_method,
        "meal_daily_amount": setting_float("payroll_meal_daily_amount", 50),
        "lodging_limit": setting_float("payroll_lodging_limit", 0),
    }


def lodging_reimbursement_limit():
    return payroll_subsidy_settings()["lodging_limit"]


def is_lodging_project_name(name):
    normalized = (name or "").strip().casefold()
    return "住宿" in normalized or "lodging" in normalized or "hotel" in normalized


def is_fuel_project_name(name):
    normalized = (name or "").strip().casefold()
    return "油费" in normalized or "加油" in normalized or "fuel" in normalized or "gas" in normalized


def validate_lodging_reimbursement(amount, label="住宿报销"):
    limit = lodging_reimbursement_limit()
    if limit > 0 and amount > limit:
        raise ValueError(f"{label}不能超过 {limit:.2f}。")


def validate_fuel_reimbursement_allowed(amount, label="油费报销"):
    if amount > 0 and payroll_subsidy_settings()["car_allowance_method"] == "mileage":
        raise ValueError(f"当前车补按里程计费，{label}不可报销。")


def get_timezone_name():
    return get_setting("app_timezone", DEFAULT_TIMEZONE) or DEFAULT_TIMEZONE


def app_timezone():
    try:
        return ZoneInfo(get_timezone_name())
    except ZoneInfoNotFoundError:
        return ZoneInfo(DEFAULT_TIMEZONE)


def get_company_profile():
    return {key: get_setting(f"company_{key}", value) for key, value in DEFAULT_COMPANY_PROFILE.items()}


def get_payment_instructions():
    return {key: get_setting(f"payment_{key}", value) for key, value in DEFAULT_PAYMENT_INSTRUCTIONS.items()}


def get_invoice_terms():
    return get_setting("invoice_terms", DEFAULT_INVOICE_TERMS)


def get_smtp_settings():
    return {key: get_setting(f"smtp_{key}", value) for key, value in DEFAULT_SMTP_SETTINGS.items()}


def get_google_maps_browser_api_key():
    return get_setting("google_maps_browser_api_key", GOOGLE_MAPS_BROWSER_API_KEY_ENV).strip()


def get_google_geocoding_api_key():
    return get_setting("google_geocoding_api_key", GOOGLE_GEOCODING_API_KEY_ENV).strip()


def save_named_attachment(uploaded, directory, table, owner_column=None, owner_id=None):
    if not uploaded or not uploaded.filename:
        return
    if not allowed_attachment(uploaded.filename):
        raise ValueError(f"附件只支持 {ALLOWED_ATTACHMENT_LABEL}。")
    source_filename = uploaded.filename or "attachment"
    extension = source_filename.rsplit(".", 1)[1].lower()
    original_filename = os.path.basename(source_filename).strip() or f"attachment.{extension}"
    stored_filename = f"{secrets.token_hex(12)}.{extension}"
    os.makedirs(directory, exist_ok=True)
    uploaded.save(os.path.join(directory, stored_filename))
    if owner_column:
        db().execute(
            f"""
            insert into {table} (
                {owner_column}, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
            ) values (?, ?, ?, ?, ?, ?)
            """,
            (owner_id, original_filename, stored_filename, uploaded.content_type, g.user["id"], now()),
        )
    else:
        db().execute(
            f"""
            insert into {table} (
                original_filename, stored_filename, content_type, uploaded_by, uploaded_at
            ) values (?, ?, ?, ?, ?)
            """,
            (original_filename, stored_filename, uploaded.content_type, g.user["id"], now()),
        )


def get_company_attachments():
    return db().execute(
        """
        select company_attachments.*, users.name as uploader_name
        from company_attachments
        left join users on users.id = company_attachments.uploaded_by
        order by uploaded_at desc, id desc
        """
    ).fetchall()


def get_user_attachments(user_id):
    return db().execute(
        """
        select user_attachments.*, users.name as uploader_name
        from user_attachments
        left join users on users.id = user_attachments.uploaded_by
        where user_attachments.user_id = ?
        order by uploaded_at desc, user_attachments.id desc
        """,
        (user_id,),
    ).fetchall()


def assigned_service_order_ids(user_id):
    return {
        row["service_order_id"]
        for row in db().execute(
            "select service_order_id from user_service_orders where user_id = ?",
            (user_id,),
        ).fetchall()
    }


def save_user_service_order_assignments(user_id, role):
    db().execute("delete from user_service_orders where user_id = ?", (user_id,))
    if role != "external_employee":
        return
    order_ids = {int(value) for value in request.form.getlist("service_order_id") if value.isdigit()}
    for order_id in order_ids:
        order = db().execute("select client_id from service_orders where id = ?", (order_id,)).fetchone()
        if order and (not is_external_manager() or order["client_id"] == g.user["client_id"]):
            db().execute(
                """
                insert into user_service_orders (user_id, service_order_id, assigned_by, assigned_at)
                values (?, ?, ?, ?)
                """,
                (user_id, order_id, g.user["id"], now()),
            )


def current_user():
    user_id = session.get("user_id")
    if not user_id:
        return None
    user = db().execute(
        """
        select users.*, employee_grades.grade_name as employee_grade_name
        from users
        left join employee_grades on employee_grades.id = users.employee_grade_id
        where users.id = ?
        """,
        (user_id,),
    ).fetchone()
    if user and not user["is_active"]:
        clear_session_preserving_language()
        return None
    return user


@app.before_request
def load_user():
    requested_language = request.args.get("lang")
    if requested_language in SUPPORTED_LANGUAGES:
        session["language"] = requested_language
    g.user = current_user()
    if g.user:
        permission_rule = required_action_for_request()
        if permission_rule and not has_action_permission(*permission_rule):
            abort(403)


def login_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not g.user:
            return redirect(url_for("login", next=request.path))
        return view(*args, **kwargs)

    return wrapped


def admin_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not g.user:
            return redirect(url_for("login", next=request.path))
        if g.user["role"] != "admin":
            abort(403)
        return view(*args, **kwargs)

    return wrapped


def is_external_user():
    return g.user and normalized_role() in {"external_manager", "external_employee"}


def is_external_manager():
    return g.user and normalized_role() == "external_manager"


def is_external_employee():
    return g.user and normalized_role() == "external_employee"


def normalized_role(role=None):
    role = role if role is not None else (g.user["role"] if g.user else "")
    if role == "internal":
        return "finance"
    if role == "user":
        return "employee"
    if role == "external":
        return "external_manager"
    return role


def requires_user_address(role):
    return normalized_role(role) not in {"external_manager", "external_employee"}


def is_internal_user():
    return g.user and normalized_role() in {"admin", "manager", "finance", "employee"}


def is_manager():
    return g.user and normalized_role() in {"admin", "manager"}


def can_approve_users():
    return g.user and has_action_permission("users", "approve")


def can_view_invoices():
    return g.user and has_action_permission("invoices", "view")


def can_view_contracts():
    return g.user and has_action_permission("contracts", "view")


def can_manage_contracts():
    return g.user and (
        has_action_permission("contracts", "create")
        or has_action_permission("contracts", "edit")
        or has_action_permission("contracts", "delete")
    )


def can_create_invoice():
    return g.user and has_action_permission("invoices", "create")


def can_create_service_order():
    return g.user and has_action_permission("service_orders", "create")


def can_delete_service_order():
    return g.user and has_action_permission("service_orders", "delete")


def can_create_expense():
    return g.user and has_action_permission("expenses", "create")


def can_manage_customer_reimbursement():
    return g.user and (
        has_action_permission("customer_reimbursements", "create")
        or has_action_permission("customer_reimbursements", "edit")
        or has_action_permission("customer_reimbursements", "delete")
        or has_action_permission("customer_reimbursements", "approve")
    )


def can_manage_employee_grades():
    return g.user and (
        has_action_permission("employee_grades", "create")
        or has_action_permission("employee_grades", "edit")
        or has_action_permission("payroll_subsidies", "edit")
    )


def can_view_labor_payroll_reports():
    return g.user and (
        has_action_permission("labor_hours_report", "view")
        or has_action_permission("payroll_report", "view")
        or has_action_permission("payroll_calendar", "view")
    )


def can_view_customer_reimbursement():
    return g.user and has_action_permission("customer_reimbursements", "view")


def can_create_service_report(order=None):
    if not g.user:
        return False
    if not has_action_permission("service_reports", "create"):
        return False
    if not is_external_user():
        return True
    if is_external_employee() and order is not None:
        return can_access_service_order(order)
    return False


def can_manage_users():
    return g.user and not is_external_user() and (
        has_action_permission("users", "create")
        or has_action_permission("users", "edit")
        or has_action_permission("users", "delete")
    )


def can_manage_company_info():
    return g.user and has_action_permission("company_info", "edit")


def can_assign_external_employees():
    return can_manage_users() or is_external_manager()


def can_manage_user_record(user):
    if can_manage_users() or user["id"] == g.user["id"]:
        return True
    return (
        is_external_manager()
        and normalized_role(user["role"]) == "external_employee"
        and user["client_id"] == g.user["client_id"]
    )


def can_manage_buyers():
    return g.user and (
        has_action_permission("buyers", "view")
        or has_action_permission("buyers", "create")
        or has_action_permission("buyers", "edit")
        or has_action_permission("buyers", "delete")
    )


def can_manage_owners():
    return g.user and (
        has_action_permission("owners", "view")
        or has_action_permission("owners", "create")
        or has_action_permission("owners", "edit")
        or has_action_permission("owners", "delete")
    )


def can_access_buyer(buyer):
    if g.user and has_action_permission("buyers", "view") and not is_external_user():
        return True
    return is_external_manager() and bool(g.user["client_id"]) and buyer["client_id"] == g.user["client_id"]


def can_view_audit_logs():
    return g.user and has_action_permission("audit_logs", "view")


def can_use_database_console():
    return g.user and has_action_permission("database_console", "view")


def menu_permission_overrides():
    if not hasattr(g, "_menu_permission_overrides"):
        rows = db().execute("select role, menu_key, is_enabled from role_menu_permissions").fetchall()
        g._menu_permission_overrides = {
            (row["role"], row["menu_key"]): bool(row["is_enabled"])
            for row in rows
        }
    return g._menu_permission_overrides


def has_menu_permission(menu_key, role=None):
    if not g.user:
        return False
    role = normalized_role(role)
    if role not in ROLE_OPTIONS:
        return False
    overrides = menu_permission_overrides()
    override_key = (role, menu_key)
    if override_key in overrides:
        return overrides[override_key]
    return role in DEFAULT_MENU_ROLES.get(menu_key, set())


def action_permission_overrides():
    if not hasattr(g, "_action_permission_overrides"):
        rows = db().execute("select role, resource_key, action_key, is_enabled from role_action_permissions").fetchall()
        g._action_permission_overrides = {
            (row["role"], row["resource_key"], row["action_key"]): bool(row["is_enabled"])
            for row in rows
        }
    return g._action_permission_overrides


def has_action_permission(resource_key, action_key, role=None):
    if not g.user:
        return False
    role = normalized_role(role)
    if role not in ROLE_OPTIONS:
        return False
    overrides = action_permission_overrides()
    override_key = (role, resource_key, action_key)
    if override_key in overrides:
        return overrides[override_key]
    return role in DEFAULT_ACTION_ROLES.get((resource_key, action_key), set())


def required_action_for_request():
    endpoint = request.endpoint or ""
    method = request.method
    if endpoint == "edit_user" and request.view_args and request.view_args.get("user_id") == g.user["id"]:
        return None
    method_rules = {
        ("clients", "GET"): ("clients", "view"),
        ("clients", "POST"): ("clients", "create"),
        ("owners", "GET"): ("owners", "view"),
        ("owners", "POST"): ("owners", "create"),
        ("buyers", "GET"): ("buyers", "view"),
        ("buyers", "POST"): ("buyers", "create"),
        ("work_order_types", "GET"): ("work_order_types", "view"),
        ("work_order_types", "POST"): ("work_order_types", "create"),
        ("projects", "GET"): ("projects", "view"),
        ("projects", "POST"): ("projects", "create"),
        ("countries", "GET"): ("countries", "view"),
        ("countries", "POST"): ("countries", "edit"),
        ("employee_grades", "GET"): ("employee_grades", "view"),
        ("employee_grades", "POST"): ("employee_grades", "edit"),
        ("payroll_subsidies", "GET"): ("payroll_subsidies", "view"),
        ("payroll_subsidies", "POST"): ("payroll_subsidies", "edit"),
        ("users", "GET"): ("users", "view"),
        ("users", "POST"): ("users", "create"),
        ("system_settings", "GET"): ("system_settings", "view"),
        ("system_settings", "POST"): ("system_settings", "edit"),
        ("database_console", "GET"): ("database_console", "view"),
        ("database_console", "POST"): ("database_console", "execute"),
        ("company_info", "GET"): ("company_info", "view"),
        ("company_info", "POST"): ("company_info", "edit"),
        ("customer_reimbursement_form", "GET"): ("customer_reimbursements", "view"),
        ("customer_reimbursement_form", "POST"): ("customer_reimbursements", "edit"),
    }
    endpoint_rules = {
        "contracts": ("contracts", "view"),
        "new_contract": ("contracts", "create"),
        "contract_detail": ("contracts", "view"),
        "edit_contract": ("contracts", "edit"),
        "delete_contract": ("contracts", "delete"),
        "delete_contract_attachment": ("contracts", "delete"),
        "edit_client": ("clients", "edit"),
        "delete_client": ("clients", "delete"),
        "edit_owner": ("owners", "edit"),
        "delete_owner": ("owners", "delete"),
        "edit_buyer": ("buyers", "edit"),
        "delete_buyer": ("buyers", "delete"),
        "edit_work_order_type": ("work_order_types", "edit"),
        "delete_work_order_type": ("work_order_types", "delete"),
        "edit_project": ("projects", "edit"),
        "delete_project": ("projects", "delete"),
        "edit_user": ("users", "edit"),
        "update_user_status": ("users", "approve"),
        "delete_user": ("users", "delete"),
        "delete_user_attachment": ("users", "edit"),
        "delete_company_attachment": ("company_info", "edit"),
        "service_orders": ("service_orders", "view"),
        "new_service_order": ("service_orders", "create"),
        "edit_service_order": ("service_orders", "edit"),
        "service_order_detail": ("service_orders", "view"),
        "delete_service_order": ("service_orders", "delete"),
        "service_order_map": ("service_orders", "view"),
        "service_report_query": ("service_reports", "view"),
        "new_service_report": ("service_reports", "create"),
        "edit_service_report": ("service_reports", "edit"),
        "delete_service_report": ("service_reports", "delete"),
        "export_service_report": ("service_reports", "export"),
        "delete_report_attachment": ("service_reports", "edit"),
        "expense_processing": ("expenses", "view"),
        "process_expense_action": ("expenses", "approve"),
        "expense_query": ("expenses", "view"),
        "new_expense": ("expenses", "create"),
        "edit_expense": ("expenses", "edit"),
        "expense_detail": ("expenses", "view"),
        "approve_expense": ("expenses", "approve"),
        "return_expense": ("expenses", "approve"),
        "delete_expense": ("expenses", "delete"),
        "delete_expense_attachment": ("expenses", "edit"),
        "invoices": ("invoices", "view"),
        "invoice_query": ("invoices", "view"),
        "new_invoice": ("invoices", "create"),
        "edit_invoice": ("invoices", "edit"),
        "invoice_detail": ("invoices", "view"),
        "delete_invoice": ("invoices", "delete"),
        "send_invoice": ("invoices", "send"),
        "mark_invoice_paid": ("invoices", "pay"),
        "unmark_invoice_paid": ("invoices", "pay"),
        "admin_update_invoice_status": ("invoices", "edit"),
        "void_invoice": ("invoices", "edit"),
        "export_invoice": ("invoices", "export"),
        "export_invoice_pdf": ("invoices", "export"),
        "delete_attachment": ("invoices", "edit"),
        "customer_reimbursement_query": ("customer_reimbursements", "view"),
        "download_customer_reimbursement": ("customer_reimbursements", "export"),
        "preview_customer_reimbursement": ("customer_reimbursements", "export"),
        "approve_customer_reimbursement": ("customer_reimbursements", "approve"),
        "return_customer_reimbursement": ("customer_reimbursements", "approve"),
        "reset_customer_reimbursement": ("customer_reimbursements", "edit"),
        "delete_customer_reimbursement": ("customer_reimbursements", "delete"),
        "delete_customer_reimbursement_attachment": ("customer_reimbursements", "edit"),
        "labor_hours_report": ("labor_hours_report", "view"),
        "payroll_report": ("payroll_report", "view"),
        "payroll_calendar": ("payroll_calendar", "view"),
        "payroll_calendar_batch": ("payroll_calendar", "view"),
        "payroll_calendar_export": ("payroll_calendar", "export"),
    }
    return method_rules.get((endpoint, method)) or endpoint_rules.get(endpoint)


def can_access_client(client_id):
    if not g.user:
        return False
    if is_internal_user():
        return True
    return is_external_manager() and g.user["client_id"] == client_id


def require_invoice_access(invoice_id):
    if not can_view_invoices():
        abort(403)
    invoice = db().execute("select * from invoices where id = ?", (invoice_id,)).fetchone()
    if not invoice:
        abort(404)
    if not can_access_client(invoice["client_id"]):
        abort(403)
    return invoice


def require_contract_access(contract_id):
    if not can_view_contracts():
        abort(403)
    contract = db().execute("select * from contracts where id = ?", (contract_id,)).fetchone()
    if not contract:
        abort(404)
    if not can_access_client(contract["client_id"]):
        abort(403)
    return contract


def client_filter_clause(alias="invoices"):
    if is_external_manager():
        if not g.user["client_id"]:
            return "1 = 0", []
        return f"{alias}.client_id = ?", [g.user["client_id"]]
    if is_external_employee():
        return (
            f"{alias}.service_order_id in (select service_order_id from user_service_orders where user_id = ?)",
            [g.user["id"]],
        )
    return "1 = 1", []


def role_label(role):
    labels = {
        "admin": "管理员",
        "manager": "经理",
        "finance": "财务",
        "employee": "员工",
        "user": "员工",
        "internal": "财务",
        "external": "外部管理员",
        "external_manager": "外部管理员",
        "external_employee": "外部员工",
    }
    return labels.get(role, role)


def money(value, currency="USD"):
    symbols = {"USD": "$", "CNY": "¥", "EUR": "€", "GBP": "£", "JPY": "¥"}
    amount = float(value or 0)
    if currency == "JPY":
        return f"{symbols.get(currency, currency + ' ')}{amount:,.0f}"
    return f"{symbols.get(currency, currency + ' ')}{amount:,.2f}"


PROJECT_COLORS = ["#0f766e", "#175cd3", "#b42318", "#7a271a", "#6941c6", "#027a48", "#b54708", "#3538cd"]


def project_color(project_id):
    return PROJECT_COLORS[int(project_id or 0) % len(PROJECT_COLORS)]


def pdf_text(value):
    return "" if value is None else str(value)


def to_float(value, default=0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return float(default)


CENT = Decimal("0.01")


def decimal_value(value, default="0"):
    try:
        return Decimal(str(value if value is not None and value != "" else default))
    except (InvalidOperation, ValueError):
        return Decimal(str(default))


def money_decimal(value):
    amount = decimal_value(value)
    if abs(amount) < CENT:
        amount = Decimal("0")
    return amount.quantize(CENT, rounding=ROUND_HALF_UP)


def money_float(value):
    return float(money_decimal(value))


def invoice_totals(invoice_id):
    rows = db().execute("select amount, tax_rate from invoice_items where invoice_id = ?", (invoice_id,)).fetchall()
    subtotal = sum(float(row["amount"] or 0) for row in rows)
    tax = sum(float(row["amount"] or 0) * float(row["tax_rate"] or 0) / 100 for row in rows)
    return {"subtotal": subtotal, "tax": tax, "total": max(subtotal + tax, 0)}


def payment_label(invoice):
    if invoice["status"] == "void":
        return "不适用"
    if invoice["paid_at"]:
        return "已核销"
    if invoice["status"] == "completed":
        return "待核销"
    return "流程中"


def local_datetime(value):
    if not value:
        return ""
    try:
        parsed = datetime.fromisoformat(str(value))
    except ValueError:
        return str(value)[:16].replace("T", " ")
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(app_timezone()).strftime("%Y-%m-%d %H:%M")


def sql_statement_kind(sql):
    stripped = (sql or "").lstrip()
    if not stripped:
        return ""
    return stripped.split(None, 1)[0].lower()


def is_read_sql(sql):
    return sql_statement_kind(sql) in {"select", "with", "pragma", "explain"}


app.jinja_env.filters["money"] = money
app.jinja_env.filters["role_label"] = role_label
app.jinja_env.filters["payment_label"] = payment_label
app.jinja_env.filters["local_datetime"] = local_datetime
app.jinja_env.globals["can_view_invoices"] = can_view_invoices
app.jinja_env.globals["can_create_invoice"] = can_create_invoice
app.jinja_env.globals["can_view_contracts"] = can_view_contracts
app.jinja_env.globals["can_manage_contracts"] = can_manage_contracts
app.jinja_env.globals["can_create_service_order"] = can_create_service_order
app.jinja_env.globals["can_delete_service_order"] = can_delete_service_order
app.jinja_env.globals["can_create_expense"] = can_create_expense
app.jinja_env.globals["can_manage_customer_reimbursement"] = can_manage_customer_reimbursement
app.jinja_env.globals["can_manage_employee_grades"] = can_manage_employee_grades
app.jinja_env.globals["can_view_labor_payroll_reports"] = can_view_labor_payroll_reports
app.jinja_env.globals["can_view_customer_reimbursement"] = can_view_customer_reimbursement
app.jinja_env.globals["can_create_service_report"] = can_create_service_report
app.jinja_env.globals["is_external_manager"] = is_external_manager
app.jinja_env.globals["is_external_employee"] = is_external_employee
app.jinja_env.globals["can_assign_external_employees"] = can_assign_external_employees
app.jinja_env.globals["can_approve_users"] = can_approve_users
app.jinja_env.globals["can_view_audit_logs"] = can_view_audit_logs
app.jinja_env.globals["can_use_database_console"] = can_use_database_console
app.jinja_env.globals["has_menu_permission"] = has_menu_permission
app.jinja_env.globals["has_action_permission"] = has_action_permission
app.jinja_env.globals["normalized_role"] = normalized_role
app.jinja_env.globals["requires_user_address"] = requires_user_address
app.jinja_env.globals["expense_labels"] = EXPENSE_STATUS_LABELS
app.jinja_env.globals["expense_payout_labels"] = EXPENSE_PAYOUT_LABELS
app.jinja_env.globals["customer_reimbursement_labels"] = CUSTOMER_REIMBURSEMENT_STATUS_LABELS
app.jinja_env.globals["project_type_labels"] = PROJECT_TYPE_LABELS
app.jinja_env.globals["contract_type_labels"] = CONTRACT_TYPE_LABELS
app.jinja_env.globals["contract_status_labels"] = CONTRACT_STATUS_LABELS


def next_client_number():
    row = db().execute(
        """
        select client_number from clients
        where client_number glob '[0-9][0-9][0-9][0-9][0-9]'
        order by client_number desc limit 1
        """
    ).fetchone()
    return "00001" if not row else f"{int(row['client_number']) + 1:05d}"


def next_buyer_number():
    row = db().execute(
        """
        select buyer_number from buyers
        where buyer_number glob 'BUY[0-9][0-9][0-9][0-9][0-9]'
        order by buyer_number desc limit 1
        """
    ).fetchone()
    return "BUY00001" if not row else f"BUY{int(row['buyer_number'][3:]) + 1:05d}"


def next_owner_number():
    row = db().execute(
        """
        select owner_number from owners
        where owner_number glob 'OWN[0-9][0-9][0-9][0-9][0-9]'
        order by owner_number desc limit 1
        """
    ).fetchone()
    return "OWN00001" if not row else f"OWN{int(row['owner_number'][3:]) + 1:05d}"


def unknown_owner():
    owner = db().execute("select * from owners where name = ?", ("未知",)).fetchone()
    if owner:
        return owner
    cursor = db().execute(
        "insert into owners (owner_number, name, created_at) values (?, ?, ?)",
        (next_owner_number(), "未知", now()),
    )
    return db().execute("select * from owners where id = ?", (cursor.lastrowid,)).fetchone()


def owner_options():
    return db().execute(
        """
        select id, owner_number, name from owners
        order by case when name = '未知' then 0 else 1 end, owner_number
        """
    ).fetchall()


def next_invoice_number():
    prefix = f"PP-{date.today():%y%m}"
    alphabet = "ABCDEFGHJKLMNPQRSTUVWXYZ23456789"
    for _ in range(30):
        code = "".join(secrets.choice(alphabet) for _ in range(6))
        invoice_number = f"{prefix}-{code}"
        if not db().execute("select id from invoices where invoice_number = ?", (invoice_number,)).fetchone():
            return invoice_number
    raise RuntimeError("Unable to generate a unique invoice number.")


def next_contract_number():
    prefix = f"CT{date.today():%y}"
    row = db().execute(
        """
        select contract_number from contracts
        where contract_number like ?
        order by contract_number desc limit 1
        """,
        (f"{prefix}%",),
    ).fetchone()
    if not row:
        return f"{prefix}001"
    suffix = row["contract_number"][len(prefix):]
    try:
        return f"{prefix}{int(suffix) + 1:03d}"
    except ValueError:
        return f"{prefix}{secrets.randbelow(900) + 100}"


def contract_attachment_dir(contract_id):
    path = os.path.join(CONTRACT_ATTACHMENTS_DIR, str(contract_id))
    os.makedirs(path, exist_ok=True)
    return path


def contract_attachment_path(attachment):
    return os.path.join(
        CONTRACT_ATTACHMENTS_DIR,
        str(attachment["contract_id"]),
        attachment["stored_filename"],
    )


def save_contract_uploads(contract_id):
    existing = {
        row["original_filename"].strip().casefold()
        for row in db().execute(
            "select original_filename from contract_attachments where contract_id = ?",
            (contract_id,),
        ).fetchall()
    }
    for uploaded in request.files.getlist("attachments"):
        if not uploaded or not uploaded.filename:
            continue
        if not allowed_attachment(uploaded.filename):
            raise ValueError(f"附件只支持 {ALLOWED_ATTACHMENT_LABEL}。")
        original_filename = os.path.basename(uploaded.filename).strip()
        if original_filename.casefold() in existing:
            raise ValueError(f"附件“{original_filename}”已经存在。")
        extension = original_filename.rsplit(".", 1)[1].lower()
        stored_filename = f"{secrets.token_hex(12)}.{extension}"
        uploaded.save(os.path.join(contract_attachment_dir(contract_id), stored_filename))
        db().execute(
            """
            insert into contract_attachments (
                contract_id, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
            ) values (?, ?, ?, ?, ?, ?)
            """,
            (
                contract_id,
                original_filename,
                stored_filename,
                uploaded.content_type,
                g.user["id"],
                now(),
            ),
        )
        existing.add(original_filename.casefold())


def contract_form_values(contract=None):
    contract_type = request.form.get("contract_type", "framework")
    status = request.form.get("status", "draft")
    if contract_type not in CONTRACT_TYPE_LABELS:
        raise ValueError("请选择有效的合同类型。")
    if status not in CONTRACT_STATUS_LABELS:
        raise ValueError("请选择有效的合同状态。")
    try:
        client_id = int(request.form.get("client_id", ""))
    except (TypeError, ValueError):
        raise ValueError("请选择客户。")
    client = db().execute("select * from clients where id = ?", (client_id,)).fetchone()
    if not client or not can_access_client(client_id):
        raise ValueError("选择的客户不存在或无权访问。")
    contract_number = request.form.get("contract_number", "").strip() or (
        contract["contract_number"] if contract else next_contract_number()
    )
    title = request.form.get("title", "").strip()
    if not title:
        raise ValueError("请填写合同名称。")
    start_date = request.form.get("start_date") or None
    end_date = request.form.get("end_date") or None
    if start_date and end_date and end_date < start_date:
        raise ValueError("合同结束日期不能早于开始日期。")
    amount_text = request.form.get("amount", "").strip()
    amount = to_float(amount_text) if amount_text else None
    if amount is not None and amount < 0:
        raise ValueError("合同金额不能小于零。")
    project_name = request.form.get("project_name", "").strip()
    rate_card = request.form.get("rate_card", "").strip()
    if contract_type == "project":
        if not project_name:
            raise ValueError("项目合同必须填写项目名称。")
        if amount is None:
            raise ValueError("项目合同必须填写合同金额。")
        rate_card = ""
    else:
        project_name = ""
    return {
        "contract_number": contract_number,
        "client_id": client_id,
        "contract_type": contract_type,
        "title": title,
        "status": status,
        "signed_date": request.form.get("signed_date") or None,
        "start_date": start_date,
        "end_date": end_date,
        "currency": request.form.get("currency", "USD"),
        "amount": amount,
        "payment_terms": request.form.get("payment_terms", "").strip(),
        "rate_card": rate_card,
        "project_name": project_name,
        "notes": request.form.get("notes", "").strip(),
    }


def next_service_order_number():
    prefix = f"SO{date.today():%y%m}"
    row = db().execute(
        """
        select order_number from service_orders
        where order_number like ?
        order by order_number desc limit 1
        """,
        (f"{prefix}%",),
    ).fetchone()
    if not row:
        return f"{prefix}001"
    suffix = row["order_number"][len(prefix):]
    try:
        return f"{prefix}{int(suffix) + 1:03d}"
    except ValueError:
        return f"{prefix}{secrets.randbelow(900) + 100}"


def next_expense_number():
    prefix = f"EX{date.today():%y%m}"
    row = db().execute(
        """
        select expense_number from expenses
        where expense_number like ?
        order by expense_number desc limit 1
        """,
        (f"{prefix}%",),
    ).fetchone()
    if not row:
        return f"{prefix}001"
    suffix = row["expense_number"][len(prefix):]
    try:
        return f"{prefix}{int(suffix) + 1:03d}"
    except ValueError:
        return f"{prefix}{secrets.randbelow(900) + 100}"


def invoice_number_exists(invoice_number, exclude_invoice_id=None):
    if exclude_invoice_id:
        return db().execute(
            "select id from invoices where invoice_number = ? and id != ?",
            (invoice_number, exclude_invoice_id),
        ).fetchone() is not None
    return db().execute("select id from invoices where invoice_number = ?", (invoice_number,)).fetchone() is not None


def allowed_attachment(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_ATTACHMENT_EXTENSIONS


def normalized_attachment_filename(filename):
    return os.path.basename(filename or "").strip().casefold()


def invoice_attachment_dir(invoice_id):
    path = os.path.join(ATTACHMENTS_DIR, str(invoice_id))
    os.makedirs(path, exist_ok=True)
    return path


def attachment_file_path(attachment):
    return os.path.join(ATTACHMENTS_DIR, str(attachment["invoice_id"]), attachment["stored_filename"])


def invoice_attachment_path(invoice_id):
    return os.path.join(ATTACHMENTS_DIR, str(invoice_id))


def existing_attachment_names(invoice_id):
    rows = db().execute(
        "select original_filename from invoice_attachments where invoice_id = ?",
        (invoice_id,),
    ).fetchall()
    return {normalized_attachment_filename(row["original_filename"]) for row in rows}


def duplicate_attachment_names(uploads, invoice_id=None):
    seen = existing_attachment_names(invoice_id) if invoice_id else set()
    duplicates = []
    for uploaded in uploads:
        original_name = os.path.basename(uploaded.filename or "").strip()
        normalized_name = normalized_attachment_filename(original_name)
        if not normalized_name:
            continue
        if normalized_name in seen:
            duplicates.append(original_name)
            continue
        seen.add(normalized_name)
    return duplicates


def validate_attachment_uploads(uploads, invoice_id=None):
    for uploaded in uploads:
        if uploaded and uploaded.filename and not allowed_attachment(uploaded.filename):
            flash(f"附件只支持 {ALLOWED_ATTACHMENT_LABEL}。", "error")
            return False
    duplicates = duplicate_attachment_names(uploads, invoice_id)
    if duplicates:
        flash(f"附件重复：{', '.join(duplicates)}。请删除重复附件后再上传。", "error")
        return False
    return True


def save_uploaded_attachment(invoice_id, uploaded):
    source_filename = uploaded.filename or "attachment"
    extension = source_filename.rsplit(".", 1)[1].lower()
    original_filename = os.path.basename(source_filename).strip() or f"attachment.{extension}"
    if "." not in original_filename:
        original_filename = f"{original_filename}.{extension}"
    if normalized_attachment_filename(original_filename) in existing_attachment_names(invoice_id):
        raise RuntimeError(f"附件已存在：{original_filename}")
    stored_filename = f"{secrets.token_hex(12)}.{extension}"
    uploaded.save(os.path.join(invoice_attachment_dir(invoice_id), stored_filename))
    db().execute(
        """
        insert into invoice_attachments (
            invoice_id, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
        ) values (?, ?, ?, ?, ?, ?)
        """,
        (invoice_id, original_filename, stored_filename, uploaded.content_type, g.user["id"], now()),
    )


def unique_invoice_attachment_filename(invoice_id, original_filename):
    original_filename = os.path.basename(original_filename or "").strip() or "attachment"
    stem, extension = os.path.splitext(original_filename)
    stem = stem or "attachment"
    existing = existing_attachment_names(invoice_id)
    candidate = original_filename
    counter = 2
    while normalized_attachment_filename(candidate) in existing:
        candidate = f"{stem}-{counter}{extension}"
        counter += 1
    return candidate


def copy_file_to_invoice_attachment(invoice_id, source_path, original_filename, content_type=None, uploaded_by=None):
    if not os.path.isfile(source_path):
        return False
    original_filename = unique_invoice_attachment_filename(invoice_id, original_filename)
    extension = os.path.splitext(original_filename)[1].lstrip(".").lower() or "dat"
    stored_filename = f"{secrets.token_hex(12)}.{extension}"
    shutil.copyfile(source_path, os.path.join(invoice_attachment_dir(invoice_id), stored_filename))
    db().execute(
        """
        insert into invoice_attachments (
            invoice_id, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
        ) values (?, ?, ?, ?, ?, ?)
        """,
        (invoice_id, original_filename, stored_filename, content_type, uploaded_by or g.user["id"], now()),
    )
    return True


def copy_mileage_proofs_to_invoice(invoice_id, service_order_id):
    rows = db().execute(
        """
        select service_report_attachments.*, service_reports.report_date
        from service_report_attachments
        join service_reports on service_reports.id = service_report_attachments.report_id
        where service_reports.service_order_id = ?
          and service_report_attachments.category = 'mileage_proof'
        order by service_reports.report_date asc, service_report_attachments.uploaded_at asc,
                 service_report_attachments.id asc
        """,
        (service_order_id,),
    ).fetchall()
    copied = 0
    for row in rows:
        original_name = row["original_filename"] or "mileage-proof"
        report_date = (row["report_date"] or "").strip()
        if report_date:
            original_name = f"里程佐证-{report_date}-{original_name}"
        if copy_file_to_invoice_attachment(
            invoice_id,
            report_attachment_path(row),
            original_name,
            row["content_type"],
            row["uploaded_by"],
        ):
            copied += 1
    return copied


def uploaded_attachments_from_request():
    return [
        file
        for file in request.files.getlist("attachments") + request.files.getlist("attachment")
        if file and file.filename
    ]


def get_invoice_attachments(invoice_id):
    return db().execute(
        """
        select invoice_attachments.*, users.name as uploader_name
        from invoice_attachments left join users on users.id = invoice_attachments.uploaded_by
        where invoice_id = ?
        order by uploaded_at desc
        """,
        (invoice_id,),
    ).fetchall()


def can_access_service_order(order):
    if not g.user:
        return False
    if is_internal_user():
        return True
    if is_external_manager():
        return bool(g.user["client_id"]) and order["client_id"] == g.user["client_id"]
    if is_external_employee():
        return db().execute(
            "select 1 from user_service_orders where user_id = ? and service_order_id = ?",
            (g.user["id"], order["id"]),
        ).fetchone() is not None
    return False


def require_service_order(order_id):
    order = db().execute(
        """
        select service_orders.*, work_order_types.name as work_order_type_name,
               coalesce(buyer_country_local.name, buyer_country_zh.name, buyer_country_en.name, buyers.country, buyers.country_code) as buyer_country,
               coalesce(owners.name, buyers.owner) as buyer_owner,
               buyers.email as buyer_email,
               buyers.site_size as buyer_site_size,
               buyers.equipment_manufacturer as buyer_equipment_manufacturer,
               clients.name as billing_client_name,
               clients.email as billing_client_email,
               contracts.contract_number,
               contracts.title as contract_title,
               contracts.contract_type,
               coalesce(country_local.name, country_zh.name, country_en.name, service_orders.country_code) as country_name,
               coalesce(country_local.region_name, country_zh.region_name, country_en.region_name, service_orders.region_code) as region_name
        from service_orders
        left join work_order_types on work_order_types.id = service_orders.work_order_type_id
        left join buyers on buyers.id = service_orders.buyer_id
        left join owners on owners.id = buyers.owner_id
        left join clients on clients.id = service_orders.client_id
        left join contracts on contracts.id = service_orders.contract_id
        left join country_translations buyer_country_local
          on buyer_country_local.country_code = buyers.country_code and buyer_country_local.language_code = ?
        left join country_translations buyer_country_zh
          on buyer_country_zh.country_code = buyers.country_code and buyer_country_zh.language_code = 'zh-CN'
        left join country_translations buyer_country_en
          on buyer_country_en.country_code = buyers.country_code and buyer_country_en.language_code = 'en'
        left join country_translations country_local
          on country_local.country_code = service_orders.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = service_orders.country_code and country_zh.language_code = 'zh-CN'
        left join country_translations country_en
          on country_en.country_code = service_orders.country_code and country_en.language_code = 'en'
        where service_orders.id = ?
        """,
        (current_language(), current_language(), order_id),
    ).fetchone()
    if not order:
        abort(404)
    if not can_access_service_order(order):
        abort(403)
    return order


def service_order_dependency_counts(order_id):
    return {
        "reports": db().execute(
            "select count(*) as count from service_reports where service_order_id = ?",
            (order_id,),
        ).fetchone()["count"],
        "invoices": db().execute(
            "select count(*) as count from invoices where service_order_id = ? and status != 'void'",
            (order_id,),
        ).fetchone()["count"],
        "customer_reimbursements": db().execute(
            "select count(*) as count from customer_reimbursements where service_order_id = ?",
            (order_id,),
        ).fetchone()["count"],
        "expenses": db().execute(
            "select count(*) as count from expenses where service_order_id = ?",
            (order_id,),
        ).fetchone()["count"],
    }


def service_order_delete_blockers(order_id):
    counts = service_order_dependency_counts(order_id)
    labels = {
        "reports": "工作日报",
        "invoices": "发票",
        "customer_reimbursements": "工单结算",
        "expenses": "员工报销",
    }
    return [labels[key] for key, count in counts.items() if count]


def require_service_order_start_date(order):
    if order["start_date"]:
        return None
    flash("请先编辑工单并维护开始日期，再新增发票、工作日报或报销。", "error")
    return redirect(url_for("edit_service_order", order_id=order["id"]))


def require_service_report(report_id):
    report = db().execute("select * from service_reports where id = ?", (report_id,)).fetchone()
    if not report:
        abort(404)
    order = require_service_order(report["service_order_id"])
    return report, order


def allowed_image(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS


def report_storage_context(report_id):
    row = db().execute(
        """
        select service_reports.report_date, service_orders.order_number
        from service_reports
        join service_orders on service_orders.id = service_reports.service_order_id
        where service_reports.id = ?
        """,
        (report_id,),
    ).fetchone()
    if not row:
        abort(404)
    order_number = secure_filename(row["order_number"]) or f"SO-{report_id}"
    try:
        report_date = datetime.strptime(row["report_date"], "%Y-%m-%d").strftime("%Y%m%d")
    except (TypeError, ValueError):
        report_date = str(row["report_date"] or "").replace("-", "") or "unknown-date"
    return order_number, report_date


def report_attachment_relative_path(report_id, category, stored_filename):
    order_number, report_date = report_storage_context(report_id)
    category_folder = REPORT_PHOTO_FOLDERS.get(category)
    if not category_folder:
        raise ValueError("未知的日报附件类别。")
    return Path(order_number, report_date, category_folder, os.path.basename(stored_filename)).as_posix()


def report_attachment_dir(report_id, category):
    relative = report_attachment_relative_path(report_id, category, "placeholder.jpg")
    path = os.path.join(REPORT_ATTACHMENTS_DIR, os.path.dirname(relative))
    os.makedirs(path, exist_ok=True)
    return path


def report_attachment_path(attachment):
    stored_filename = str(attachment["stored_filename"])
    stored_path = Path(stored_filename)
    if len(stored_path.parts) > 1:
        return os.path.join(REPORT_ATTACHMENTS_DIR, *stored_path.parts)
    return os.path.join(REPORT_ATTACHMENTS_DIR, str(attachment["report_id"]), stored_filename)


def prune_empty_report_folders(path):
    root = Path(REPORT_ATTACHMENTS_DIR).resolve()
    current = Path(path).resolve().parent
    while current != root:
        try:
            current.relative_to(root)
            current.rmdir()
        except (OSError, ValueError):
            break
        current = current.parent


def relocate_report_attachments(report_id):
    attachments = db().execute(
        "select * from service_report_attachments where report_id = ?",
        (report_id,),
    ).fetchall()
    for attachment in attachments:
        old_path = report_attachment_path(attachment)
        stored_filename = os.path.basename(attachment["stored_filename"])
        relative_path = report_attachment_relative_path(report_id, attachment["category"], stored_filename)
        new_path = os.path.join(REPORT_ATTACHMENTS_DIR, *Path(relative_path).parts)
        if os.path.normcase(os.path.abspath(old_path)) != os.path.normcase(os.path.abspath(new_path)):
            os.makedirs(os.path.dirname(new_path), exist_ok=True)
            if os.path.isfile(old_path):
                shutil.move(old_path, new_path)
                prune_empty_report_folders(old_path)
            db().execute(
                "update service_report_attachments set stored_filename = ? where id = ?",
                (relative_path, attachment["id"]),
            )


def uploaded_report_files(field_name):
    return [file for file in request.files.getlist(field_name) if file and file.filename]


def compress_report_image(source, target_path):
    compress_image(source, target_path)


def file_sha256(path):
    digest = hashlib.sha256()
    with open(path, "rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalized_report_photo_name(filename):
    name = os.path.basename(filename or "").strip()
    stem, _ = os.path.splitext(name)
    stem = re.sub(r"\s*\(\d+\)$", "", stem).strip()
    return stem.casefold()


def is_duplicate_report_photo(report_id, category, original_filename, candidate_hash):
    candidate_name = normalized_report_photo_name(original_filename)
    rows = db().execute(
        """
        select * from service_report_attachments
        where report_id = ? and category = ?
        """,
        (report_id, category),
    ).fetchall()
    for row in rows:
        existing_path = report_attachment_path(row)
        if os.path.isfile(existing_path):
            try:
                if file_sha256(existing_path) == candidate_hash:
                    return True
            except OSError:
                pass
        elif candidate_name and normalized_report_photo_name(row["original_filename"]) == candidate_name:
            return True
    return False


def save_report_attachment(report_id, uploaded, category):
    if not uploaded or not uploaded.filename:
        return False
    if not allowed_image(uploaded.filename):
        raise ValueError("日报照片仅支持 PNG、JPG、JPEG、WEBP、GIF。")
    source_filename = uploaded.filename or "photo"
    extension = source_filename.rsplit(".", 1)[1].lower()
    original_filename = os.path.basename(source_filename).strip() or f"photo.{extension}"
    target_dir = report_attachment_dir(report_id, category)
    temporary_filename = f".{secrets.token_hex(12)}.jpg.tmp"
    temporary_path = os.path.join(target_dir, temporary_filename)
    compress_report_image(uploaded.stream, temporary_path)
    candidate_hash = file_sha256(temporary_path)
    if is_duplicate_report_photo(report_id, category, original_filename, candidate_hash):
        os.remove(temporary_path)
        return False
    image_filename = f"{secrets.token_hex(12)}.jpg"
    stored_filename = report_attachment_relative_path(report_id, category, image_filename)
    os.replace(temporary_path, os.path.join(target_dir, image_filename))
    content_type = "image/jpeg"
    db().execute(
        """
        insert into service_report_attachments (
            report_id, category, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
        ) values (?, ?, ?, ?, ?, ?, ?)
        """,
        (report_id, category, original_filename, stored_filename, content_type, g.user["id"], now()),
    )
    return True


def save_report_file_attachment(report_id, uploaded, category):
    if not uploaded or not uploaded.filename:
        return False
    if not allowed_attachment(uploaded.filename):
        raise ValueError(f"里程佐证只支持 {ALLOWED_ATTACHMENT_LABEL}。")
    source_filename = uploaded.filename or "attachment"
    extension = source_filename.rsplit(".", 1)[1].lower()
    original_filename = os.path.basename(source_filename).strip() or f"attachment.{extension}"
    if "." not in original_filename:
        original_filename = f"{original_filename}.{extension}"
    target_dir = report_attachment_dir(report_id, category)
    temporary_filename = f".{secrets.token_hex(12)}.{extension}.tmp"
    temporary_path = os.path.join(target_dir, temporary_filename)
    uploaded.save(temporary_path)
    candidate_hash = file_sha256(temporary_path)
    if is_duplicate_report_photo(report_id, category, original_filename, candidate_hash):
        os.remove(temporary_path)
        return False
    stored_basename = f"{secrets.token_hex(12)}.{extension}"
    stored_filename = report_attachment_relative_path(report_id, category, stored_basename)
    os.replace(temporary_path, os.path.join(target_dir, stored_basename))
    db().execute(
        """
        insert into service_report_attachments (
            report_id, category, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
        ) values (?, ?, ?, ?, ?, ?, ?)
        """,
        (report_id, category, original_filename, stored_filename, uploaded.content_type, g.user["id"], now()),
    )
    return True


def shared_photos_root():
    return Path(SHARED_PHOTOS_DIR).resolve()


def ensure_service_order_picture_folder(order_number):
    folder = shared_photos_root() / secure_filename(order_number) / "pictures"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def service_order_incoming_folder(order_number, source_name="uploads"):
    folder = shared_photos_root() / secure_filename(order_number) / "incoming" / secure_filename(source_name)
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def is_google_photos_share_url(value):
    parsed = urlsplit((value or "").strip())
    host = parsed.netloc.lower().split(":")[0]
    return parsed.scheme in {"http", "https"} and host in GOOGLE_PHOTOS_ALLOWED_HOSTS


def google_photos_page_html(share_url):
    request_headers = {
        "User-Agent": "Mozilla/5.0 (compatible; PrasinosPowerInvoiceTool/1.0)",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }
    try:
        with urlopen(Request(share_url, headers=request_headers), timeout=30) as response:
            content_type = response.headers.get_content_type()
            if content_type and "html" not in content_type:
                raise ValueError("链接返回的不是 Google Photos 分享页面。")
            return response.read(12 * 1024 * 1024).decode("utf-8", errors="ignore")
    except HTTPError as error:
        if error.code in {401, 403}:
            raise ValueError("这个 Google Photos 分享链接需要登录或没有访问权限，服务器无法直接导入。") from error
        raise ValueError(f"读取 Google Photos 分享链接失败：HTTP {error.code}") from error
    except URLError as error:
        raise ValueError(f"读取 Google Photos 分享链接失败：{error.reason}") from error


def normalize_embedded_google_url(value):
    text = html.unescape(value or "")
    replacements = {
        "\\/": "/",
        "\\u003d": "=",
        "\\u0026": "&",
        "\\u003f": "?",
        "\\u0025": "%",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def google_image_score(url):
    size_match = re.search(r"=w(\d+)-h(\d+)", url)
    if size_match:
        return int(size_match.group(1)) * int(size_match.group(2))
    square_match = re.search(r"=s(\d+)", url)
    if square_match:
        size = int(square_match.group(1))
        return size * size
    return len(url)


def google_image_dimensions(url):
    size_match = re.search(r"=w(\d+)-h(\d+)", url)
    if size_match:
        return int(size_match.group(1)), int(size_match.group(2))
    square_match = re.search(r"=s(\d+)", url)
    if square_match:
        size = int(square_match.group(1))
        return size, size
    return None


def is_likely_google_account_image(url):
    parsed = urlsplit(url)
    lower = url.casefold()
    account_keywords = (
        "avatar",
        "profile",
        "userphoto",
        "photo.jpg",
        "account",
        "person",
        "people",
        "anonymous",
        "googleusercontent.com/a/",
    )
    if any(keyword in lower for keyword in account_keywords):
        return True
    dimensions = google_image_dimensions(url)
    if dimensions:
        width, height = dimensions
        if width <= 300 and height <= 300:
            return True
    return False


def google_photo_url_key(url):
    return url.split("=", 1)[0]


def google_photo_download_candidates(url):
    base = google_photo_url_key(url)
    candidates = [
        f"{base}=d",
        f"{base}=s0",
        f"{base}=s0-d",
        url,
    ]
    return list(dict.fromkeys(candidates))


def extract_google_photo_urls(page_html):
    normalized = normalize_embedded_google_url(page_html)
    matches = re.findall(r"https://lh3\.googleusercontent\.com/[^\s\"'<>\\\])]+", normalized)
    by_photo = {}
    for raw_url in matches:
        url = raw_url.rstrip(".,;")
        if is_likely_google_account_image(url):
            continue
        photo_key = google_photo_url_key(url)
        if not photo_key:
            continue
        if photo_key not in by_photo or google_image_score(url) > google_image_score(by_photo[photo_key]):
            by_photo[photo_key] = url
    urls = list(by_photo.values())
    return urls[:GOOGLE_PHOTOS_MAX_IMPORT], len(urls)


def downloaded_image_extension(content_type, url):
    mapping = {
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/webp": ".webp",
        "image/gif": ".gif",
        "image/heic": ".heic",
        "image/heif": ".heif",
    }
    if content_type in mapping:
        return mapping[content_type]
    suffix = Path(urlsplit(url).path).suffix.lower()
    if suffix.lstrip(".") in ALLOWED_IMAGE_EXTENSIONS:
        return suffix
    return ".jpg"


def unique_download_path(folder, filename):
    candidate = folder / secure_filename(filename)
    if not candidate.exists():
        return candidate
    counter = 2
    while True:
        next_candidate = candidate.with_name(f"{candidate.stem}-{counter}{candidate.suffix}")
        if not next_candidate.exists():
            return next_candidate
        counter += 1


def download_google_photo_image(url, target_dir, index):
    request_headers = {
        "User-Agent": "Mozilla/5.0 (compatible; PrasinosPowerInvoiceTool/1.0)",
        "Accept": "image/avif,image/webp,image/apng,image/*,*/*;q=0.8",
    }
    last_error = None
    for candidate_url in google_photo_download_candidates(url):
        try:
            with urlopen(Request(candidate_url, headers=request_headers), timeout=45) as response:
                content_type = response.headers.get_content_type()
                if not (content_type or "").startswith("image/"):
                    raise ValueError(f"不是图片内容：{content_type or 'unknown'}")
                extension = downloaded_image_extension(content_type, candidate_url)
                target_path = unique_download_path(target_dir, f"google-photo-{index:03d}{extension}")
                total = 0
                digest = hashlib.sha256()
                with open(target_path, "wb") as file:
                    while True:
                        chunk = response.read(1024 * 1024)
                        if not chunk:
                            break
                        total += len(chunk)
                        if total > GOOGLE_PHOTOS_MAX_IMAGE_BYTES:
                            file.close()
                            try:
                                target_path.unlink()
                            except FileNotFoundError:
                                pass
                            raise ValueError("图片超过 30MB，已跳过。")
                        digest.update(chunk)
                        file.write(chunk)
                return target_path, digest.hexdigest(), total
        except ValueError as error:
            if "超过 30MB" in str(error):
                raise
            last_error = error
        except (HTTPError, URLError, OSError) as error:
            last_error = error
    raise ValueError(f"无法下载原图：{last_error}")


def import_google_photos_share_to_incoming(order, share_url):
    if not is_google_photos_share_url(share_url):
        raise ValueError("请输入 Google Photos 分享链接。")
    page_html = google_photos_page_html(share_url)
    image_urls, total_found = extract_google_photo_urls(page_html)
    if not image_urls:
        raise ValueError("没有从这个分享链接中识别到可下载照片。链接可能需要登录，或 Google 页面结构发生变化。")
    target_dir = service_order_incoming_folder(order["order_number"], "google-photos")
    imported = []
    skipped = []
    seen_hashes = set()
    for index, image_url in enumerate(image_urls, start=1):
        try:
            path, digest, size = download_google_photo_image(image_url, target_dir, index)
            if digest in seen_hashes:
                path.unlink(missing_ok=True)
                skipped.append(f"第 {index} 张：重复照片")
                continue
            seen_hashes.add(digest)
            imported.append({"path": path, "size": size})
        except Exception as error:
            skipped.append(f"第 {index} 张：{error}")
    if not imported and skipped:
        raise ValueError("照片下载失败：" + "；".join(skipped[:5]))
    truncated_count = max(0, total_found - len(image_urls))
    return imported, skipped, target_dir, total_found, truncated_count


def send_google_photos_import_result(user, order, share_url, imported, skipped, target_dir, total_found=0, truncated_count=0, error=None):
    success_count = len(imported or [])
    skipped_count = len(skipped or [])
    if error:
        title = f"Google Photos 导入失败：{order['order_number']}"
        lines = [
            f"工单：{order['order_number']}",
            f"链接：{share_url}",
            f"失败原因：{error}",
        ]
    else:
        title = f"Google Photos 导入完成：{order['order_number']}"
        lines = [
            f"工单：{order['order_number']}",
            f"链接：{share_url}",
            f"识别到照片：{total_found} 张",
            f"尝试导入：{success_count + skipped_count} 张",
            f"导入照片：{success_count} 张",
            f"跳过/失败：{skipped_count} 张",
            f"保存位置：{target_dir}",
            "照片已放入 incoming/google-photos，photo_worker 会自动处理到 pictures。",
        ]
        if truncated_count:
            lines.append(f"注意：识别到的照片超过系统单次导入上限 {GOOGLE_PHOTOS_MAX_IMPORT} 张，仍有 {truncated_count} 张未导入。")
        if skipped:
            lines.append("前几条跳过/失败信息：")
            lines.extend(skipped[:10])
    body = "\n".join(lines)
    create_message(user["id"], title, body, f"/service-orders/{order['id']}")
    email = (user["email"] or "").strip()
    if email:
        try:
            send_email(
                to=email,
                subject=title,
                html=f"<p>{html.escape(body).replace(chr(10), '<br>')}</p>",
            )
        except Exception:
            app.logger.exception("Unable to send Google Photos import result email to %s", email)


def run_google_photos_import_job(order_id, user_id, share_url):
    with app.app_context():
        order = db().execute("select * from service_orders where id = ?", (order_id,)).fetchone()
        user = db().execute("select * from users where id = ?", (user_id,)).fetchone()
        if not order or not user:
            return
        try:
            imported, skipped, target_dir, total_found, truncated_count = import_google_photos_share_to_incoming(order, share_url)
            send_google_photos_import_result(user, order, share_url, imported, skipped, target_dir, total_found, truncated_count)
            db().commit()
        except Exception as error:
            db().rollback()
            try:
                send_google_photos_import_result(user, order, share_url, [], [], "", error=str(error))
                db().commit()
            except Exception:
                db().rollback()
                app.logger.exception("Unable to notify Google Photos import result for order %s", order["order_number"])


def start_google_photos_import_job(order_id, user_id, share_url):
    thread = threading.Thread(
        target=run_google_photos_import_job,
        args=(order_id, user_id, share_url),
        daemon=True,
    )
    thread.start()


def resolve_shared_photo(relative_path="", require_file=False, allow_missing=False):
    root = shared_photos_root()
    candidate = (root / str(relative_path or "")).resolve()
    try:
        candidate.relative_to(root)
    except ValueError:
        abort(403)
    if allow_missing and not candidate.exists():
        return candidate
    if require_file:
        if not candidate.is_file() or candidate.suffix.lower().lstrip(".") not in ALLOWED_IMAGE_EXTENSIONS:
            abort(404)
    elif not candidate.is_dir():
        abort(404)
    return candidate


def shared_photo_relative(path):
    return path.resolve().relative_to(shared_photos_root()).as_posix()


def resolve_shared_picture(relative_path):
    source_path = resolve_shared_photo(relative_path, require_file=True)
    relative_parts = source_path.relative_to(shared_photos_root()).parts
    if len(relative_parts) < 3 or relative_parts[1].casefold() != "pictures":
        abort(403)
    return source_path, relative_parts


def count_shared_images(path, excluded_dirs=None):
    if not path.is_dir():
        return 0
    excluded = {name.casefold() for name in (excluded_dirs or set())}
    return sum(
        1
        for entry in path.rglob("*")
        if entry.is_file()
        and not entry.name.startswith(".")
        and entry.suffix.lower().lstrip(".") in ALLOWED_IMAGE_EXTENSIONS
        and not ({part.casefold() for part in entry.relative_to(path).parts} & ({"@eadir"} | excluded))
        and valid_image_file(entry)
    )


def valid_image_file(path):
    try:
        with Image.open(path) as image:
            image.verify()
        return True
    except (OSError, ValueError):
        return False


def order_photo_status(order_dir):
    state_names = {"incoming", "pictures", "thumbnails", "processing", "failed", "original_backup"}
    waiting = count_shared_images(order_dir / "incoming")
    if order_dir.is_dir():
        waiting += sum(
            1
            for entry in order_dir.rglob("*")
            if entry.is_file()
            and not entry.name.startswith(".")
            and entry.suffix.lower().lstrip(".") in ALLOWED_IMAGE_EXTENSIONS
            and entry.relative_to(order_dir).parts[0].casefold() not in state_names
            and "@eadir" not in {part.casefold() for part in entry.relative_to(order_dir).parts}
        )
    return {
        "waiting": waiting,
        "processing": count_shared_images(order_dir / "processing"),
        "completed": count_shared_images(order_dir / "pictures"),
        "failed": count_shared_images(order_dir / "failed", excluded_dirs={"orphaned_pictures"}),
    }


def seed_customer_reimbursement_projects(connection):
    for name, unit_price in CUSTOMER_REIMBURSEMENT_PROJECTS.items():
        row = connection.execute(
            "select id, project_type from projects where name_key = ? and project_type = 'customer_expense'",
            (project_name_key(name),),
        ).fetchone()
        if row:
            connection.execute(
                """
                update projects
                set name = ?, name_key = ?, unit_price = ?, is_active = 1
                where id = ?
                """,
                (name, project_name_key(name), unit_price, row["id"]),
            )
            continue
        connection.execute(
            """
            insert into projects (name, name_key, project_type, default_amount, unit_price, tax_rate, is_active, created_at)
            values (?, ?, 'customer_expense', 0, ?, 0, 1, ?)
            """,
            (name, project_name_key(name), unit_price, now()),
        )
    for name in CUSTOMER_REIMBURSEMENT_INVOICE_PROJECTS:
        row = connection.execute(
            "select id from projects where name_key = ? and project_type = 'invoice'",
            (project_name_key(name),),
        ).fetchone()
        if row:
            continue
        connection.execute(
            """
            insert into projects (name, name_key, project_type, default_amount, unit_price, tax_rate, is_active, created_at)
            values (?, ?, 'invoice', 0, 0, 0, 1, ?)
            """,
            (name, project_name_key(name), now()),
        )


def save_shared_report_photo(report_id, relative_path, category):
    source_path = resolve_shared_photo(relative_path, require_file=True)
    order_number, _ = report_storage_context(report_id)
    processed_root = (shared_photos_root() / order_number / "pictures").resolve()
    try:
        source_path.relative_to(processed_root)
    except ValueError:
        raise ValueError("只能选择已完成处理的 NAS 照片。")
    if source_path.suffix.lower() not in {".jpg", ".jpeg"}:
        raise ValueError("NAS 照片尚未完成处理。")
    if is_duplicate_report_photo(report_id, category, source_path.name, file_sha256(source_path)):
        return False
    image_filename = f"{secrets.token_hex(12)}.jpg"
    stored_filename = report_attachment_relative_path(report_id, category, image_filename)
    shutil.copyfile(source_path, os.path.join(report_attachment_dir(report_id, category), image_filename))
    db().execute(
        """
        insert into service_report_attachments (
            report_id, category, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
        ) values (?, ?, ?, ?, 'image/jpeg', ?, ?)
        """,
        (report_id, category, source_path.name, stored_filename, g.user["id"], now()),
    )
    return True


def save_report_uploads(report_id):
    for field_name, category in (
        ("self_check_photos", "self_check"),
        ("site_photos", "site"),
        ("arrival_photos", "arrival"),
        ("departure_photos", "departure"),
    ):
        for uploaded in uploaded_report_files(field_name):
            save_report_attachment(report_id, uploaded, category)
        for relative_path in request.form.getlist(f"shared_photo_{category}"):
            save_shared_report_photo(report_id, relative_path, category)
    for uploaded in uploaded_report_files("mileage_proof_attachments"):
        save_report_file_attachment(report_id, uploaded, "mileage_proof")


def claim_report_save_token(token, report_id=None):
    token = str(token or "").strip()
    if not token:
        raise ValueError("保存令牌无效，请刷新页面后重试。")
    cursor = db().execute(
        """
        insert or ignore into service_report_save_tokens (token, report_id, created_at)
        values (?, ?, ?)
        """,
        (token, report_id, now()),
    )
    return cursor.rowcount == 1


def finish_report_save_token(token, report_id):
    db().execute(
        "update service_report_save_tokens set report_id = ? where token = ?",
        (report_id, token),
    )


def claim_invoice_save_token(token, invoice_id=None):
    token = str(token or "").strip()
    if not token:
        raise ValueError("保存令牌无效，请刷新页面后重试。")
    cursor = db().execute(
        """
        insert or ignore into invoice_save_tokens (token, invoice_id, created_at)
        values (?, ?, ?)
        """,
        (token, invoice_id, now()),
    )
    return cursor.rowcount == 1


def finish_invoice_save_token(token, invoice_id):
    db().execute(
        "update invoice_save_tokens set invoice_id = ? where token = ?",
        (invoice_id, token),
    )


def claim_expense_save_token(token, expense_id=None):
    token = str(token or "").strip()
    if not token:
        raise ValueError("保存令牌无效，请刷新页面后重试。")
    cursor = db().execute(
        """
        insert or ignore into expense_save_tokens (token, expense_id, created_at)
        values (?, ?, ?)
        """,
        (token, expense_id, now()),
    )
    return cursor.rowcount == 1


def finish_expense_save_token(token, expense_id):
    db().execute(
        "update expense_save_tokens set expense_id = ? where token = ?",
        (expense_id, token),
    )


def posted_report_time(prefix):
    hour = request.form.get(f"{prefix}_hour", "").strip()
    minute = request.form.get(f"{prefix}_minute", "").strip()
    if not hour and not minute:
        return request.form.get(prefix, "").strip()
    try:
        hour_value = int(hour)
        minute_value = int(minute)
    except (TypeError, ValueError):
        raise ValueError("请选择有效的现场时间。")
    if not 0 <= hour_value <= 23 or not 0 <= minute_value <= 59:
        raise ValueError("请选择有效的现场时间。")
    return f"{hour_value:02d}:{minute_value:02d}"


def report_worker_ids_from_form():
    return list(dict.fromkeys(value for value in request.form.getlist("worker_user_id") if value))


def get_report_attachments(report_id):
    rows = db().execute(
        """
        select service_report_attachments.*, users.name as uploader_name
        from service_report_attachments left join users on users.id = service_report_attachments.uploaded_by
        where report_id = ?
        order by category asc, uploaded_at asc, id asc
        """,
        (report_id,),
    ).fetchall()
    grouped = {"self_check": [], "site": [], "arrival": [], "departure": [], "mileage_proof": []}
    for row in rows:
        grouped.setdefault(row["category"], []).append(row)
    return grouped


def customer_reimbursement_dir(reimbursement_id):
    path = os.path.join(CUSTOMER_REIMBURSEMENT_DIR, str(reimbursement_id))
    os.makedirs(path, exist_ok=True)
    return path


def customer_reimbursement_attachment_dir(reimbursement_id):
    path = os.path.join(customer_reimbursement_dir(reimbursement_id), "attachments")
    os.makedirs(path, exist_ok=True)
    return path


def customer_reimbursement_file_path(reimbursement):
    return os.path.join(CUSTOMER_REIMBURSEMENT_DIR, str(reimbursement["id"]), reimbursement["stored_filename"])


def customer_reimbursement_attachment_path(attachment):
    return os.path.join(CUSTOMER_REIMBURSEMENT_DIR, str(attachment["customer_reimbursement_id"]), "attachments", attachment["stored_filename"])


def require_customer_reimbursement(reimbursement_id):
    if not can_view_customer_reimbursement():
        abort(403)
    reimbursement = db().execute("select * from customer_reimbursements where id = ?", (reimbursement_id,)).fetchone()
    if not reimbursement:
        abort(404)
    order = require_service_order(reimbursement["service_order_id"])
    return reimbursement, order


def customer_reimbursement_rates():
    rows = db().execute(
        """
        select name, unit_price
        from projects
        where project_type = 'customer_expense' and is_active = 1
        """
    ).fetchall()
    rates = {name: 0.0 for name in CUSTOMER_REIMBURSEMENT_PROJECTS}
    rates.update({row["name"]: float(row["unit_price"] or 0) for row in rows})
    return rates


def employee_grade_options(include_inactive=False):
    active_clause = "" if include_inactive else "where is_active = 1"
    return db().execute(
        f"""
        select *
        from employee_grades
        {active_clause}
        order by is_active desc, grade_name
        """
    ).fetchall()


def parse_report_minutes(value):
    if not value or ":" not in str(value):
        return None
    try:
        hour_text, minute_text = str(value).split(":", 1)
        return int(hour_text) * 60 + int(minute_text)
    except ValueError:
        return None


def rounded_report_service_hours(arrival_time, departure_time):
    arrival = parse_report_minutes(arrival_time)
    departure = parse_report_minutes(departure_time)
    if arrival is None or departure is None or departure <= arrival:
        return 0
    duration_minutes = departure - arrival
    rounded_minutes = (duration_minutes // 15) * 15
    if duration_minutes % 15 > 7:
        rounded_minutes += 15
    return round(rounded_minutes / 60, 2)


def format_report_hours(value):
    return f"{float(value):.2f}".rstrip("0").rstrip(".")


def calculated_report_total_time():
    return format_report_hours(
        to_float(request.form.get("travel_hours")) + to_float(request.form.get("public_transport_hours"))
    )


def calculated_report_total_service_hours(arrival_time, departure_time):
    return rounded_report_service_hours(arrival_time, departure_time) * len(report_worker_ids_from_form())


def observed_us_holidays(year):
    def observed(day):
        if day.weekday() == 5:
            return day - timedelta(days=1)
        if day.weekday() == 6:
            return day + timedelta(days=1)
        return day

    def nth_weekday(month, weekday, nth):
        current = date(year, month, 1)
        while current.weekday() != weekday:
            current += timedelta(days=1)
        return current + timedelta(days=7 * (nth - 1))

    def last_weekday(month, weekday):
        current = date(year, month + 1, 1) - timedelta(days=1) if month < 12 else date(year, 12, 31)
        while current.weekday() != weekday:
            current -= timedelta(days=1)
        return current

    return {
        observed(date(year, 1, 1)),
        nth_weekday(1, 0, 3),
        nth_weekday(2, 0, 3),
        last_weekday(5, 0),
        observed(date(year, 6, 19)),
        observed(date(year, 7, 4)),
        nth_weekday(9, 0, 1),
        nth_weekday(10, 0, 2),
        observed(date(year, 11, 11)),
        nth_weekday(11, 3, 4),
        observed(date(year, 12, 25)),
    }


def is_us_weekend_or_holiday(day):
    return day.weekday() >= 5 or day in observed_us_holidays(day.year)


def report_duration_hours(report, worker_count=1):
    rounded_hours = rounded_report_service_hours(report["arrival_time"], report["departure_time"])
    if rounded_hours:
        return rounded_hours
    fallback_hours = float(report["total_service_hours"] or 0)
    return round(fallback_hours / max(worker_count, 1), 2)


def report_actual_date(report):
    return report["actual_work_date"] or report["report_date"]


def split_report_labor_hours(report, worker_count=1):
    work_hours = report_duration_hours(report, worker_count)
    transport_hours = float(report["public_transport_hours"] or 0)
    try:
        work_day = date.fromisoformat(report_actual_date(report))
    except (TypeError, ValueError):
        work_day = None
    if work_day and is_us_weekend_or_holiday(work_day):
        return {
            "standard_hours": 0,
            "transport_hours": transport_hours,
            "overtime_hours": 0,
            "holiday_hours": work_hours,
        }
    return {
        "standard_hours": min(work_hours, 8),
        "transport_hours": transport_hours,
        "overtime_hours": max(work_hours - 8, 0),
        "holiday_hours": 0,
    }


def customer_reimbursement_seed_rows(order_id):
    rates = customer_reimbursement_rates()
    reports = db().execute(
        """
        select service_reports.*
        from service_reports
        where service_reports.service_order_id = ?
        order by service_reports.report_date asc, service_reports.id asc
        """,
        (order_id,),
    ).fetchall()
    rows = []
    for report in reports:
        workers = db().execute(
            """
            select users.name
            from service_report_workers
            join users on users.id = service_report_workers.user_id
            where service_report_workers.report_id = ?
            order by users.name
            """,
            (report["id"],),
        ).fetchall()
        if not workers:
            continue
        labor_hours = split_report_labor_hours(report, len(workers))
        for worker in workers:
            row = {
                "worker_name": worker["name"],
                "project_date": report_actual_date(report),
                "standard_hours": labor_hours["standard_hours"],
                "transport_hours": labor_hours["transport_hours"],
                "overtime_hours": labor_hours["overtime_hours"],
                "holiday_hours": labor_hours["holiday_hours"],
                "standard_rate": rates["标准工时"],
                "transport_rate": rates["交通工时"],
                "overtime_rate": rates["加班工时"],
                "holiday_rate": rates["节假日工时"],
                "lodging": 0,
                "airfare": 0,
                "baggage": 0,
                "rental_car": 0,
                "fuel": 0,
                "parking": 0,
                "taxi": 0,
                "miles": 0,
                "mileage_rate": rates["里程费"],
                "other": 0,
            }
            rows.append(calculate_customer_reimbursement_item(row, len(rows)))
    return rows


def calculate_customer_reimbursement_item(row, sort_order=0):
    labor_total = (
        decimal_value(row.get("standard_hours")) * money_decimal(row.get("standard_rate"))
        + decimal_value(row.get("transport_hours")) * money_decimal(row.get("transport_rate"))
        + decimal_value(row.get("overtime_hours")) * money_decimal(row.get("overtime_rate"))
        + decimal_value(row.get("holiday_hours")) * money_decimal(row.get("holiday_rate"))
    )
    mileage_total = decimal_value(row.get("miles")) * money_decimal(row.get("mileage_rate"))
    travel_total = sum((money_decimal(row.get(key)) for key in ("lodging", "airfare", "baggage", "rental_car", "fuel", "parking", "taxi", "other")), Decimal("0"))
    row["labor_total"] = money_float(labor_total)
    row["mileage_total"] = money_float(mileage_total)
    row["total"] = money_float(money_decimal(labor_total) + money_decimal(mileage_total) + travel_total)
    row["sort_order"] = sort_order
    return row


def customer_reimbursement_totals(items):
    labor_total = sum((money_decimal(item["labor_total"]) for item in items), Decimal("0"))
    lodging_total = sum((money_decimal(item["lodging"]) for item in items), Decimal("0"))
    travel_total = sum(
        money_decimal(item[key])
        for item in items
        for key in ("lodging", "airfare", "baggage", "rental_car", "fuel", "parking", "taxi", "other")
    ) or Decimal("0")
    mileage_total = sum((money_decimal(item["mileage_total"]) for item in items), Decimal("0"))
    total_amount = sum((money_decimal(item["total"]) for item in items), Decimal("0"))
    return {
        "labor_total": money_float(labor_total),
        "lodging_total": money_float(lodging_total),
        "travel_total": money_float(travel_total),
        "mileage_total": money_float(mileage_total),
        "total_amount": money_float(total_amount),
    }


def customer_reimbursement_items(reimbursement_id):
    return db().execute(
        """
        select *
        from customer_reimbursement_items
        where customer_reimbursement_id = ?
        order by sort_order, id
        """,
        (reimbursement_id,),
    ).fetchall()


def get_customer_reimbursement_attachments(reimbursement_id):
    return db().execute(
        """
        select customer_reimbursement_attachments.*, users.name as uploader_name
        from customer_reimbursement_attachments
        left join users on users.id = customer_reimbursement_attachments.uploaded_by
        where customer_reimbursement_id = ?
        order by uploaded_at desc, id desc
        """,
        (reimbursement_id,),
    ).fetchall()


def save_customer_reimbursement_attachment(reimbursement_id, uploaded):
    if not uploaded or not uploaded.filename:
        return
    if not allowed_attachment(uploaded.filename):
        raise ValueError(f"附件只支持 {ALLOWED_ATTACHMENT_LABEL}。")
    source_filename = uploaded.filename or "attachment"
    extension = source_filename.rsplit(".", 1)[1].lower()
    original_filename = os.path.basename(source_filename).strip() or f"attachment.{extension}"
    stored_filename = f"{secrets.token_hex(12)}.{extension}"
    uploaded.save(os.path.join(customer_reimbursement_attachment_dir(reimbursement_id), stored_filename))
    db().execute(
        """
        insert into customer_reimbursement_attachments (
            customer_reimbursement_id, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
        ) values (?, ?, ?, ?, ?, ?)
        """,
        (reimbursement_id, original_filename, stored_filename, uploaded.content_type, g.user["id"], now()),
    )


def save_customer_reimbursement_uploads(reimbursement_id):
    for uploaded in uploaded_attachments_from_request():
        save_customer_reimbursement_attachment(reimbursement_id, uploaded)


def customer_reimbursement_items_from_form():
    field_names = [
        "worker_name", "project_date", "standard_hours", "transport_hours", "overtime_hours", "holiday_hours",
        "lodging", "airfare", "baggage", "rental_car", "fuel", "parking", "taxi", "miles", "other",
    ]
    money_fields = {"lodging", "airfare", "baggage", "rental_car", "fuel", "parking", "taxi", "other"}
    rates = customer_reimbursement_rates()
    posted = {name: request.form.getlist(name) for name in field_names}
    count = max((len(values) for values in posted.values()), default=0)
    rows = []
    for index in range(count):
        worker_name = posted["worker_name"][index].strip() if index < len(posted["worker_name"]) else ""
        project_date = posted["project_date"][index].strip() if index < len(posted["project_date"]) else ""
        if not worker_name and not project_date:
            continue
        row = {"worker_name": worker_name, "project_date": project_date}
        for name in field_names[2:]:
            value = posted[name][index] if index < len(posted[name]) else 0
            row[name] = money_float(value) if name in money_fields else to_float(value)
        row.update(
            {
                "standard_rate": rates["标准工时"],
                "transport_rate": rates["交通工时"],
                "overtime_rate": rates["加班工时"],
                "holiday_rate": rates["节假日工时"],
                "mileage_rate": rates["里程费"],
            }
        )
        if not row["worker_name"] or not row["project_date"]:
            raise ValueError("工单结算每一行都必须有姓名和项目时间。")
        validate_lodging_reimbursement(row["lodging"], "工单结算住宿报销")
        validate_fuel_reimbursement_allowed(row["fuel"], "工单结算油费")
        rows.append(calculate_customer_reimbursement_item(row, len(rows)))
    if not rows:
        raise ValueError("工单结算至少需要一行明细。")
    return rows


def save_customer_reimbursement_items(reimbursement_id, rows):
    db().execute("delete from customer_reimbursement_items where customer_reimbursement_id = ?", (reimbursement_id,))
    for row in rows:
        db().execute(
            """
            insert into customer_reimbursement_items (
                customer_reimbursement_id, worker_name, project_date, standard_hours, transport_hours,
                overtime_hours, holiday_hours, standard_rate, transport_rate, overtime_rate, holiday_rate,
                labor_total, lodging, airfare, baggage, rental_car, fuel, parking, taxi, miles, mileage_rate,
                mileage_total, other, total, sort_order
            ) values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                reimbursement_id, row["worker_name"], row["project_date"], row["standard_hours"], row["transport_hours"],
                row["overtime_hours"], row["holiday_hours"], row["standard_rate"], row["transport_rate"],
                row["overtime_rate"], row["holiday_rate"], row["labor_total"], row["lodging"], row["airfare"],
                row["baggage"], row["rental_car"], row["fuel"], row["parking"], row["taxi"], row["miles"],
                row["mileage_rate"], row["mileage_total"], row["other"], row["total"], row["sort_order"],
            ),
        )


def latest_customer_reimbursement(order_id):
    return db().execute(
        """
        select customer_reimbursements.*, users.name as creator_name
        from customer_reimbursements
        left join users on users.id = customer_reimbursements.created_by
        where service_order_id = ?
        order by created_at desc, id desc
        limit 1
        """,
        (order_id,),
    ).fetchone()


def service_order_active_invoice(order_id):
    return db().execute(
        """
        select *
        from invoices
        where service_order_id = ? and status != 'void'
        order by issue_date desc, id desc
        limit 1
        """,
        (order_id,),
    ).fetchone()


def customer_reimbursement_linked_invoice(reimbursement, order_id):
    if reimbursement and reimbursement["invoice_id"]:
        invoice = db().execute(
            "select * from invoices where id = ?",
            (reimbursement["invoice_id"],),
        ).fetchone()
        if invoice:
            return invoice
    return service_order_active_invoice(order_id)


def customer_reimbursement_invoice_items(reimbursement):
    projects = db().execute(
        """
        select *
        from projects
        where project_type = 'invoice' and name in (?, ?, ?)
        """,
        tuple(CUSTOMER_REIMBURSEMENT_INVOICE_PROJECTS.keys()),
    ).fetchall()
    projects_by_name = {project["name"]: project for project in projects}
    missing = [name for name in CUSTOMER_REIMBURSEMENT_INVOICE_PROJECTS if name not in projects_by_name]
    if missing:
        raise ValueError(f"请先在项目维护中创建发票项目：{', '.join(missing)}。")
    items = []
    for project_name, amount_field in CUSTOMER_REIMBURSEMENT_INVOICE_PROJECTS.items():
        amount = float(reimbursement[amount_field] or 0)
        if amount <= 0:
            continue
        project = projects_by_name[project_name]
        items.append(
            {
                "project_id": project["id"],
                "amount": round(amount, 2),
                "tax_rate": project["tax_rate"],
            }
        )
    if not items:
        raise ValueError("工单结算金额为 0，无法生成发票。")
    return items


def create_customer_reimbursement(order):
    file_name = f"费用报销单{order['client_order_number']}.pdf"
    cursor = db().execute(
        """
        insert into customer_reimbursements (
            service_order_id, file_name, stored_filename, created_by, created_at
        ) values (?, ?, ?, ?, ?)
        """,
        (order["id"], file_name, f"{secrets.token_hex(12)}.pdf", g.user["id"], now()),
    )
    reimbursement_id = cursor.lastrowid
    rows = customer_reimbursement_seed_rows(order["id"])
    if not rows:
        raise ValueError("这个工单还没有可用于生成工单结算的工作日报。")
    save_customer_reimbursement_items(reimbursement_id, rows)
    update_customer_reimbursement_totals(reimbursement_id, rows)
    return db().execute("select * from customer_reimbursements where id = ?", (reimbursement_id,)).fetchone()


def update_customer_reimbursement_totals(reimbursement_id, rows=None):
    rows = rows if rows is not None else customer_reimbursement_items(reimbursement_id)
    totals = customer_reimbursement_totals(rows)
    db().execute(
        """
        update customer_reimbursements
        set labor_total = ?, lodging_total = ?, travel_total = ?, mileage_total = ?, total_amount = ?
        where id = ?
        """,
        (
            totals["labor_total"], totals["lodging_total"], totals["travel_total"],
            totals["mileage_total"], totals["total_amount"], reimbursement_id,
        ),
    )
    return totals


def ensure_customer_reimbursement_pdf_record(reimbursement, order):
    if reimbursement["file_name"].lower().endswith(".pdf") and reimbursement["stored_filename"].lower().endswith(".pdf"):
        return reimbursement
    db().execute(
        """
        update customer_reimbursements
        set file_name = ?, stored_filename = ?
        where id = ?
        """,
        (
            f"费用报销单{order['client_order_number']}.pdf",
            f"{secrets.token_hex(12)}.pdf",
            reimbursement["id"],
        ),
    )
    return db().execute("select * from customer_reimbursements where id = ?", (reimbursement["id"],)).fetchone()


def reimbursement_number(value):
    number = float(value or 0)
    return f"{number:,.2f}" if number else ""


def reimbursement_paragraph(value, style):
    text = "" if value is None else str(value)
    return Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>"), style)


def build_customer_reimbursement_pdf(reimbursement, order, rows):
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    path = customer_reimbursement_file_path(reimbursement)
    os.makedirs(os.path.dirname(path), exist_ok=True)

    title_style = ParagraphStyle(
        "CustomerReimbursementTitle",
        fontName="STSong-Light",
        fontSize=16,
        leading=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#10233F"),
    )
    meta_style = ParagraphStyle(
        "CustomerReimbursementMeta",
        fontName="STSong-Light",
        fontSize=8,
        leading=11,
        alignment=TA_LEFT,
        textColor=colors.HexColor("#475467"),
    )
    header_style = ParagraphStyle(
        "CustomerReimbursementHeader",
        fontName="STSong-Light",
        fontSize=6,
        leading=7,
        alignment=TA_CENTER,
        textColor=colors.white,
    )
    cell_style = ParagraphStyle(
        "CustomerReimbursementCell",
        fontName="STSong-Light",
        fontSize=6,
        leading=7,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#101828"),
    )
    total_style = ParagraphStyle(
        "CustomerReimbursementTotal",
        parent=cell_style,
        fontSize=6.5,
    )

    document = SimpleDocTemplate(
        path,
        pagesize=landscape(A4),
        leftMargin=5 * mm,
        rightMargin=5 * mm,
        topMargin=7 * mm,
        bottomMargin=7 * mm,
        title=reimbursement["file_name"],
        author=get_company_profile()["name"],
    )

    header_group = [
        "姓名", "项目时间", "工时", "", "", "", "", "", "", "", "", "住宿",
        "交通", "", "", "", "", "", "里程", "", "", "其他", "合计",
    ]
    header_detail = [
        "", "", "标准\n工时", "交通\n工时", "加班\n工时", "节假日\n工时",
        "标准工时费/小时", "交通工时费/小时", "加班工时费/小时", "节假日工时费/小时",
        "工时费\n合计", "", "机票", "行李", "租车", "加油", "停车", "打车",
        "英里数", "里程费/英里", "里程费", "", "",
    ]
    table_data = [
        [reimbursement_paragraph(value, header_style) for value in header_group],
        [reimbursement_paragraph(value, header_style) for value in header_detail],
    ]

    for item in rows:
        values = [
            item["worker_name"],
            item["project_date"],
            reimbursement_number(item["standard_hours"]),
            reimbursement_number(item["transport_hours"]),
            reimbursement_number(item["overtime_hours"]),
            reimbursement_number(item["holiday_hours"]),
            reimbursement_number(item["standard_rate"]),
            reimbursement_number(item["transport_rate"]),
            reimbursement_number(item["overtime_rate"]),
            reimbursement_number(item["holiday_rate"]),
            reimbursement_number(item["labor_total"]),
            reimbursement_number(item["lodging"]),
            reimbursement_number(item["airfare"]),
            reimbursement_number(item["baggage"]),
            reimbursement_number(item["rental_car"]),
            reimbursement_number(item["fuel"]),
            reimbursement_number(item["parking"]),
            reimbursement_number(item["taxi"]),
            reimbursement_number(item["miles"]),
            reimbursement_number(item["mileage_rate"]),
            reimbursement_number(item["mileage_total"]),
            reimbursement_number(item["other"]),
            reimbursement_number(item["total"]),
        ]
        table_data.append([reimbursement_paragraph(value, cell_style) for value in values])

    sum_fields = {
        2: "standard_hours", 3: "transport_hours", 4: "overtime_hours", 5: "holiday_hours",
        10: "labor_total", 11: "lodging", 12: "airfare", 13: "baggage", 14: "rental_car",
        15: "fuel", 16: "parking", 17: "taxi", 18: "miles", 20: "mileage_total",
        21: "other", 22: "total",
    }
    total_values = ["Total", ""] + [""] * 21
    for column_index, field_name in sum_fields.items():
        total_values[column_index] = reimbursement_number(sum(float(item[field_name] or 0) for item in rows))
    table_data.append([reimbursement_paragraph(value, total_style) for value in total_values])

    column_widths = [
        width * 0.8 * mm
        for width in (
            22, 21, 12, 12, 12, 13, 18, 18, 18, 18, 18, 13,
            13, 13, 13, 13, 13, 13, 13, 15, 13, 13, 16,
        )
    ]
    table = LongTable(table_data, colWidths=column_widths, repeatRows=2, hAlign="CENTER")
    last_row = len(table_data) - 1
    table.setStyle(
        TableStyle(
            [
                ("SPAN", (0, 0), (0, 1)),
                ("SPAN", (1, 0), (1, 1)),
                ("SPAN", (2, 0), (10, 0)),
                ("SPAN", (11, 0), (11, 1)),
                ("SPAN", (12, 0), (17, 0)),
                ("SPAN", (18, 0), (20, 0)),
                ("SPAN", (21, 0), (21, 1)),
                ("SPAN", (22, 0), (22, 1)),
                ("BACKGROUND", (0, 0), (-1, 1), colors.HexColor("#0F766E")),
                ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#166F68")),
                ("BACKGROUND", (0, last_row), (-1, last_row), colors.HexColor("#E7F6F2")),
                ("TEXTCOLOR", (0, last_row), (-1, last_row), colors.HexColor("#10233F")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#98A2B3")),
                ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#475467")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )

    story = [
        Paragraph("费用报销单", title_style),
        Spacer(1, 3 * mm),
        Paragraph(
            f"工单编号：{order['order_number']}　　服务订单号码：{order['client_order_number']}　　站点：{order['client_name']}",
            meta_style,
        ),
        Spacer(1, 3 * mm),
        table,
        Spacer(1, 4 * mm),
        Paragraph("备注：", meta_style),
    ]

    def draw_page_number(page_canvas, doc):
        page_canvas.saveState()
        page_canvas.setFont("STSong-Light", 7)
        page_canvas.setFillColor(colors.HexColor("#667085"))
        page_canvas.drawRightString(landscape(A4)[0] - 5 * mm, 3 * mm, f"第 {doc.page} 页")
        page_canvas.restoreState()

    document.build(story, onFirstPage=draw_page_number, onLaterPages=draw_page_number)
    return path


def remove_customer_reimbursement_pdf(reimbursement):
    path = customer_reimbursement_file_path(reimbursement)
    try:
        os.remove(path)
    except FileNotFoundError:
        pass


def service_report_docx_filename(order, report, used_filenames=None):
    filename = secure_filename(
        f"{order['client_order_number']}-{report['report_date']}-report.docx"
    ) or f"service-report-{report['id']}.docx"
    if used_filenames is None or filename not in used_filenames:
        return filename
    stem, extension = os.path.splitext(filename)
    return f"{stem}-{report['id']}{extension or '.docx'}"


def service_order_report_word_attachments(order):
    reports = db().execute(
        """
        select * from service_reports
        where service_order_id = ?
        order by report_date, id
        """,
        (order["id"],),
    ).fetchall()
    attachments = []
    used_filenames = set()
    for report in reports:
        try:
            content = build_service_report_docx(report, order)
        except Exception as error:
            raise ValueError(f"工作日报 {report['report_date']} Word 文件生成失败，请检查日报内容或照片。") from error
        filename = service_report_docx_filename(order, report, used_filenames)
        used_filenames.add(filename)
        attachments.append(
            {
                "filename": filename,
                "content": content,
                "maintype": "application",
                "subtype": "vnd.openxmlformats-officedocument.wordprocessingml.document",
            }
        )
    return attachments


def customer_reimbursement_outgoing_attachments(reimbursement, order):
    reimbursement = ensure_customer_reimbursement_pdf_record(reimbursement, order)
    pdf_path = build_customer_reimbursement_pdf(
        reimbursement,
        order,
        customer_reimbursement_items(reimbursement["id"]),
    )
    attachments = []
    with open(pdf_path, "rb") as file:
        attachments.append(
            {
                "filename": reimbursement["file_name"],
                "content": file.read(),
                "maintype": "application",
                "subtype": "pdf",
            }
        )
    for attachment in get_customer_reimbursement_attachments(reimbursement["id"]):
        path = customer_reimbursement_attachment_path(attachment)
        if not os.path.exists(path):
            continue
        maintype, _, subtype = (attachment["content_type"] or "application/octet-stream").partition("/")
        with open(path, "rb") as file:
            attachments.append(
                {
                    "filename": attachment["original_filename"],
                    "content": file.read(),
                    "maintype": maintype or "application",
                    "subtype": subtype or "octet-stream",
                }
            )
    attachments.extend(service_order_report_word_attachments(order))
    return reimbursement, attachments


def deliver_customer_reimbursement_email(reimbursement, order):
    client = db().execute("select * from clients where id = ?", (order["client_id"],)).fetchone() if order["client_id"] else None
    if not client:
        raise ValueError("这个工单还没有关联客户，请先编辑工单选择客户。")
    recipient = (client["email"] or "").strip()
    if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", recipient):
        raise ValueError(f"客户 {client['name']} 没有有效邮箱，请先维护客户主数据。")
    subject = f"Customer Reimbursement Statement {order['client_order_number']}"
    body = (
        f"Dear {client['contact_name'] or client['name']},\n\n"
        "Please find the customer reimbursement statement and supporting documents attached."
    )
    reimbursement, attachments = customer_reimbursement_outgoing_attachments(reimbursement, order)
    send_email(
        to=recipient,
        subject=subject,
        html=f"<p>{html.escape(body).replace(chr(10), '<br>')}</p>",
        attachments=attachments,
    )
    log_action(
        "send",
        "customer_reimbursement",
        reimbursement["id"],
        reimbursement["file_name"],
        f"发送至：{recipient}；附件：{len(attachments)} 个",
    )
    return recipient


def can_access_expense(expense):
    if not g.user:
        return False
    if normalized_role() in {"admin", "manager", "finance"}:
        return True
    return expense["created_by"] == g.user["id"]


def require_expense(expense_id):
    expense = db().execute("select * from expenses where id = ?", (expense_id,)).fetchone()
    if not expense:
        abort(404)
    if not can_access_expense(expense):
        abort(403)
    order = require_service_order(expense["service_order_id"])
    return expense, order


def expense_attachment_dir(expense_id):
    path = os.path.join(EXPENSE_ATTACHMENTS_DIR, str(expense_id))
    os.makedirs(path, exist_ok=True)
    return path


def expense_attachment_path(attachment):
    return os.path.join(EXPENSE_ATTACHMENTS_DIR, str(attachment["expense_id"]), attachment["stored_filename"])


def save_expense_attachment(expense_id, uploaded):
    if not uploaded or not uploaded.filename:
        return
    if not allowed_attachment(uploaded.filename):
        raise ValueError(f"附件只支持 {ALLOWED_ATTACHMENT_LABEL}。")
    source_filename = uploaded.filename or "attachment"
    extension = source_filename.rsplit(".", 1)[1].lower()
    original_filename = os.path.basename(source_filename).strip() or f"attachment.{extension}"
    stored_filename = f"{secrets.token_hex(12)}.{extension}"
    uploaded.save(os.path.join(expense_attachment_dir(expense_id), stored_filename))
    db().execute(
        """
        insert into expense_attachments (
            expense_id, original_filename, stored_filename, content_type, uploaded_by, uploaded_at
        ) values (?, ?, ?, ?, ?, ?)
        """,
        (expense_id, original_filename, stored_filename, uploaded.content_type, g.user["id"], now()),
    )


def save_expense_uploads(expense_id):
    for uploaded in uploaded_attachments_from_request():
        save_expense_attachment(expense_id, uploaded)


def get_expense_attachments(expense_id):
    return db().execute(
        """
        select expense_attachments.*, users.name as uploader_name
        from expense_attachments left join users on users.id = expense_attachments.uploaded_by
        where expense_id = ?
        order by uploaded_at desc, id desc
        """,
        (expense_id,),
    ).fetchall()


def service_report_workers(report_id):
    return db().execute(
        """
        select users.id, users.name, users.email, service_report_workers.driving_miles
        from service_report_workers
        join users on users.id = service_report_workers.user_id
        where service_report_workers.report_id = ?
        order by users.name
        """,
        (report_id,),
    ).fetchall()


def service_report_worker_options(order, report_id=None):
    params = [order["country_code"]]
    historical_clause = ""
    if report_id:
        historical_clause = """
          or users.id in (
              select user_id from service_report_workers where report_id = ?
          )
        """
        params.append(report_id)
    role_clause = (
        "users.id = ?"
        if is_external_employee()
        else "users.role in ('manager', 'finance', 'employee')"
    )
    if is_external_employee():
        params.insert(0, g.user["id"])
    return db().execute(
        f"""
        select users.id, users.name, users.email, users.country_code
        from users
        where {role_clause}
          and users.is_active = 1
          and (
              users.country_code = ?
              {historical_clause}
          )
        order by users.name
        """,
        params,
    ).fetchall()


def report_parts(table, report_id):
    return db().execute(f"select * from {table} where report_id = ? order by sort_order, id", (report_id,)).fetchall()


def parse_part_rows(prefix, fields):
    values = [request.form.getlist(f"{prefix}_{field}") for field in fields]
    max_len = max([len(items) for items in values] + [0])
    rows = []
    for index in range(max_len):
        row = {field: values[pos][index].strip() if index < len(values[pos]) else "" for pos, field in enumerate(fields)}
        if any(row.values()):
            rows.append(row)
    return rows


def report_form_defaults(report=None, order=None):
    if report:
        data = dict(report)
        data["actual_work_date"] = data.get("actual_work_date") or data.get("report_date")
        return data
    return {
        "report_date": date.today().isoformat(),
        "actual_work_date": date.today().isoformat(),
        "total_service_hours": "",
        "travel_hours": "",
        "public_transport_hours": "",
        "driving_miles": "",
        "departure_address": "",
        "site_address": order["site_address"] if order else "",
        "total_time": "",
        "cabinet_number": "",
        "arrival_time": "",
        "departure_time": "",
        "service_description": "",
    }


def save_report_detail_rows(report_id):
    worker_ids = report_worker_ids_from_form()
    if not worker_ids:
        raise ValueError("服务人员清单至少需要选择一人。")
    placeholders = ",".join("?" for _ in worker_ids)
    report_order = db().execute(
        """
        select service_orders.country_code
        from service_reports
        join service_orders on service_orders.id = service_reports.service_order_id
        where service_reports.id = ?
        """,
        (report_id,),
    ).fetchone()
    if not report_order:
        raise ValueError("工作日报关联的工单不存在。")
    valid_workers = db().execute(
        f"""
        select id from users
        where role in ('manager', 'finance', 'employee', 'external_employee')
          and is_active = 1
          and (
              country_code = ?
              or id in (
                  select user_id from service_report_workers where report_id = ?
              )
          )
          and id in ({placeholders})
        """,
        (report_order["country_code"], report_id, *worker_ids),
    ).fetchall()
    valid_worker_ids = {str(row["id"]) for row in valid_workers}
    if valid_worker_ids != set(worker_ids):
        raise ValueError("服务人员清单包含无效用户，请重新选择。")

    db().execute("delete from service_report_workers where report_id = ?", (report_id,))
    for user_id in worker_ids:
        worker_miles = to_float(request.form.get(f"worker_driving_miles_{user_id}"))
        db().execute(
            "insert into service_report_workers (report_id, user_id, driving_miles) values (?, ?, ?)",
            (report_id, user_id, worker_miles),
        )

    db().execute("delete from service_report_saved_parts where report_id = ?", (report_id,))
    for index, row in enumerate(parse_part_rows("saved", ["part_number", "part_name", "quantity", "status"])):
        db().execute(
            """
            insert into service_report_saved_parts (report_id, part_number, part_name, quantity, status, sort_order)
            values (?, ?, ?, ?, ?, ?)
            """,
            (report_id, row["part_number"], row["part_name"], row["quantity"], row["status"], index),
        )

    db().execute("delete from service_report_replaced_parts where report_id = ?", (report_id,))
    for index, row in enumerate(parse_part_rows("replaced", ["part_number", "part_name", "old_serial_number", "new_serial_number", "quantity"])):
        db().execute(
            """
            insert into service_report_replaced_parts (
                report_id, part_number, part_name, old_serial_number, new_serial_number, quantity, sort_order
            ) values (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                report_id,
                row["part_number"],
                row["part_name"],
                row["old_serial_number"],
                row["new_serial_number"],
                row["quantity"],
                index,
            ),
        )


def load_invoice(invoice_id):
    invoice = require_invoice_access(invoice_id)
    client = db().execute("select * from clients where id = ?", (invoice["client_id"],)).fetchone()
    items = db().execute("select * from invoice_items where invoice_id = ?", (invoice_id,)).fetchall()
    return invoice, client, items


def create_message(user_id, title, body, link=None):
    db().execute(
        "insert into messages (user_id, title, body, link, created_at) values (?, ?, ?, ?, ?)",
        (user_id, title, body, link, now()),
    )


def log_action(action, entity_type, entity_id, entity_label, summary=""):
    if not g.user:
        return
    db().execute(
        """
        insert into audit_logs (
            user_id, user_name, action, entity_type, entity_id, entity_label, summary, created_at
        ) values (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            g.user["id"],
            g.user["name"],
            action,
            entity_type,
            entity_id,
            str(entity_label or ""),
            str(summary or ""),
            now(),
        ),
    )


def request_ip_address():
    forwarded_for = request.headers.get("X-Forwarded-For", "").split(",", 1)[0].strip()
    return forwarded_for or request.headers.get("X-Real-IP", "").strip() or request.remote_addr or ""


def log_login_action(user):
    db().execute(
        """
        insert into audit_logs (
            user_id, user_name, action, entity_type, entity_id, entity_label, summary, created_at
        ) values (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            user["id"],
            user["name"],
            "login",
            "user",
            user["id"],
            user["email"] or user["name"],
            f"角色：{role_label(user['role'])}；IP：{request_ip_address() or '-'}",
            now(),
        ),
    )


def normalized_address(value):
    return " ".join(str(value or "").strip().split()).casefold()


class TemporaryGeocodingError(Exception):
    pass


def fetch_geocoder_json(request_url):
    global _last_geocode_request_at
    with _geocode_lock:
        wait_seconds = 1.05 - (time.monotonic() - _last_geocode_request_at)
        if wait_seconds > 0:
            time.sleep(wait_seconds)
        request_object = Request(
            request_url,
            headers={
                "User-Agent": NOMINATIM_USER_AGENT,
                "Accept": "application/json",
                "Accept-Language": "en-US,en;q=0.8",
            },
        )
        try:
            with urlopen(request_object, timeout=8) as response:
                payload = json.loads(response.read().decode("utf-8"))
        except HTTPError as error:
            if error.code == 429 or error.code >= 500:
                raise TemporaryGeocodingError(str(error)) from error
            raise
        except (URLError, TimeoutError, OSError) as error:
            raise TemporaryGeocodingError(str(error)) from error
        finally:
            _last_geocode_request_at = time.monotonic()
    return payload


def configured_country_codes():
    return {code.strip().casefold() for code in NOMINATIM_COUNTRY_CODES.split(",") if code.strip()}


def address_expectations(address):
    uppercase_address = str(address or "").upper()
    state_codes = [token for token in re.findall(r"\b[A-Z]{2}\b", uppercase_address) if token in US_STATE_CODES]
    zip_codes = re.findall(r"\b(\d{5})(?:-\d{4})?\b", uppercase_address)
    return state_codes[-1] if state_codes else None, zip_codes[-1] if zip_codes else None


def census_result_matches_address(match, address):
    expected_state, expected_zip = address_expectations(address)
    components = match.get("addressComponents", {})
    result_state = str(components.get("state", "")).upper()
    result_zip = str(components.get("zip", ""))[:5]
    if expected_state and result_state != expected_state:
        return False
    if expected_zip and result_zip and result_zip != expected_zip:
        return False
    return True


def geocode_with_google(address):
    api_key = get_google_geocoding_api_key()
    if not api_key:
        return None
    expected_state, expected_zip = address_expectations(address)
    query = {
        "address": address,
        "key": api_key,
        "region": "us",
    }
    try:
        payload = fetch_geocoder_json(f"{GOOGLE_GEOCODING_URL}?{urlencode(query)}")
    except TemporaryGeocodingError:
        raise
    except (HTTPError, ValueError, json.JSONDecodeError):
        return None
    if not isinstance(payload, dict):
        return None
    status = payload.get("status")
    if status in {"OVER_QUERY_LIMIT", "RESOURCE_EXHAUSTED", "UNKNOWN_ERROR"}:
        raise TemporaryGeocodingError(f"Google geocoding status: {status}")
    if status != "OK":
        app.logger.info("Google geocoding returned %s for address %s", status, address)
        return None
    for result in payload.get("results", []):
        components = result.get("address_components", [])
        result_state = ""
        result_zip = ""
        for component in components:
            types = set(component.get("types", []))
            if "administrative_area_level_1" in types:
                result_state = str(component.get("short_name", "")).upper()
            if "postal_code" in types:
                result_zip = str(component.get("long_name", ""))[:5]
        if expected_state and result_state and result_state != expected_state:
            continue
        if expected_zip and result_zip and result_zip != expected_zip:
            continue
        try:
            location = result["geometry"]["location"]
            return float(location["lat"]), float(location["lng"])
        except (KeyError, TypeError, ValueError):
            continue
    return None


def geocode_with_census(address):
    query = {
        "address": address,
        "benchmark": "Public_AR_Current",
        "format": "json",
    }
    try:
        payload = fetch_geocoder_json(f"{CENSUS_GEOCODER_URL}?{urlencode(query)}")
    except TemporaryGeocodingError:
        raise
    except (HTTPError, ValueError, json.JSONDecodeError):
        return None
    matches = payload.get("result", {}).get("addressMatches", []) if isinstance(payload, dict) else []
    for match in matches:
        if not census_result_matches_address(match, address):
            continue
        try:
            coordinates = match["coordinates"]
            return float(coordinates["y"]), float(coordinates["x"])
        except (KeyError, TypeError, ValueError):
            continue
    return None


def nominatim_address_candidates(address):
    candidates = [address]
    parts = [part.strip() for part in address.split(",") if part.strip()]
    if len(parts) >= 3:
        candidates.append(", ".join(parts[-3:]))
    zip_match = re.search(r"\b\d{5}(?:-\d{4})?\b", address)
    if zip_match:
        candidates.append(f"{zip_match.group(0)}, USA")
    unique_candidates = []
    seen = set()
    for candidate in candidates:
        key = normalized_address(candidate)
        if key and key not in seen:
            seen.add(key)
            unique_candidates.append(candidate)
    return unique_candidates


def geocode_with_nominatim(address):
    country_codes = configured_country_codes()
    expected_state, expected_zip = address_expectations(address)
    for candidate in nominatim_address_candidates(address):
        query = {
            "q": candidate,
            "format": "jsonv2",
            "limit": "1",
            "addressdetails": "1",
        }
        if NOMINATIM_COUNTRY_CODES:
            query["countrycodes"] = NOMINATIM_COUNTRY_CODES
        try:
            payload = fetch_geocoder_json(f"{NOMINATIM_URL}?{urlencode(query)}")
        except TemporaryGeocodingError:
            raise
        except (HTTPError, ValueError, json.JSONDecodeError):
            continue
        if not payload:
            continue
        result = payload[0]
        result_address = result.get("address", {})
        result_country = str(result_address.get("country_code", "")).casefold()
        if country_codes and result_country and result_country not in country_codes:
            continue
        result_state = str(result_address.get("ISO3166-2-lvl4", "")).upper().removeprefix("US-")
        result_zip = str(result_address.get("postcode", ""))[:5]
        if expected_state and result_state and result_state != expected_state:
            continue
        if expected_zip and result_zip and result_zip != expected_zip:
            continue
        try:
            return float(result["lat"]), float(result["lon"])
        except (KeyError, TypeError, ValueError, IndexError):
            continue
    return None


def geocode_address(address):
    if not GEOCODING_ENABLED:
        return None
    temporary_error = None
    try:
        coordinates = geocode_with_google(address)
        if coordinates:
            return coordinates
    except TemporaryGeocodingError as error:
        temporary_error = error
    if "us" in configured_country_codes():
        try:
            coordinates = geocode_with_census(address)
            if coordinates:
                return coordinates
        except TemporaryGeocodingError as error:
            temporary_error = error
    try:
        coordinates = geocode_with_nominatim(address)
        if coordinates:
            return coordinates
    except TemporaryGeocodingError as error:
        temporary_error = error
    if temporary_error:
        raise temporary_error
    return None


def geocode_service_order(order_id, force=False):
    order = db().execute("select * from service_orders where id = ?", (order_id,)).fetchone()
    if not order:
        return None
    address = str(order["site_address"] or "").strip()
    address_key = normalized_address(address)
    if not address_key:
        db().execute(
            """
            update service_orders
            set latitude = null, longitude = null, geocode_address = ?, geocode_status = 'failed',
                geocode_attempted_at = ?, geocode_version = ?
            where id = ?
            """,
            (address, now(), GEOCODER_VERSION, order_id),
        )
        return None
    if (
        not force
        and order["geocode_status"] == "success"
        and order["latitude"] is not None
        and order["longitude"] is not None
        and normalized_address(order["geocode_address"]) == address_key
        and order["geocode_version"] == GEOCODER_VERSION
    ):
        return float(order["latitude"]), float(order["longitude"])
    cached = db().execute(
        """
        select latitude, longitude
        from service_orders
        where id != ? and geocode_status = 'success'
          and latitude is not null and longitude is not null
          and lower(trim(geocode_address)) = lower(trim(?))
          and geocode_version = ?
        order by geocode_attempted_at desc, id desc
        limit 1
        """,
        (order_id, address, GEOCODER_VERSION),
    ).fetchone()
    if cached:
        coordinates = float(cached["latitude"]), float(cached["longitude"])
    else:
        try:
            coordinates = geocode_address(address)
        except TemporaryGeocodingError:
            raise
        except (HTTPError, URLError, TimeoutError, OSError, ValueError, json.JSONDecodeError) as error:
            app.logger.warning("Unable to geocode service order %s: %s", order_id, error)
            coordinates = None
    if coordinates:
        db().execute(
            """
            update service_orders
            set latitude = ?, longitude = ?, geocode_address = ?, geocode_status = 'success',
                geocode_attempted_at = ?, geocode_version = ?
            where id = ?
            """,
            (coordinates[0], coordinates[1], address, now(), GEOCODER_VERSION, order_id),
        )
        return coordinates
    db().execute(
        """
        update service_orders
        set latitude = null, longitude = null, geocode_address = ?, geocode_status = 'failed',
            geocode_attempted_at = ?, geocode_version = ?
        where id = ?
        """,
        (address, now(), GEOCODER_VERSION, order_id),
    )
    return None


def geocode_buyer(buyer_id, force=False):
    buyer = db().execute("select * from buyers where id = ?", (buyer_id,)).fetchone()
    if not buyer:
        return None
    address = str(buyer["detailed_address"] or "").strip()
    address_key = normalized_address(address)
    if not address_key:
        db().execute(
            """
            update buyers
            set latitude = null, longitude = null, geocode_address = ?, geocode_status = 'failed',
                geocode_attempted_at = ?, geocode_version = ?
            where id = ?
            """,
            (address, now(), GEOCODER_VERSION, buyer_id),
        )
        return None
    if (
        not force
        and buyer["geocode_status"] == "success"
        and buyer["latitude"] is not None
        and buyer["longitude"] is not None
        and normalized_address(buyer["geocode_address"]) == address_key
        and buyer["geocode_version"] == GEOCODER_VERSION
    ):
        return float(buyer["latitude"]), float(buyer["longitude"])
    cached = db().execute(
        """
        select latitude, longitude
        from buyers
        where id != ? and geocode_status = 'success'
          and latitude is not null and longitude is not null
          and lower(trim(geocode_address)) = lower(trim(?))
          and geocode_version = ?
        order by geocode_attempted_at desc, id desc
        limit 1
        """,
        (buyer_id, address, GEOCODER_VERSION),
    ).fetchone()
    if cached:
        coordinates = float(cached["latitude"]), float(cached["longitude"])
    else:
        try:
            coordinates = geocode_address(address)
        except TemporaryGeocodingError:
            raise
        except (HTTPError, URLError, TimeoutError, OSError, ValueError, json.JSONDecodeError) as error:
            app.logger.warning("Unable to geocode buyer %s: %s", buyer_id, error)
            coordinates = None
    if coordinates:
        db().execute(
            """
            update buyers
            set latitude = ?, longitude = ?, geocode_address = ?, geocode_status = 'success',
                geocode_attempted_at = ?, geocode_version = ?
            where id = ?
            """,
            (coordinates[0], coordinates[1], address, now(), GEOCODER_VERSION, buyer_id),
        )
        return coordinates
    db().execute(
        """
        update buyers
        set latitude = null, longitude = null, geocode_address = ?, geocode_status = 'failed',
            geocode_attempted_at = ?, geocode_version = ?
        where id = ?
        """,
        (address, now(), GEOCODER_VERSION, buyer_id),
    )
    return None


def notify_role(roles, title, body, link=None, exclude_user_ids=None):
    excluded = {int(user_id) for user_id in (exclude_user_ids or []) if user_id is not None}
    placeholders = ",".join("?" for _ in roles)
    rows = db().execute(f"select id from users where role in ({placeholders})", list(roles)).fetchall()
    for row in rows:
        if row["id"] in excluded:
            continue
        create_message(row["id"], title, body, link)


def review_message_body(user_name, invoice, client, total, is_resubmission=False):
    action = "重新提交了发票" if is_resubmission else "提交了发票"
    client_label = client["short_name"] or client["name"]
    return f"{user_name}{action} {invoice['invoice_number']} 客户:{client_label} 金额:{money(total, invoice['currency'])} 请审核。"


def return_message_body(invoice, client, total, reason):
    client_label = client["short_name"] or client["name"]
    return f"发票 {invoice['invoice_number']} 已被经理退回。客户:{client_label} 金额:{money(total, invoice['currency'])} 原因：{reason}"


def unread_message_count():
    if not g.user:
        return 0
    return db().execute(
        "select count(*) as count from messages where user_id = ? and is_read = 0",
        (g.user["id"],),
    ).fetchone()["count"]


@app.context_processor
def inject_globals():
    return {
        "unread_message_count": unread_message_count(),
        "current_language": current_language(),
        "supported_languages": SUPPORTED_LANGUAGES,
        "can_manage_company_info": can_manage_company_info,
    }


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        language = request.form.get("language")
        if language in SUPPORTED_LANGUAGES:
            session["language"] = language
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        user = db().execute("select * from users where email = ?", (email,)).fetchone()
        if user and check_password_hash(user["password_hash"], password):
            if not user["is_active"]:
                flash("账号尚未启用，请等待管理员或经理批准。", "error")
                return redirect(url_for("login"))
            clear_session_preserving_language()
            session["user_id"] = user["id"]
            log_login_action(user)
            db().commit()
            return redirect(request.args.get("next") or url_for("dashboard"))
        flash("邮箱或密码不正确。", "error")
    return render_template("login.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("registration_email", "").strip().lower()
        password = request.form.get("registration_password", "")
        password_confirm = request.form.get("registration_password_confirm", "")
        account_type = request.form.get("account_type", "external_employee")
        country = country_from_form()
        if account_type not in {"employee", "external_employee"}:
            account_type = "external_employee"
        if "�" in name:
            flash("姓名包含损坏字符，请重新输入正确姓名。", "error")
            return redirect(url_for("register"))
        if len(password) < 8:
            flash("密码至少需要 8 位。", "error")
            return redirect(url_for("register"))
        if password != password_confirm:
            flash("两次输入的密码不一致。", "error")
            return redirect(url_for("register"))
        try:
            cursor = db().execute(
                """
                insert into users (
                    name, email, password_hash, role, is_active, region_code, country_code, created_at
                )
                values (?, ?, ?, ?, 0, ?, ?, ?)
                """,
                (
                    name or email,
                    email,
                    generate_password_hash(password),
                    account_type,
                    country["region_code"],
                    country["code"],
                    now(),
                ),
            )
            user_id = cursor.lastrowid
            account_label = "员工" if account_type == "employee" else "外部员工"
            notify_role(
                ["admin", "manager"],
                f"新{account_label}账号待批准",
                f"{name or email} 已注册{account_label}账号，请审核并启用。",
                url_for("users"),
            )
            db().commit()
            flash("注册成功，请等待管理员或经理批准启用。", "success")
            return redirect(url_for("login"))
        except sqlite3.IntegrityError:
            flash("这个邮箱已经注册。", "error")
    return render_template("register.html", countries=country_rows())


@app.route("/logout")
def logout():
    clear_session_preserving_language()
    return redirect(url_for("login"))


@app.post("/language")
def switch_language():
    language = request.form.get("language", "")
    if language in SUPPORTED_LANGUAGES:
        session["language"] = language
    next_url = request.form.get("next", "").strip()
    if not next_url.startswith("/") or next_url.startswith("//"):
        next_url = url_for("dashboard") if g.user else url_for("login")
    else:
        parsed = urlsplit(next_url)
        cleaned_query = urlencode(
            [(key, value) for key, value in parse_qsl(parsed.query, keep_blank_values=True) if key != "lang"],
            doseq=True,
        )
        next_url = urlunsplit(("", "", parsed.path, cleaned_query, parsed.fragment))
    return redirect(next_url)


@app.route("/messages")
@login_required
def messages():
    rows = db().execute(
        "select * from messages where user_id = ? order by is_read asc, datetime(created_at) desc, id desc",
        (g.user["id"],),
    ).fetchall()
    return render_template("messages.html", messages=rows)


@app.route("/messages/<int:message_id>")
@login_required
def message_detail(message_id):
    message = db().execute(
        "select * from messages where id = ? and user_id = ?",
        (message_id, g.user["id"]),
    ).fetchone()
    if not message:
        abort(404)
    db().execute("update messages set is_read = 1 where id = ?", (message_id,))
    db().commit()
    link = message["link"]
    if link:
        parts = link.strip("/").split("/")
        if len(parts) >= 2 and parts[0] == "invoices" and parts[1].isdigit():
            invoice_exists = db().execute("select id from invoices where id = ?", (int(parts[1]),)).fetchone()
            if not invoice_exists:
                flash("这条消息对应的发票已经被删除。", "error")
                return redirect(url_for("messages"))
        return redirect(link)
    return redirect(url_for("messages"))


@app.route("/")
@login_required
def dashboard():
    if not can_view_invoices() or is_external_user():
        return redirect(url_for("service_orders"))
    metrics = get_metrics()
    selected_project_ids = request.args.getlist("project_id")
    project_options = dashboard_project_options()
    project_filter_submitted = "project_filter" in request.args
    available_project_ids = {str(project["id"]) for project in project_options}
    selected_project_ids = [project_id for project_id in selected_project_ids if project_id in available_project_ids]
    if not project_filter_submitted and not selected_project_ids:
        selected_project_ids = [str(project["id"]) for project in project_options]
    chart = monthly_project_chart(selected_project_ids)
    paid_chart = monthly_paid_chart()
    access_clause, access_params = client_filter_clause("invoices")
    recent = db().execute(
        f"""
        select invoices.*, clients.name as client_name
        from invoices join clients on clients.id = invoices.client_id
        where invoices.status = 'completed' and {access_clause}
        order by invoices.created_at desc limit 8
        """,
        access_params,
    ).fetchall()
    return render_template(
        "dashboard.html",
        metrics=metrics,
        chart=chart,
        paid_chart=paid_chart,
        recent=recent,
        labels=STATUS_LABELS,
        project_options=project_options,
        selected_project_ids=selected_project_ids,
    )


@app.route("/contracts")
@login_required
def contracts():
    if not can_view_contracts():
        abort(403)
    q = request.args.get("q", "").strip()
    contract_type = request.args.get("contract_type", "").strip()
    status = request.args.get("status", "").strip()
    clauses = ["1 = 1"]
    params = []
    if is_external_manager():
        clauses.append("contracts.client_id = ?")
        params.append(g.user["client_id"] or 0)
    if q:
        clauses.append(
            "(contracts.contract_number like ? or contracts.title like ? "
            "or clients.name like ? or contracts.project_name like ?)"
        )
        params.extend([f"%{q}%"] * 4)
    if contract_type in CONTRACT_TYPE_LABELS:
        clauses.append("contracts.contract_type = ?")
        params.append(contract_type)
    if status in CONTRACT_STATUS_LABELS:
        clauses.append("contracts.status = ?")
        params.append(status)
    rows = db().execute(
        f"""
        select contracts.*, clients.name as client_name,
               count(distinct service_orders.id) as work_order_count
        from contracts
        join clients on clients.id = contracts.client_id
        left join service_orders on service_orders.contract_id = contracts.id
        where {" and ".join(clauses)}
        group by contracts.id
        order by
          case contracts.status when 'active' then 0 when 'signed' then 1 when 'review' then 2 else 3 end,
          coalesce(contracts.end_date, '9999-12-31'), contracts.id desc
        """,
        params,
    ).fetchall()
    return render_template(
        "contracts.html",
        contracts=rows,
        q=q,
        selected_type=contract_type,
        selected_status=status,
    )


@app.route("/contracts/new", methods=["GET", "POST"])
@login_required
def new_contract():
    if not can_manage_contracts():
        abort(403)
    clients_rows = db().execute("select * from clients order by client_number, name").fetchall()
    if not clients_rows:
        flash("请先创建客户。", "error")
        return redirect(url_for("clients"))
    if request.method == "POST":
        try:
            values = contract_form_values()
            cursor = db().execute(
                """
                insert into contracts (
                    contract_number, client_id, contract_type, title, status, signed_date,
                    start_date, end_date, currency, amount, payment_terms, rate_card,
                    project_name, notes, created_by, created_at, updated_at
                ) values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    values["contract_number"], values["client_id"], values["contract_type"],
                    values["title"], values["status"], values["signed_date"], values["start_date"],
                    values["end_date"], values["currency"], values["amount"],
                    values["payment_terms"], values["rate_card"], values["project_name"],
                    values["notes"], g.user["id"], now(), now(),
                ),
            )
            contract_id = cursor.lastrowid
            save_contract_uploads(contract_id)
            log_action(
                "create",
                "contract",
                contract_id,
                values["contract_number"],
                f"{CONTRACT_TYPE_LABELS[values['contract_type']]} · {values['title']}",
            )
            db().commit()
            flash("合同已创建。", "success")
            return redirect(url_for("contract_detail", contract_id=contract_id))
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("合同编号已经存在，请更换合同编号。", "error")
    defaults = dict(request.form) if request.method == "POST" else {
        "contract_number": next_contract_number(),
        "contract_type": "framework",
        "status": "draft",
        "currency": "USD",
    }
    return render_template(
        "contract_form.html",
        contract=None,
        clients=clients_rows,
        defaults=defaults,
        attachments=[],
        form_title="新建合同",
    )


@app.route("/contracts/<int:contract_id>")
@login_required
def contract_detail(contract_id):
    contract = require_contract_access(contract_id)
    client = db().execute("select * from clients where id = ?", (contract["client_id"],)).fetchone()
    attachments = db().execute(
        """
        select contract_attachments.*, users.name as uploader_name
        from contract_attachments
        left join users on users.id = contract_attachments.uploaded_by
        where contract_attachments.contract_id = ?
        order by contract_attachments.uploaded_at desc, contract_attachments.id desc
        """,
        (contract_id,),
    ).fetchall()
    work_orders = db().execute(
        """
        select service_orders.*, work_order_types.name as work_order_type_name
        from service_orders
        left join work_order_types on work_order_types.id = service_orders.work_order_type_id
        where service_orders.contract_id = ?
        order by service_orders.created_at desc, service_orders.id desc
        """,
        (contract_id,),
    ).fetchall()
    invoices = db().execute(
        """
        select distinct invoices.*
        from invoices
        join service_orders on service_orders.id = invoices.service_order_id
        where service_orders.contract_id = ? and invoices.status != 'void'
        order by invoices.issue_date desc, invoices.id desc
        """,
        (contract_id,),
    ).fetchall()
    return render_template(
        "contract_detail.html",
        contract=contract,
        client=client,
        attachments=attachments,
        work_orders=work_orders,
        invoices=invoices,
        labels=STATUS_LABELS,
    )


@app.route("/contracts/<int:contract_id>/edit", methods=["GET", "POST"])
@login_required
def edit_contract(contract_id):
    if not can_manage_contracts():
        abort(403)
    contract = require_contract_access(contract_id)
    clients_rows = db().execute("select * from clients order by client_number, name").fetchall()
    if request.method == "POST":
        try:
            values = contract_form_values(contract)
            linked_clients = db().execute(
                """
                select count(*) as count from service_orders
                where contract_id = ? and client_id != ?
                """,
                (contract_id, values["client_id"]),
            ).fetchone()["count"]
            if linked_clients:
                raise ValueError("合同已有其他客户的工单，不能更改合同客户。")
            db().execute(
                """
                update contracts
                set contract_number = ?, client_id = ?, contract_type = ?, title = ?, status = ?,
                    signed_date = ?, start_date = ?, end_date = ?, currency = ?, amount = ?,
                    payment_terms = ?, rate_card = ?, project_name = ?, notes = ?, updated_at = ?
                where id = ?
                """,
                (
                    values["contract_number"], values["client_id"], values["contract_type"],
                    values["title"], values["status"], values["signed_date"], values["start_date"],
                    values["end_date"], values["currency"], values["amount"],
                    values["payment_terms"], values["rate_card"], values["project_name"],
                    values["notes"], now(), contract_id,
                ),
            )
            save_contract_uploads(contract_id)
            log_action("update", "contract", contract_id, values["contract_number"], "修改合同资料")
            db().commit()
            flash("合同已更新。", "success")
            return redirect(url_for("contract_detail", contract_id=contract_id))
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("合同编号已经存在，请更换合同编号。", "error")
    attachments = db().execute(
        "select * from contract_attachments where contract_id = ? order by uploaded_at desc, id desc",
        (contract_id,),
    ).fetchall()
    defaults = dict(request.form) if request.method == "POST" else dict(contract)
    return render_template(
        "contract_form.html",
        contract=contract,
        clients=clients_rows,
        defaults=defaults,
        attachments=attachments,
        form_title="编辑合同",
    )


@app.post("/contracts/<int:contract_id>/delete")
@login_required
def delete_contract(contract_id):
    if not can_manage_contracts():
        abort(403)
    contract = require_contract_access(contract_id)
    work_order_count = db().execute(
        "select count(*) as count from service_orders where contract_id = ?",
        (contract_id,),
    ).fetchone()["count"]
    if work_order_count:
        flash("合同已有工单关联，不能删除；可以将状态改为已终止。", "error")
        return redirect(url_for("contract_detail", contract_id=contract_id))
    shutil.rmtree(os.path.join(CONTRACT_ATTACHMENTS_DIR, str(contract_id)), ignore_errors=True)
    db().execute("delete from contracts where id = ?", (contract_id,))
    log_action("delete", "contract", contract_id, contract["contract_number"], contract["title"])
    db().commit()
    flash("合同已删除。", "success")
    return redirect(url_for("contracts"))


@app.route("/contract-attachments/<int:attachment_id>")
@login_required
def preview_contract_attachment(attachment_id):
    attachment = db().execute(
        "select * from contract_attachments where id = ?",
        (attachment_id,),
    ).fetchone()
    if not attachment:
        abort(404)
    require_contract_access(attachment["contract_id"])
    return send_file(
        contract_attachment_path(attachment),
        mimetype=attachment["content_type"] or "application/octet-stream",
        download_name=attachment["original_filename"],
    )


@app.route("/contract-attachments/<int:attachment_id>/download")
@login_required
def download_contract_attachment(attachment_id):
    attachment = db().execute(
        "select * from contract_attachments where id = ?",
        (attachment_id,),
    ).fetchone()
    if not attachment:
        abort(404)
    require_contract_access(attachment["contract_id"])
    return send_file(
        contract_attachment_path(attachment),
        as_attachment=True,
        download_name=attachment["original_filename"],
    )


@app.post("/contract-attachments/<int:attachment_id>/delete")
@login_required
def delete_contract_attachment(attachment_id):
    if not can_manage_contracts():
        abort(403)
    attachment = db().execute(
        "select * from contract_attachments where id = ?",
        (attachment_id,),
    ).fetchone()
    if not attachment:
        abort(404)
    require_contract_access(attachment["contract_id"])
    try:
        os.remove(contract_attachment_path(attachment))
    except FileNotFoundError:
        pass
    db().execute("delete from contract_attachments where id = ?", (attachment_id,))
    db().commit()
    flash("合同附件已删除。", "success")
    return redirect(url_for("edit_contract", contract_id=attachment["contract_id"]))


@app.route("/clients", methods=["GET", "POST"])
@login_required
def clients():
    if normalized_role() in {"employee", "external_employee"}:
        abort(403)
    if is_external_manager() and not g.user["client_id"]:
        flash("外部用户尚未绑定客户。", "error")
        return redirect(url_for("dashboard"))
    if request.method == "POST":
        if is_external_user():
            abort(403)
        try:
            db().execute(
                """
                insert into clients (client_number, name, short_name, contact_name, email, address, country, created_at)
                values (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    next_client_number(),
                    request.form.get("name", "").strip(),
                    request.form.get("short_name", "").strip() or request.form.get("name", "").strip(),
                    request.form.get("contact_name", "").strip(),
                    request.form.get("email", "").strip(),
                    request.form.get("address", "").strip(),
                    request.form.get("country", "China").strip() or "China",
                    now(),
                ),
            )
            db().commit()
            flash("客户已创建。", "success")
        except sqlite3.IntegrityError:
            flash("客户编号重复，请重试。", "error")
        return redirect(url_for("clients"))
    q = request.args.get("q", "").strip()
    if is_external_manager():
        rows = db().execute("select * from clients where id = ?", (g.user["client_id"],)).fetchall()
    else:
        params = []
        where = ""
        if q:
            where = "where client_number like ? or name like ? or short_name like ? or contact_name like ? or email like ?"
            params = [f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"]
        rows = db().execute(f"select * from clients {where} order by client_number asc", params).fetchall()
    return render_template("clients.html", clients=rows, q=q)


@app.route("/clients/<int:client_id>/edit", methods=["GET", "POST"])
@login_required
def edit_client(client_id):
    if normalized_role() == "employee":
        abort(403)
    if is_external_user():
        abort(403)
    client = db().execute("select * from clients where id = ?", (client_id,)).fetchone()
    if not client:
        abort(404)
    if request.method == "POST":
        client_number = request.form.get("client_number", "").strip()
        if len(client_number) != 5 or not client_number.isdigit():
            flash("客户编号必须是 5 位数字。", "error")
            return redirect(url_for("edit_client", client_id=client_id))
        try:
            db().execute(
                """
                update clients
                set client_number = ?, name = ?, short_name = ?, contact_name = ?, email = ?, address = ?, country = ?
                where id = ?
                """,
                (
                    client_number,
                    request.form.get("name", "").strip(),
                    request.form.get("short_name", "").strip() or request.form.get("name", "").strip(),
                    request.form.get("contact_name", "").strip(),
                    request.form.get("email", "").strip(),
                    request.form.get("address", "").strip(),
                    request.form.get("country", "China").strip() or "China",
                    client_id,
                ),
            )
            db().commit()
            flash("客户资料已更新。", "success")
        except sqlite3.IntegrityError:
            flash("客户编号重复，请换一个编号。", "error")
            return redirect(url_for("edit_client", client_id=client_id))
        return redirect(url_for("clients"))
    return render_template("client_form.html", client=client)


@app.post("/clients/<int:client_id>/delete")
@login_required
def delete_client(client_id):
    if normalized_role() == "employee":
        abort(403)
    if is_external_user():
        abort(403)
    invoice_count = db().execute("select count(*) as count from invoices where client_id = ?", (client_id,)).fetchone()["count"]
    contract_count = db().execute("select count(*) as count from contracts where client_id = ?", (client_id,)).fetchone()["count"]
    if invoice_count or contract_count:
        flash("这个客户已有合同或发票记录，不能删除。可以编辑客户资料以保留历史记录。", "error")
        return redirect(url_for("clients"))
    db().execute("delete from clients where id = ?", (client_id,))
    db().commit()
    flash("客户已删除。", "success")
    return redirect(url_for("clients"))


@app.route("/owners", methods=["GET", "POST"])
@login_required
def owners():
    if not can_manage_owners():
        abort(403)
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        if not name:
            flash("请填写业主名称。", "error")
            return redirect(url_for("owners"))
        if db().execute("select 1 from owners where lower(trim(name)) = lower(trim(?))", (name,)).fetchone():
            flash("业主名称已存在。", "error")
            return redirect(url_for("owners"))
        try:
            db().execute(
                "insert into owners (owner_number, name, created_at) values (?, ?, ?)",
                (next_owner_number(), name, now()),
            )
            db().commit()
            flash("业主已创建。", "success")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("业主编号重复，请重试。", "error")
        return redirect(url_for("owners"))
    q = request.args.get("q", "").strip()
    params = []
    where = ""
    if q:
        where = "where owner_number like ? or name like ?"
        params = [f"%{q}%", f"%{q}%"]
    rows = db().execute(
        f"select * from owners {where} order by case when name = '未知' then 0 else 1 end, owner_number asc",
        params,
    ).fetchall()
    return render_template("owners.html", owners=rows, q=q)


@app.route("/owners/<int:owner_id>/edit", methods=["POST"])
@login_required
def edit_owner(owner_id):
    if not can_manage_owners():
        abort(403)
    owner = db().execute("select * from owners where id = ?", (owner_id,)).fetchone()
    if not owner:
        abort(404)
    owner_number = request.form.get("owner_number", "").strip().upper()
    name = request.form.get("name", "").strip()
    if not owner_number or not name:
        flash("请填写业主编号和名称。", "error")
        return redirect(url_for("owners"))
    if db().execute(
        "select 1 from owners where id != ? and lower(trim(name)) = lower(trim(?))",
        (owner_id, name),
    ).fetchone():
        flash("业主名称已存在。", "error")
        return redirect(url_for("owners"))
    try:
        db().execute(
            "update owners set owner_number = ?, name = ? where id = ?",
            (owner_number, name, owner_id),
        )
        db().execute(
            "update buyers set owner = ? where owner_id = ?",
            (name, owner_id),
        )
        db().commit()
        flash("业主资料已更新。", "success")
    except sqlite3.IntegrityError:
        db().rollback()
        flash("业主编号重复，请换一个编号。", "error")
    return redirect(url_for("owners"))


@app.post("/owners/<int:owner_id>/delete")
@login_required
def delete_owner(owner_id):
    if not can_manage_owners():
        abort(403)
    owner = db().execute("select * from owners where id = ?", (owner_id,)).fetchone()
    if not owner:
        abort(404)
    if owner["name"] == "未知":
        flash("默认业主不能删除。", "error")
        return redirect(url_for("owners"))
    used = db().execute(
        "select count(*) as count from buyers where owner_id = ?",
        (owner_id,),
    ).fetchone()["count"]
    if used:
        flash("这个业主已有站点，不能删除。可以编辑业主资料。", "error")
        return redirect(url_for("owners"))
    db().execute("delete from owners where id = ?", (owner_id,))
    db().commit()
    flash("业主已删除。", "success")
    return redirect(url_for("owners"))


@app.route("/buyers", methods=["GET", "POST"])
@login_required
def buyers():
    if not can_manage_buyers():
        abort(403)
    owners_rows = owner_options()
    countries_rows = country_rows()
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        country = country_from_form()
        owner = None
        owner_id = int(request.form["owner_id"]) if request.form.get("owner_id", "").isdigit() else None
        if owner_id:
            owner = db().execute("select * from owners where id = ?", (owner_id,)).fetchone()
            if not owner:
                flash("请选择有效的业主。", "error")
                return redirect(url_for("buyers"))
        else:
            owner = unknown_owner()
            owner_id = owner["id"]
        detailed_address = request.form.get("detailed_address", "").strip()
        if not name or not detailed_address:
            flash("请填写站点名称和详细地址。", "error")
            return redirect(url_for("buyers"))
        client_id = g.user["client_id"] if is_external_manager() else (
            int(request.form["client_id"]) if request.form.get("client_id", "").isdigit() else None
        )
        existing_site = db().execute(
            """
            select buyer_number from buyers
            where coalesce(client_id, 0) = coalesce(?, 0)
              and lower(trim(name)) = lower(trim(?))
              and lower(trim(detailed_address)) = lower(trim(?))
            """,
            (client_id, name, detailed_address),
        ).fetchone()
        if existing_site:
            flash(f"站点已存在，编号：{existing_site['buyer_number']}。", "error")
            return redirect(url_for("buyers"))
        try:
            db().execute(
                """
                insert into buyers (
                    buyer_number, client_id, country, country_code, name, owner_id, owner, contact_name, contact_details,
                    email, site_size, detailed_address, equipment_manufacturer, created_at
                ) values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    next_buyer_number(),
                    client_id,
                    country["code"],
                    country["code"],
                    name,
                    owner_id,
                    owner["name"] if owner else "",
                    request.form.get("contact_name", "").strip(),
                    request.form.get("contact_details", "").strip(),
                    request.form.get("email", "").strip(),
                    request.form.get("site_size", "").strip(),
                    detailed_address,
                    request.form.get("equipment_manufacturer", "").strip(),
                    now(),
                ),
            )
            db().commit()
            flash("站点已创建。", "success")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("站点编号重复，请重试。", "error")
        return redirect(url_for("buyers"))
    q = request.args.get("q", "").strip()
    sort = request.args.get("sort", "owner_name")
    direction = request.args.get("direction", "asc").lower()
    buyer_sort_columns = {
        "buyer_number": "buyers.buyer_number",
        "country_name": "coalesce(country_local.name, country_zh.name, country_en.name, buyers.country, buyers.country_code)",
        "name": "buyers.name",
        "owner_name": "coalesce(owners.name, buyers.owner)",
        "contact_name": "buyers.contact_name",
        "contact_details": "buyers.contact_details",
        "email": "buyers.email",
        "site_size": "buyers.site_size",
        "detailed_address": "buyers.detailed_address",
    }
    if sort not in buyer_sort_columns:
        sort = "owner_name"
    if direction not in {"asc", "desc"}:
        direction = "asc"
    order_direction = "desc" if direction == "desc" else "asc"
    order_by = f"lower(coalesce({buyer_sort_columns[sort]}, '')) {order_direction}, buyers.name asc, buyers.buyer_number asc"
    params = []
    clauses = []
    if is_external_manager():
        clauses.append("buyers.client_id = ?")
        params.append(g.user["client_id"])
    if q:
        clauses.append(
            """(
            buyers.buyer_number like ? or buyers.name like ? or coalesce(owners.name, buyers.owner) like ? or buyers.contact_name like ?
            or buyers.contact_details like ? or buyers.email like ? or buyers.site_size like ? or buyers.detailed_address like ? or buyers.equipment_manufacturer like ?
            )"""
        )
        params = [f"%{q}%"] * 9
        if is_external_manager():
            params.insert(0, g.user["client_id"])
    where = f"where {' and '.join(clauses)}" if clauses else ""
    rows = db().execute(
        f"""
        select buyers.*, coalesce(owners.name, buyers.owner) as owner_name,
               owners.owner_number as owner_number,
               coalesce(country_local.name, country_zh.name, country_en.name, buyers.country, buyers.country_code) as country_name
        from buyers
        left join owners on owners.id = buyers.owner_id
        left join country_translations country_local
          on country_local.country_code = buyers.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = buyers.country_code and country_zh.language_code = 'zh-CN'
        left join country_translations country_en
          on country_en.country_code = buyers.country_code and country_en.language_code = 'en'
        {where}
        order by {order_by}
        """,
        [current_language(), *params],
    ).fetchall()
    clients_rows = db().execute("select id, client_number, name from clients order by client_number").fetchall()
    return render_template(
        "buyers.html",
        buyers=rows,
        clients=clients_rows,
        owners=owners_rows,
        countries=countries_rows,
        q=q,
        sort=sort,
        direction=direction,
    )


@app.route("/buyers/<int:buyer_id>/edit", methods=["GET", "POST"])
@login_required
def edit_buyer(buyer_id):
    if not can_manage_buyers():
        abort(403)
    buyer = db().execute("select * from buyers where id = ?", (buyer_id,)).fetchone()
    if not buyer:
        abort(404)
    if not can_access_buyer(buyer):
        abort(403)
    if request.method == "POST":
        buyer_number = request.form.get("buyer_number", "").strip().upper()
        name = request.form.get("name", "").strip()
        country = country_from_form(buyer["country_code"] or "US")
        owner = None
        owner_id = int(request.form["owner_id"]) if request.form.get("owner_id", "").isdigit() else None
        if owner_id:
            owner = db().execute("select * from owners where id = ?", (owner_id,)).fetchone()
            if not owner:
                flash("请选择有效的业主。", "error")
                return redirect(url_for("edit_buyer", buyer_id=buyer_id))
        else:
            owner = unknown_owner()
            owner_id = owner["id"]
        detailed_address = request.form.get("detailed_address", "").strip()
        if not buyer_number or not name or not detailed_address:
            flash("请填写编号、站点名称和详细地址。", "error")
            return redirect(url_for("edit_buyer", buyer_id=buyer_id))
        address_changed = normalized_address(buyer["detailed_address"]) != normalized_address(detailed_address)
        try:
            db().execute(
                """
                update buyers
                set buyer_number = ?, client_id = ?, country = ?, country_code = ?, name = ?, owner_id = ?, owner = ?, contact_name = ?,
                    contact_details = ?, email = ?, site_size = ?, detailed_address = ?, equipment_manufacturer = ?
                where id = ?
                """,
                (
                    buyer_number,
                    g.user["client_id"] if is_external_manager() else (
                        int(request.form["client_id"]) if request.form.get("client_id", "").isdigit() else buyer["client_id"]
                    ),
                    country["code"],
                    country["code"],
                    name,
                    owner_id,
                    owner["name"] if owner else "",
                    request.form.get("contact_name", "").strip(),
                    request.form.get("contact_details", "").strip(),
                    request.form.get("email", "").strip(),
                    request.form.get("site_size", "").strip(),
                    detailed_address,
                    request.form.get("equipment_manufacturer", "").strip(),
                    buyer_id,
                ),
            )
            db().execute(
                """
                update service_orders
                set client_name = ?, buyer_contact_name = ?, buyer_contact_details = ?
                where buyer_id = ?
                """,
                (
                    name,
                    request.form.get("contact_name", "").strip(),
                    request.form.get("contact_details", "").strip(),
                    buyer_id,
                ),
            )
            if address_changed:
                db().execute(
                    """
                    update buyers
                    set latitude = null, longitude = null, geocode_address = null,
                        geocode_status = 'pending', geocode_attempted_at = null, geocode_version = null
                    where id = ?
                    """,
                    (buyer_id,),
                )
            db().commit()
            flash("站点资料已更新。", "success")
            return redirect(url_for("buyers"))
        except sqlite3.IntegrityError:
            db().rollback()
            flash("站点编号重复，请换一个编号。", "error")
    clients_rows = db().execute("select id, client_number, name from clients order by client_number").fetchall()
    owners_rows = owner_options()
    return render_template(
        "buyer_form.html",
        buyer=buyer,
        clients=clients_rows,
        owners=owners_rows,
        countries=country_rows(),
    )


@app.post("/buyers/<int:buyer_id>/delete")
@login_required
def delete_buyer(buyer_id):
    if not can_manage_buyers():
        abort(403)
    buyer = db().execute("select * from buyers where id = ?", (buyer_id,)).fetchone()
    if not buyer:
        abort(404)
    if not can_access_buyer(buyer):
        abort(403)
    used = db().execute(
        "select count(*) as count from service_orders where buyer_id = ?",
        (buyer_id,),
    ).fetchone()["count"]
    if used:
        flash("这个站点已有工单，不能删除。可以编辑站点资料。", "error")
        return redirect(url_for("buyers"))
    db().execute("delete from buyers where id = ?", (buyer_id,))
    db().commit()
    flash("站点已删除。", "success")
    return redirect(url_for("buyers"))


@app.route("/work-order-types", methods=["GET", "POST"])
@login_required
def work_order_types():
    if request.method == "POST":
        code = request.form.get("code", "").strip().upper()
        name = request.form.get("name", "").strip()
        if not code or not name:
            flash("请填写工单类型编码和名称。", "error")
            return redirect(url_for("work_order_types"))
        try:
            db().execute(
                """
                insert into work_order_types (code, name, description, is_active, created_at)
                values (?, ?, ?, 1, ?)
                """,
                (code, name, request.form.get("description", "").strip(), now()),
            )
            db().commit()
            flash("工单类型已创建。", "success")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("工单类型编码重复。", "error")
        return redirect(url_for("work_order_types"))
    rows = db().execute("select * from work_order_types order by is_active desc, code").fetchall()
    return render_template("work_order_types.html", work_order_types=rows)


@app.route("/work-order-types/<int:type_id>/edit", methods=["GET", "POST"])
@login_required
def edit_work_order_type(type_id):
    work_order_type = db().execute("select * from work_order_types where id = ?", (type_id,)).fetchone()
    if not work_order_type:
        abort(404)
    if request.method == "POST":
        try:
            db().execute(
                """
                update work_order_types
                set code = ?, name = ?, description = ?, is_active = ?
                where id = ?
                """,
                (
                    request.form.get("code", "").strip().upper(),
                    request.form.get("name", "").strip(),
                    request.form.get("description", "").strip(),
                    1 if request.form.get("is_active", "1") == "1" else 0,
                    type_id,
                ),
            )
            db().commit()
            flash("工单类型已更新。", "success")
            return redirect(url_for("work_order_types"))
        except sqlite3.IntegrityError:
            db().rollback()
            flash("工单类型编码重复。", "error")
    return render_template("work_order_type_form.html", work_order_type=work_order_type)


@app.post("/work-order-types/<int:type_id>/delete")
@login_required
def delete_work_order_type(type_id):
    used = db().execute(
        "select count(*) as count from service_orders where work_order_type_id = ?",
        (type_id,),
    ).fetchone()["count"]
    if used:
        flash("这个工单类型已有工单使用，不能删除，可以停用。", "error")
        return redirect(url_for("work_order_types"))
    db().execute("delete from work_order_types where id = ?", (type_id,))
    db().commit()
    flash("工单类型已删除。", "success")
    return redirect(url_for("work_order_types"))


@app.route("/projects", methods=["GET", "POST"])
@login_required
def projects():
    if request.method == "POST":
        name = normalized_project_name(request.form.get("name"))
        project_type = request.form.get("project_type", "invoice")
        if project_type not in PROJECT_TYPE_LABELS:
            project_type = "invoice"
        if not name:
            flash("项目名称不能为空。", "error")
            return redirect(url_for("projects"))
        if project_name_exists(project_type, name):
            flash("同一项目类型下已存在同名项目，不能重复创建。", "error")
            return redirect(url_for("projects"))
        try:
            db().execute(
                """
                insert into projects (name, name_key, project_type, default_amount, unit_price, tax_rate, is_active, created_at)
                values (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    name,
                    project_name_key(name),
                    project_type,
                    to_float(request.form.get("default_amount")) if project_type == "invoice" else 0,
                    to_float(request.form.get("unit_price")),
                    to_float(request.form.get("tax_rate")) if project_type == "invoice" else 0,
                    1 if request.form.get("is_active", "1") == "1" else 0,
                    now(),
                ),
            )
            db().commit()
            flash("项目已创建。", "success")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("同一项目类型下已存在同名项目，不能重复创建。", "error")
        return redirect(url_for("projects"))
    merge_duplicate_projects(db())
    db().commit()
    rows = db().execute("select * from projects order by is_active desc, name").fetchall()
    return render_template("projects.html", projects=rows)


@app.route("/projects/<int:project_id>/edit", methods=["GET", "POST"])
@login_required
def edit_project(project_id):
    project = db().execute("select * from projects where id = ?", (project_id,)).fetchone()
    if not project:
        abort(404)
    if request.method == "POST":
        name = normalized_project_name(request.form.get("name"))
        project_type = request.form.get("project_type", "invoice")
        if project_type not in PROJECT_TYPE_LABELS:
            project_type = "invoice"
        if not name:
            flash("项目名称不能为空。", "error")
            return redirect(url_for("edit_project", project_id=project_id))
        if project_name_exists(project_type, name, excluded_project_id=project_id):
            flash("同一项目类型下已存在同名项目，不能重复保存。", "error")
            return redirect(url_for("edit_project", project_id=project_id))
        invoice_used = db().execute(
            "select count(*) as count from invoice_items where project_id = ?",
            (project_id,),
        ).fetchone()["count"]
        expense_used = db().execute(
            "select count(*) as count from expense_items where project_id = ?",
            (project_id,),
        ).fetchone()["count"]
        if project_type != project["project_type"] and (invoice_used or expense_used):
            flash("已被发票或报销使用的项目不能切换类型，可以新建另一个项目。", "error")
            return redirect(url_for("edit_project", project_id=project_id))
        try:
            db().execute(
                """
                update projects set name = ?, name_key = ?, project_type = ?, default_amount = ?, unit_price = ?, tax_rate = ?, is_active = ?
                where id = ?
                """,
                (
                    name,
                    project_name_key(name),
                    project_type,
                    to_float(request.form.get("default_amount")) if project_type == "invoice" else 0,
                    to_float(request.form.get("unit_price")),
                    to_float(request.form.get("tax_rate")) if project_type == "invoice" else 0,
                    1 if request.form.get("is_active", "1") == "1" else 0,
                    project_id,
                ),
            )
            db().commit()
            flash("项目已更新。", "success")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("同一项目类型下已存在同名项目，不能重复保存。", "error")
        return redirect(url_for("projects"))
    return render_template("project_form.html", project=project)


@app.post("/projects/<int:project_id>/delete")
@login_required
def delete_project(project_id):
    invoice_used = db().execute("select count(*) as count from invoice_items where project_id = ?", (project_id,)).fetchone()["count"]
    expense_used = db().execute("select count(*) as count from expense_items where project_id = ?", (project_id,)).fetchone()["count"]
    if invoice_used or expense_used:
        flash("这个项目已有发票或报销记录，不能删除。可以停用该项目。", "error")
        return redirect(url_for("projects"))
    db().execute("delete from projects where id = ?", (project_id,))
    db().commit()
    flash("项目已删除。", "success")
    return redirect(url_for("projects"))


@app.route("/countries", methods=["GET", "POST"])
@login_required
def countries():
    if request.method == "POST":
        code = request.form.get("code", "").strip().upper()
        region_code = request.form.get("region_code", "").strip().lower()
        if not re.fullmatch(r"[A-Z]{2,3}", code) or not re.fullmatch(r"[a-z][a-z0-9_-]*", region_code):
            flash("国家代码应为 2-3 位大写字母，区域代码应使用小写字母、数字、下划线或连字符。", "error")
            return redirect(url_for("countries"))
        translations = []
        seen_languages = set()
        for language_code, name, region_name in zip(
            request.form.getlist("language_code"),
            request.form.getlist("translation_name"),
            request.form.getlist("translation_region_name"),
        ):
            language_code = language_code.strip()
            name = name.strip()
            region_name = region_name.strip()
            if not language_code or not name or not region_name or language_code in seen_languages:
                continue
            seen_languages.add(language_code)
            translations.append((language_code, name, region_name))
        if not translations:
            flash("请至少填写一种语言的国家名称和区域名称。", "error")
            return redirect(url_for("countries"))
        db().execute(
            """
            insert into countries (code, region_code, is_active, sort_order, created_at)
            values (?, ?, ?, ?, ?)
            on conflict(code) do update set
                region_code = excluded.region_code,
                is_active = excluded.is_active,
                sort_order = excluded.sort_order
            """,
            (
                code,
                region_code,
                1 if request.form.get("is_active", "1") == "1" else 0,
                int(request.form.get("sort_order") or 0),
                now(),
            ),
        )
        db().execute("delete from country_translations where country_code = ?", (code,))
        db().executemany(
            """
            insert into country_translations (country_code, language_code, name, region_name)
            values (?, ?, ?, ?)
            """,
            [(code, language_code, name, region_name) for language_code, name, region_name in translations],
        )
        db().commit()
        flash("国家配置已保存。", "success")
        return redirect(url_for("countries"))
    rows = country_rows(include_inactive=True)
    translations = {country["code"]: country_translations(country["code"]) for country in rows}
    return render_template("countries.html", countries=rows, translations=translations)


@app.route("/employee-grades", methods=["GET", "POST"])
@login_required
def employee_grades():
    if not can_manage_employee_grades():
        abort(403)
    if request.method == "POST":
        grade_id = request.form.get("grade_id")
        grade_name = request.form.get("grade_name", "").strip()
        if not grade_name:
            flash("请填写员工等级。", "error")
            return redirect(url_for("employee_grades"))
        values = (
            grade_name,
            to_float(request.form.get("base_salary")),
            to_float(request.form.get("standard_hourly_rate")),
            to_float(request.form.get("transport_hourly_rate")),
            to_float(request.form.get("overtime_hourly_rate")),
            to_float(request.form.get("holiday_hourly_rate")),
            1 if request.form.get("is_active", "1") == "1" else 0,
        )
        try:
            if grade_id and str(grade_id).isdigit():
                db().execute(
                    """
                    update employee_grades
                    set grade_name = ?, base_salary = ?, standard_hourly_rate = ?, transport_hourly_rate = ?,
                        overtime_hourly_rate = ?, holiday_hourly_rate = ?, is_active = ?
                    where id = ?
                    """,
                    (*values, int(grade_id)),
                )
                log_action("update", "employee_grade", int(grade_id), grade_name, "修改员工等级")
            else:
                cursor = db().execute(
                    """
                    insert into employee_grades (
                        grade_name, base_salary, standard_hourly_rate, transport_hourly_rate,
                        overtime_hourly_rate, holiday_hourly_rate, is_active, created_at
                    ) values (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (*values, now()),
                )
                log_action("create", "employee_grade", cursor.lastrowid, grade_name, "新增员工等级")
            db().commit()
            flash("员工等级已保存。", "success")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("员工等级名称已存在。", "error")
        return redirect(url_for("employee_grades"))
    return render_template("employee_grades.html", grades=employee_grade_options(include_inactive=True))


@app.route("/payroll/subsidies", methods=["GET", "POST"])
@login_required
def payroll_subsidies():
    if not can_manage_employee_grades():
        abort(403)
    if request.method == "POST":
        car_method = request.form.get("car_allowance_method", "daily")
        meal_method = request.form.get("meal_allowance_method", "daily")
        if car_method not in {"daily", "mileage"}:
            car_method = "daily"
        if meal_method not in {"daily", "provided"}:
            meal_method = "daily"
        values = {
            "payroll_cycle_start": request.form.get("cycle_start") or "2026-07-06",
            "payroll_car_allowance_method": car_method,
            "payroll_car_daily_amount": max(to_float(request.form.get("car_daily_amount")), 0),
            "payroll_car_mileage_rate": max(to_float(request.form.get("car_mileage_rate")), 0),
            "payroll_meal_allowance_method": meal_method,
            "payroll_meal_daily_amount": max(to_float(request.form.get("meal_daily_amount")), 0),
            "payroll_lodging_limit": max(to_float(request.form.get("lodging_limit")), 0),
        }
        try:
            date.fromisoformat(values["payroll_cycle_start"])
        except ValueError:
            values["payroll_cycle_start"] = "2026-07-06"
        for key, value in values.items():
            set_setting(key, str(value))
        log_action("update", "payroll_subsidies", None, "薪酬补贴参数", "修改薪酬补贴参数")
        db().commit()
        flash("补贴参数已保存。", "success")
        return redirect(url_for("payroll_subsidies"))
    return render_template("payroll_subsidies.html", settings=payroll_subsidy_settings())


@app.route("/users", methods=["GET", "POST"])
@login_required
def users():
    if request.method == "POST" and not can_assign_external_employees():
        abort(403)
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        name = request.form.get("name", "").strip() or email
        address = request.form.get("address", "").strip()
        password = request.form.get("password", "")
        role = request.form.get("role", "employee")
        country = country_from_form()
        employee_grade_id = (
            int(request.form["employee_grade_id"])
            if can_manage_users() and request.form.get("employee_grade_id", "").isdigit()
            else None
        )
        if is_external_manager():
            role = "external_employee"
        if role not in ROLE_OPTIONS:
            role = "employee"
        client_id = (
            g.user["client_id"]
            if is_external_manager()
            else int(request.form["client_id"]) if role in {"external_manager", "external_employee"} and request.form.get("client_id", "").isdigit()
            else None
        )
        if "�" in name:
            flash("姓名包含损坏字符，请重新输入正确姓名。", "error")
            return redirect(url_for("users"))
        if requires_user_address(role) and not address:
            flash("内部员工必须填写地址。", "error")
            return redirect(url_for("users"))
        if len(password) < 8:
            flash("密码至少需要 8 位。", "error")
            return redirect(url_for("users"))
        try:
            cursor = db().execute(
                """
                insert into users (
                    name, email, password_hash, role, address, employee_grade_id, client_id, region_code, country_code, created_at
                )
                values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    name,
                    email,
                    generate_password_hash(password),
                    role,
                    address,
                    employee_grade_id,
                    client_id,
                    country["region_code"],
                    country["code"],
                    now(),
                ),
            )
            user_id = cursor.lastrowid
            save_user_service_order_assignments(user_id, role)
            for uploaded in uploaded_attachments_from_request():
                save_named_attachment(
                    uploaded,
                    os.path.join(USER_ATTACHMENT_DIR, str(user_id)),
                    "user_attachments",
                    "user_id",
                    user_id,
                )
            db().commit()
            flash("用户已创建。", "success")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("这个邮箱已经存在。", "error")
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
        return redirect(url_for("users"))
    if can_manage_users():
        rows = db().execute(
            """
            select users.*, clients.name as client_name,
                   employee_grades.grade_name as employee_grade_name,
                   coalesce(country_local.name, country_zh.name, users.country_code) as country_name,
                   coalesce(country_local.region_name, country_zh.region_name, users.region_code) as region_name
            from users left join clients on clients.id = users.client_id
            left join employee_grades on employee_grades.id = users.employee_grade_id
            left join country_translations country_local
              on country_local.country_code = users.country_code and country_local.language_code = ?
            left join country_translations country_zh
              on country_zh.country_code = users.country_code and country_zh.language_code = 'zh-CN'
            order by users.created_at desc
            """,
            (current_language(),),
        ).fetchall()
    elif normalized_role() == "manager":
        rows = db().execute(
            """
            select users.*, clients.name as client_name,
                   employee_grades.grade_name as employee_grade_name,
                   coalesce(country_local.name, country_zh.name, users.country_code) as country_name,
                   coalesce(country_local.region_name, country_zh.region_name, users.region_code) as region_name
            from users left join clients on clients.id = users.client_id
            left join employee_grades on employee_grades.id = users.employee_grade_id
            left join country_translations country_local
              on country_local.country_code = users.country_code and country_local.language_code = ?
            left join country_translations country_zh
              on country_zh.country_code = users.country_code and country_zh.language_code = 'zh-CN'
            where users.id = ? or users.role in ('employee', 'external_employee')
            order by users.is_active asc, users.created_at desc
            """,
            (current_language(), g.user["id"]),
        ).fetchall()
    elif is_external_manager():
        rows = db().execute(
            """
            select users.*, clients.name as client_name,
                   employee_grades.grade_name as employee_grade_name,
                   coalesce(country_local.name, country_zh.name, users.country_code) as country_name,
                   coalesce(country_local.region_name, country_zh.region_name, users.region_code) as region_name
            from users left join clients on clients.id = users.client_id
            left join employee_grades on employee_grades.id = users.employee_grade_id
            left join country_translations country_local
              on country_local.country_code = users.country_code and country_local.language_code = ?
            left join country_translations country_zh
              on country_zh.country_code = users.country_code and country_zh.language_code = 'zh-CN'
            where users.id = ?
               or (users.role = 'external_employee' and users.client_id = ?)
            order by users.created_at desc
            """,
            (current_language(), g.user["id"], g.user["client_id"]),
        ).fetchall()
    else:
        rows = db().execute(
            """
            select users.*, null as client_name,
                   employee_grades.grade_name as employee_grade_name,
                   coalesce(country_local.name, country_zh.name, users.country_code) as country_name,
                   coalesce(country_local.region_name, country_zh.region_name, users.region_code) as region_name
            from users
            left join employee_grades on employee_grades.id = users.employee_grade_id
            left join country_translations country_local
              on country_local.country_code = users.country_code and country_local.language_code = ?
            left join country_translations country_zh
              on country_zh.country_code = users.country_code and country_zh.language_code = 'zh-CN'
            where users.id = ?
            """,
            (current_language(), g.user["id"]),
        ).fetchall()
    clients_rows = (
        db().execute("select id, client_number, name from clients where id = ?", (g.user["client_id"],)).fetchall()
        if is_external_manager()
        else db().execute("select id, client_number, name from clients order by client_number").fetchall()
    )
    service_orders_rows = (
        db().execute(
            "select id, order_number, client_name, status from service_orders where client_id = ? order by status != 'closed' desc, created_at desc, id desc",
            (g.user["client_id"],),
        ).fetchall()
        if is_external_manager()
        else db().execute("select id, order_number, client_name, status from service_orders order by status != 'closed' desc, created_at desc, id desc").fetchall()
    )
    return render_template(
        "users.html",
        users=rows,
        clients=clients_rows,
        service_orders=service_orders_rows,
        user_attachments={user["id"]: get_user_attachments(user["id"]) for user in rows},
        user_order_ids={user["id"]: assigned_service_order_ids(user["id"]) for user in rows},
        role_options=ROLE_OPTIONS,
        can_manage=can_manage_users(),
        can_assign=can_assign_external_employees(),
        can_approve=can_approve_users(),
        countries=country_rows(),
        employee_grades=employee_grade_options(),
    )


@app.route("/users/<int:user_id>/edit", methods=["GET", "POST"])
@login_required
def edit_user(user_id):
    if not can_manage_users() and user_id != g.user["id"] and not is_external_manager():
        abort(403)
    user = db().execute(
        """
        select users.*, employee_grades.grade_name as employee_grade_name
        from users
        left join employee_grades on employee_grades.id = users.employee_grade_id
        where users.id = ?
        """,
        (user_id,),
    ).fetchone()
    if not user:
        abort(404)
    if is_external_manager() and user_id != g.user["id"]:
        if normalized_role(user["role"]) != "external_employee" or user["client_id"] != g.user["client_id"]:
            abort(403)
    if request.method == "POST":
        email = request.form.get("email", user["email"]).strip().lower() if can_manage_users() else user["email"]
        name = request.form.get("name", "").strip() or email
        role = request.form.get("role", normalized_role(user["role"])) if can_manage_users() else user["role"]
        address = request.form.get("address", "").strip()
        if role not in ROLE_OPTIONS and can_manage_users():
            role = "employee"
        client_id = (
            user["client_id"]
            if not can_manage_users()
            else int(request.form["client_id"]) if role in {"external_manager", "external_employee"} and request.form.get("client_id", "").isdigit()
            else None
        )
        employee_grade_id = (
            int(request.form["employee_grade_id"])
            if can_manage_users() and request.form.get("employee_grade_id", "").isdigit()
            else user["employee_grade_id"]
        )
        password = request.form.get("password", "")
        country = country_from_form(user["country_code"]) if can_assign_external_employees() else country_by_code(
            user["country_code"], include_inactive=True
        )
        admin_count = db().execute("select count(*) as count from users where role = 'admin'").fetchone()["count"]
        if "�" in name:
            flash("姓名包含损坏字符，请重新输入正确姓名。", "error")
            return redirect(url_for("edit_user", user_id=user_id))
        if can_manage_users() and user["role"] == "admin" and role != "admin" and admin_count <= 1:
            flash("至少需要保留一个管理员。", "error")
            return redirect(url_for("edit_user", user_id=user_id))
        if requires_user_address(role) and not address:
            flash("内部员工必须填写地址。", "error")
            return redirect(url_for("edit_user", user_id=user_id))
        try:
            db().execute(
                """
                update users
                set name = ?, email = ?, role = ?, address = ?, employee_grade_id = ?, client_id = ?, region_code = ?, country_code = ?
                where id = ?
                """,
                (name, email, role, address, employee_grade_id, client_id, country["region_code"], country["code"], user_id),
            )
            if can_assign_external_employees():
                save_user_service_order_assignments(user_id, role)
            if password:
                if len(password) < 8:
                    flash("新密码至少需要 8 位。", "error")
                    return redirect(url_for("edit_user", user_id=user_id))
                db().execute("update users set password_hash = ? where id = ?", (generate_password_hash(password), user_id))
            for uploaded in uploaded_attachments_from_request():
                save_named_attachment(
                    uploaded,
                    os.path.join(USER_ATTACHMENT_DIR, str(user_id)),
                    "user_attachments",
                    "user_id",
                    user_id,
                )
            db().commit()
            flash("用户资料已更新。", "success")
        except sqlite3.IntegrityError:
            db().rollback()
            flash("这个邮箱已经存在。", "error")
            return redirect(url_for("edit_user", user_id=user_id))
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("edit_user", user_id=user_id))
        return redirect(url_for("users"))
    clients_rows = (
        db().execute("select id, client_number, name from clients where id = ?", (g.user["client_id"],)).fetchall()
        if is_external_manager()
        else db().execute("select id, client_number, name from clients order by client_number").fetchall()
    )
    service_orders_rows = (
        db().execute(
            "select id, order_number, client_name, status from service_orders where client_id = ? order by status != 'closed' desc, created_at desc, id desc",
            (g.user["client_id"],),
        ).fetchall()
        if is_external_manager()
        else db().execute("select id, order_number, client_name, status from service_orders order by status != 'closed' desc, created_at desc, id desc").fetchall()
    )
    return render_template(
        "user_form.html",
        user=user,
        clients=clients_rows,
        service_orders=service_orders_rows,
        selected_order_ids=assigned_service_order_ids(user_id),
        attachments=get_user_attachments(user_id),
        role_options=ROLE_OPTIONS,
        can_manage=can_manage_users(),
        can_assign=can_assign_external_employees(),
        can_approve=can_approve_users(),
        countries=country_rows(),
        employee_grades=employee_grade_options(),
    )


@app.post("/users/<int:user_id>/status")
@login_required
def update_user_status(user_id):
    if not can_approve_users():
        abort(403)
    user = db().execute("select * from users where id = ?", (user_id,)).fetchone()
    if not user:
        abort(404)
    if normalized_role(user["role"]) not in {"employee", "external_employee"}:
        abort(403)
    is_active = 1 if request.form.get("is_active") == "1" else 0
    db().execute("update users set is_active = ? where id = ?", (is_active, user_id))
    if is_active:
        create_message(
            user_id,
            "账号已启用",
            f"{g.user['name']} 已批准并启用你的账号。",
        )
    db().commit()
    flash("用户状态已更新。", "success")
    return redirect(url_for("users"))


@app.get("/user-attachments/<int:attachment_id>/download")
@login_required
def download_user_attachment(attachment_id):
    attachment = db().execute("select * from user_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    owner = db().execute("select * from users where id = ?", (attachment["user_id"],)).fetchone()
    if not owner or not can_manage_user_record(owner):
        abort(403)
    return send_file(
        os.path.join(USER_ATTACHMENT_DIR, str(attachment["user_id"]), attachment["stored_filename"]),
        as_attachment=True,
        download_name=attachment["original_filename"],
    )


@app.get("/user-attachments/<int:attachment_id>/preview")
@login_required
def preview_user_attachment(attachment_id):
    attachment = db().execute("select * from user_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    owner = db().execute("select * from users where id = ?", (attachment["user_id"],)).fetchone()
    if not owner or not can_manage_user_record(owner):
        abort(403)
    return send_file(
        os.path.join(USER_ATTACHMENT_DIR, str(attachment["user_id"]), attachment["stored_filename"]),
        as_attachment=False,
        download_name=attachment["original_filename"],
        mimetype=attachment["content_type"] or None,
        conditional=True,
    )


@app.post("/user-attachments/<int:attachment_id>/delete")
@login_required
def delete_user_attachment(attachment_id):
    attachment = db().execute("select * from user_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    owner = db().execute("select * from users where id = ?", (attachment["user_id"],)).fetchone()
    if not owner or not can_manage_user_record(owner):
        abort(403)
    try:
        os.remove(os.path.join(USER_ATTACHMENT_DIR, str(attachment["user_id"]), attachment["stored_filename"]))
    except FileNotFoundError:
        pass
    db().execute("delete from user_attachments where id = ?", (attachment_id,))
    db().commit()
    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify({"ok": True, "attachment_id": attachment_id})
    flash("用户附件已删除。", "success")
    return redirect(url_for("users"))


@app.post("/users/<int:user_id>/delete")
@login_required
def delete_user(user_id):
    if user_id == g.user["id"]:
        flash("不能删除当前登录用户。", "error")
        return redirect(url_for("users"))
    user = db().execute("select * from users where id = ?", (user_id,)).fetchone()
    if not user:
        abort(404)
    admin_count = db().execute("select count(*) as count from users where role = 'admin'").fetchone()["count"]
    if user["role"] == "admin" and admin_count <= 1:
        flash("至少需要保留一个管理员。", "error")
        return redirect(url_for("users"))
    invoice_count = db().execute("select count(*) as count from invoices where created_by = ?", (user_id,)).fetchone()["count"]
    if invoice_count:
        flash("这个用户已经创建过发票，不能删除。", "error")
        return redirect(url_for("users"))
    db().execute("delete from messages where user_id = ?", (user_id,))
    db().execute("delete from users where id = ?", (user_id,))
    db().commit()
    flash("用户已删除。", "success")
    return redirect(url_for("users"))


@app.route("/settings/system", methods=["GET", "POST"])
@login_required
def system_settings():
    if request.method == "POST":
        set_setting("company_tax_note", request.form.get("company_tax_note", "").strip())
        for key in DEFAULT_PAYMENT_INSTRUCTIONS:
            set_setting(f"payment_{key}", request.form.get(f"payment_{key}", "").strip())
        for key in DEFAULT_SMTP_SETTINGS:
            set_setting(f"smtp_{key}", request.form.get(f"smtp_{key}", "").strip())
        set_setting("invoice_terms", request.form.get("invoice_terms", "").strip())
        set_setting(
            "google_maps_browser_api_key",
            request.form.get("google_maps_browser_api_key", "").strip(),
        )
        set_setting(
            "google_geocoding_api_key",
            request.form.get("google_geocoding_api_key", "").strip(),
        )
        db().commit()
        flash("系统设置已保存。", "success")
        return redirect(url_for("system_settings"))
    return render_template(
        "company_settings.html",
        company=get_company_profile(),
        payment=get_payment_instructions(),
        smtp=get_smtp_settings(),
        terms=get_invoice_terms(),
        google_maps_browser_api_key=get_google_maps_browser_api_key(),
        google_geocoding_api_key=get_google_geocoding_api_key(),
    )


@app.route("/settings/menu-permissions", methods=["GET", "POST"])
@admin_required
def menu_permissions():
    if request.method == "POST":
        menu_keys = [item["key"] for group in MENU_PERMISSION_GROUPS for item in group["items"]]
        action_pairs = [
            (item["key"], action_key)
            for group in ROLE_ACTION_PERMISSION_GROUPS
            for item in group["items"]
            for action_key in item["actions"]
        ]
        timestamp = now()
        db().execute("delete from role_menu_permissions")
        db().execute("delete from role_action_permissions")
        for role in ROLE_OPTIONS:
            for menu_key in menu_keys:
                enabled = 1 if request.form.get(f"menu__{role}__{menu_key}") == "1" else 0
                db().execute(
                    """
                    insert into role_menu_permissions (
                        role, menu_key, is_enabled, updated_by, updated_at
                    ) values (?, ?, ?, ?, ?)
                    """,
                    (role, menu_key, enabled, g.user["id"], timestamp),
                )
            for resource_key, action_key in action_pairs:
                enabled = 1 if request.form.get(f"action__{role}__{resource_key}__{action_key}") == "1" else 0
                db().execute(
                    """
                    insert into role_action_permissions (
                        role, resource_key, action_key, is_enabled, updated_by, updated_at
                    ) values (?, ?, ?, ?, ?, ?)
                    """,
                    (role, resource_key, action_key, enabled, g.user["id"], timestamp),
                )
        log_action("update", "menu_permissions", None, "系统权限", "修改菜单和操作权限")
        db().commit()
        g._menu_permission_overrides = {
            (role, menu_key): request.form.get(f"menu__{role}__{menu_key}") == "1"
            for role in ROLE_OPTIONS
            for menu_key in menu_keys
        }
        g._action_permission_overrides = {
            (role, resource_key, action_key): request.form.get(f"action__{role}__{resource_key}__{action_key}") == "1"
            for role in ROLE_OPTIONS
            for resource_key, action_key in action_pairs
        }
        flash("权限配置已保存。", "success")
        return redirect(url_for("menu_permissions"))
    return render_template(
        "menu_permissions.html",
        role_options=ROLE_OPTIONS,
        menu_groups=MENU_PERMISSION_GROUPS,
        action_groups=ROLE_ACTION_PERMISSION_GROUPS,
        action_labels=ACTION_LABELS,
        permission_tree_groups=permission_tree_groups(),
        users_by_role={
            role: db().execute(
                "select id, name, email, is_active from users where role = ? order by is_active desc, name",
                (role,),
            ).fetchall()
            for role in ROLE_OPTIONS
        },
    )


@app.route("/settings/database", methods=["GET", "POST"])
@login_required
def database_console():
    sql = request.form.get("sql", "").strip() if request.method == "POST" else ""
    result = None
    columns = []
    rows = []
    row_count = None
    statement_kind = sql_statement_kind(sql)
    if request.method == "POST":
        if not sql:
            flash("请输入 SQL 语句。", "error")
        elif ";" in sql.rstrip(";"):
            flash("为了降低误操作风险，每次只能执行一条 SQL 语句。", "error")
        else:
            read_query = is_read_sql(sql)
            confirmed = request.form.get("confirm_write") == "1"
            if not read_query and not confirmed:
                flash("执行数据更改语句前，请先勾选确认。", "error")
            else:
                try:
                    cursor = db().execute(sql)
                    if cursor.description:
                        columns = [description[0] for description in cursor.description]
                        rows = cursor.fetchmany(500)
                        result = "query"
                        if cursor.fetchone() is not None:
                            flash("查询结果超过 500 行，仅显示前 500 行。", "warning")
                        if not read_query:
                            log_action(
                                "sql",
                                "database",
                                None,
                                "SQL Console",
                                sql[:500],
                            )
                            db().commit()
                            flash("SQL 已执行。", "success")
                    else:
                        row_count = cursor.rowcount
                        result = "write"
                        log_action(
                            "sql",
                            "database",
                            None,
                            "SQL Console",
                            sql[:500],
                        )
                        db().commit()
                        flash("SQL 已执行。", "success")
                except sqlite3.Error as error:
                    db().rollback()
                    flash(f"SQL 执行失败：{error}", "error")
    tables = db().execute(
        """
        select name, type
        from sqlite_master
        where type in ('table', 'view') and name not like 'sqlite_%'
        order by type, name
        """
    ).fetchall()
    table_dictionary = []
    for table in tables:
        quoted_name = '"' + table["name"].replace('"', '""') + '"'
        table_dictionary.append(
            {
                "name": table["name"],
                "type": table["type"],
                "columns": db().execute(f"pragma table_info({quoted_name})").fetchall(),
            }
        )
    return render_template(
        "database_console.html",
        sql=sql,
        statement_kind=statement_kind,
        result=result,
        columns=columns,
        rows=rows,
        row_count=row_count,
        tables=tables,
        table_dictionary=table_dictionary,
    )


@app.route("/company-info", methods=["GET", "POST"])
@login_required
def company_info():
    if not is_internal_user():
        abort(403)
    if request.method == "POST" and not can_manage_company_info():
        abort(403)
    if request.method == "POST":
        for key in ("name", "address", "email", "phone", "registration_number", "ein"):
            set_setting(f"company_{key}", request.form.get(f"company_{key}", "").strip())
        timezone_name = request.form.get("app_timezone", DEFAULT_TIMEZONE).strip() or DEFAULT_TIMEZONE
        try:
            ZoneInfo(timezone_name)
        except ZoneInfoNotFoundError:
            flash("系统时区无效，请使用类似 America/Chicago 的时区名称。", "error")
            return redirect(url_for("company_info"))
        set_setting("app_timezone", timezone_name)
        try:
            for uploaded in uploaded_attachments_from_request():
                save_named_attachment(uploaded, COMPANY_ATTACHMENT_DIR, "company_attachments")
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("company_info"))
        db().commit()
        flash("公司信息已保存。", "success")
        return redirect(url_for("company_info"))
    return render_template(
        "company_info.html",
        company=get_company_profile(),
        timezone_name=get_timezone_name(),
        attachments=get_company_attachments(),
    )


@app.get("/company-attachments/<int:attachment_id>/download")
@login_required
def download_company_attachment(attachment_id):
    if not is_internal_user():
        abort(403)
    attachment = db().execute("select * from company_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    return send_file(
        os.path.join(COMPANY_ATTACHMENT_DIR, attachment["stored_filename"]),
        as_attachment=True,
        download_name=attachment["original_filename"],
    )


@app.get("/company-attachments/<int:attachment_id>/preview")
@login_required
def preview_company_attachment(attachment_id):
    if not is_internal_user():
        abort(403)
    attachment = db().execute("select * from company_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    return send_file(
        os.path.join(COMPANY_ATTACHMENT_DIR, attachment["stored_filename"]),
        as_attachment=False,
        download_name=attachment["original_filename"],
        mimetype=attachment["content_type"] or None,
        conditional=True,
    )


@app.post("/company-attachments/<int:attachment_id>/delete")
@login_required
def delete_company_attachment(attachment_id):
    attachment = db().execute("select * from company_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    try:
        os.remove(os.path.join(COMPANY_ATTACHMENT_DIR, attachment["stored_filename"]))
    except FileNotFoundError:
        pass
    db().execute("delete from company_attachments where id = ?", (attachment_id,))
    db().commit()
    flash("公司附件已删除。", "success")
    return redirect(url_for("company_info"))


@app.get("/settings/company")
@login_required
def legacy_company_settings():
    return redirect(url_for("system_settings"))


@app.route("/invoices")
@login_required
def invoices():
    if not can_view_invoices():
        abort(403)
    status = request.args.get("status", "")
    paid_status = request.args.get("paid_status", "")
    q = request.args.get("q", "").strip()
    date_from = request.args.get("date_from", "")
    date_to = request.args.get("date_to", "")
    created_by = request.args.get("created_by", "")
    access_clause, access_params = client_filter_clause("invoices")
    clauses = [access_clause]
    params = list(access_params)
    if status:
        clauses.append("invoices.status = ?")
        params.append(status)
    if paid_status == "paid":
        clauses.append("invoices.paid_at is not null")
    elif paid_status == "unpaid":
        clauses.append("invoices.status = 'completed' and invoices.paid_at is null")
    if q:
        clauses.append("(invoices.invoice_number like ? or clients.name like ? or clients.short_name like ? or clients.client_number like ?)")
        params.extend([f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"])
    if date_from:
        clauses.append("invoices.issue_date >= ?")
        params.append(date_from)
    if date_to:
        clauses.append("invoices.issue_date <= ?")
        params.append(date_to)
    if created_by and not is_external_user():
        clauses.append("invoices.created_by = ?")
        params.append(created_by)
    rows = db().execute(
        f"""
        select invoices.*, clients.name as client_name, clients.short_name as client_short_name,
               clients.client_number, users.name as creator_name,
               service_orders.order_number as service_order_number
        from invoices
        join clients on clients.id = invoices.client_id
        left join users on users.id = invoices.created_by
        left join service_orders on service_orders.id = invoices.service_order_id
        where {" and ".join(clauses)}
        order by invoices.issue_date desc, invoices.id desc
        """,
        params,
    ).fetchall()
    totals = {row["id"]: invoice_totals(row["id"]) for row in rows}
    users_rows = db().execute("select id, name, email from users order by name").fetchall() if not is_external_user() else []
    return render_template(
        "invoices.html",
        invoices=rows,
        totals=totals,
        labels=STATUS_LABELS,
        users=users_rows,
        status=status,
        paid_status=paid_status,
        q=q,
        date_from=date_from,
        date_to=date_to,
        created_by=created_by,
    )


@app.route("/reports/invoices")
@login_required
def invoice_query():
    if not can_view_invoices():
        abort(403)
    client_q = request.args.get("client", "").strip()
    short_q = request.args.get("short_name", "").strip()
    date_from = request.args.get("date_from", "")
    date_to = request.args.get("date_to", "")
    region_code, country_code, countries, regions, location_clauses, location_params = report_location_filters()
    available_statuses = set(STATUS_LABELS.keys())
    selected_statuses = [value for value in request.args.getlist("status") if value in available_statuses]
    selected_paid_statuses = [value for value in request.args.getlist("paid_status") if value in {"paid", "unpaid"}]
    project_options = report_project_options()
    available_project_ids = {str(project["id"]) for project in project_options}
    selected_project_ids = [value for value in request.args.getlist("project_id") if value in available_project_ids]
    effective_project_ids = selected_project_ids or [str(project["id"]) for project in project_options]
    access_clause, access_params = client_filter_clause("invoices")
    clauses = [access_clause]
    params = list(access_params)
    clauses.extend(location_clauses)
    params.extend(location_params)
    if selected_statuses:
        placeholders = ",".join("?" for _ in selected_statuses)
        clauses.append(f"invoices.status in ({placeholders})")
        params.extend(selected_statuses)
    else:
        clauses.append("invoices.status != 'void'")
    if selected_paid_statuses and len(selected_paid_statuses) < 2:
        if selected_paid_statuses[0] == "paid":
            clauses.append("invoices.paid_at is not null")
        elif selected_paid_statuses[0] == "unpaid":
            clauses.append("invoices.status = 'completed' and invoices.paid_at is null")
    if client_q:
        clauses.append("(clients.name like ? or clients.client_number like ?)")
        params.extend([f"%{client_q}%", f"%{client_q}%"])
    if short_q:
        clauses.append("clients.short_name like ?")
        params.append(f"%{short_q}%")
    if date_from:
        clauses.append("invoices.issue_date >= ?")
        params.append(date_from)
    if date_to:
        clauses.append("invoices.issue_date <= ?")
        params.append(date_to)
    if effective_project_ids:
        placeholders = ",".join("?" for _ in effective_project_ids)
        clauses.append(f"invoice_items.project_id in ({placeholders})")
        params.extend(effective_project_ids)
    rows = db().execute(
        f"""
        select invoices.id, invoices.invoice_number, invoices.issue_date, invoices.currency, invoices.status, invoices.paid_at,
               clients.client_number, clients.name as client_name, clients.short_name,
               invoice_items.description as project_name, invoice_items.amount, invoice_items.tax_rate
        from invoice_items
        join invoices on invoices.id = invoice_items.invoice_id
        join clients on clients.id = invoices.client_id
        left join service_orders on service_orders.id = invoices.service_order_id
        where {" and ".join(clauses)}
        order by invoices.issue_date desc, invoices.id desc, invoice_items.id asc
        """,
        params,
    ).fetchall()
    report_rows = []
    subtotal = tax_total = grand_total = 0
    for row in rows:
        amount = float(row["amount"] or 0)
        tax = amount * float(row["tax_rate"] or 0) / 100
        line_total = amount + tax
        subtotal += amount
        tax_total += tax
        grand_total += line_total
        report_rows.append({"row": row, "tax": tax, "line_total": line_total})
    return render_template(
        "reports.html",
        rows=report_rows,
        project_options=project_options,
        selected_project_ids=effective_project_ids,
        client_q=client_q,
        short_q=short_q,
        date_from=date_from,
        date_to=date_to,
        selected_statuses=selected_statuses or [value for value in STATUS_LABELS if value != "void"],
        selected_paid_statuses=selected_paid_statuses or ["paid", "unpaid"],
        subtotal=subtotal,
        tax_total=tax_total,
        grand_total=grand_total,
        labels=STATUS_LABELS,
        countries=countries,
        regions=regions,
        region_code=region_code,
        country_code=country_code,
    )


@app.route("/reports")
@login_required
def report_center():
    return redirect(url_for("service_order_query"))


@app.route("/reports/service-orders")
@login_required
def service_order_query():
    if not is_internal_user():
        abort(403)
    q = request.args.get("q", "").strip()
    status = request.args.get("status", "")
    region_code, country_code, countries, regions, location_clauses, location_params = report_location_filters()
    clauses = ["1 = 1"]
    params = []
    clauses.extend(location_clauses)
    params.extend(location_params)
    if q:
        clauses.append(
            "(service_orders.order_number like ? or service_orders.client_name like ? or coalesce(owners.name, buyers.owner) like ? or service_orders.client_order_number like ?)"
        )
        params.extend([f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"])
    if status:
        clauses.append("service_orders.status = ?")
        params.append(status)
    rows = db().execute(
        f"""
        select service_orders.*,
               work_order_types.name as work_order_type_name,
               coalesce(owners.name, buyers.owner) as buyer_owner,
               buyers.equipment_manufacturer as buyer_equipment_manufacturer,
               coalesce(country_local.name, country_zh.name, service_orders.country_code) as country_name,
               coalesce(country_local.region_name, country_zh.region_name, service_orders.region_code) as region_name,
               count(distinct service_reports.id) as report_count,
               count(distinct invoices.id) as invoice_count,
               coalesce(sum(case when expenses.status = 'approved' then expenses.amount else 0 end), 0) as approved_expense_total
        from service_orders
        left join work_order_types on work_order_types.id = service_orders.work_order_type_id
        left join buyers on buyers.id = service_orders.buyer_id
        left join owners on owners.id = buyers.owner_id
        left join country_translations country_local
          on country_local.country_code = service_orders.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = service_orders.country_code and country_zh.language_code = 'zh-CN'
        left join service_reports on service_reports.service_order_id = service_orders.id
        left join invoices on invoices.service_order_id = service_orders.id and invoices.status != 'void'
        left join expenses on expenses.service_order_id = service_orders.id
        where {" and ".join(clauses)}
        group by service_orders.id
        order by service_orders.created_at desc, service_orders.id desc
        """,
        [current_language(), *params],
    ).fetchall()
    return render_template(
        "service_order_query.html",
        orders=rows,
        q=q,
        status=status,
        countries=countries,
        regions=regions,
        region_code=region_code,
        country_code=country_code,
    )


@app.route("/reports/buyers")
@login_required
def buyer_query():
    if not is_internal_user():
        abort(403)
    q = request.args.get("q", "").strip()
    region_code, country_code, countries, regions, _, _ = report_location_filters()
    clauses = ["1 = 1"]
    params = []
    if q:
        clauses.append(
            """
            (buyers.buyer_number like ? or buyers.name like ? or coalesce(owners.name, buyers.owner) like ? or buyers.contact_name like ?
             or buyers.contact_details like ? or buyers.email like ? or buyers.detailed_address like ?
             or buyers.equipment_manufacturer like ? or buyers.site_size like ?)
            """
        )
        params.extend([f"%{q}%"] * 9)
    if region_code:
        clauses.append("exists (select 1 from countries where countries.code = buyers.country_code and countries.region_code = ?)")
        params.append(region_code)
    if country_code:
        clauses.append("buyers.country_code = ?")
        params.append(country_code)
    rows = buyer_map_rows(" and ".join(clauses), params)
    return render_template(
        "buyer_query.html",
        buyers=rows,
        countries=countries,
        regions=regions,
        q=q,
        region_code=region_code,
        country_code=country_code,
    )


@app.route("/reports/service-reports")
@login_required
def service_report_query():
    if not is_internal_user():
        abort(403)
    q = request.args.get("q", "").strip()
    date_from = request.args.get("date_from", "")
    date_to = request.args.get("date_to", "")
    region_code, country_code, countries, regions, location_clauses, location_params = report_location_filters()
    clauses = ["1 = 1"]
    params = []
    clauses.extend(location_clauses)
    params.extend(location_params)
    if q:
        clauses.append(
            "(service_orders.order_number like ? or service_orders.client_name like ? or coalesce(owners.name, buyers.owner) like ? or service_reports.cabinet_number like ?)"
        )
        params.extend([f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"])
    if date_from:
        clauses.append("coalesce(service_reports.actual_work_date, service_reports.report_date) >= ?")
        params.append(date_from)
    if date_to:
        clauses.append("coalesce(service_reports.actual_work_date, service_reports.report_date) <= ?")
        params.append(date_to)
    rows = db().execute(
        f"""
        select service_reports.*, service_orders.order_number, service_orders.client_name,
               coalesce(owners.name, buyers.owner) as buyer_owner,
               coalesce(country_local.name, country_zh.name, service_orders.country_code) as country_name,
               coalesce(country_local.region_name, country_zh.region_name, service_orders.region_code) as region_name,
               group_concat(users.name, ', ') as worker_names
        from service_reports
        join service_orders on service_orders.id = service_reports.service_order_id
        left join buyers on buyers.id = service_orders.buyer_id
        left join owners on owners.id = buyers.owner_id
        left join country_translations country_local
          on country_local.country_code = service_orders.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = service_orders.country_code and country_zh.language_code = 'zh-CN'
        left join service_report_workers on service_report_workers.report_id = service_reports.id
        left join users on users.id = service_report_workers.user_id
        where {" and ".join(clauses)}
        group by service_reports.id
        order by coalesce(service_reports.actual_work_date, service_reports.report_date) desc, service_reports.id desc
        """,
        [current_language(), *params],
    ).fetchall()
    return render_template(
        "service_report_query.html",
        rows=rows,
        q=q,
        date_from=date_from,
        date_to=date_to,
        countries=countries,
        regions=regions,
        region_code=region_code,
        country_code=country_code,
    )


def labor_report_entries(date_from="", date_to="", worker_id="", date_mode="actual"):
    if date_mode == "attendance":
        report_date_expr = "date(service_reports.updated_at)"
    elif date_mode == "report":
        report_date_expr = "date(service_reports.report_date)"
    else:
        report_date_expr = "coalesce(service_reports.actual_work_date, service_reports.report_date)"
    clauses = ["1 = 1"]
    params = []
    if date_from:
        clauses.append(f"{report_date_expr} >= ?")
        params.append(date_from)
    if date_to:
        clauses.append(f"{report_date_expr} <= ?")
        params.append(date_to)
    if worker_id and str(worker_id).isdigit():
        clauses.append("users.id = ?")
        params.append(int(worker_id))
    rows = db().execute(
        f"""
        with worker_counts as (
            select report_id, count(*) as worker_count
            from service_report_workers
            group by report_id
        )
        select service_reports.id as report_id,
               service_reports.report_date,
               {report_date_expr} as actual_work_date,
               {report_date_expr} as attendance_date,
               service_reports.total_service_hours,
               service_reports.travel_hours,
               service_reports.public_transport_hours,
               service_report_workers.driving_miles as worker_driving_miles,
               service_reports.arrival_time,
               service_reports.departure_time,
               service_orders.id as service_order_id,
               service_orders.order_number,
               service_orders.client_name,
               coalesce(owners.name, buyers.owner) as buyer_owner,
               users.id as worker_id,
               users.name as worker_name,
               employee_grades.grade_name,
               employee_grades.base_salary,
               employee_grades.standard_hourly_rate,
               employee_grades.transport_hourly_rate,
               employee_grades.overtime_hourly_rate,
               employee_grades.holiday_hourly_rate,
               coalesce(worker_counts.worker_count, 1) as worker_count
        from service_reports
        join service_orders on service_orders.id = service_reports.service_order_id
        join service_report_workers on service_report_workers.report_id = service_reports.id
        join users on users.id = service_report_workers.user_id
        left join employee_grades on employee_grades.id = users.employee_grade_id
        left join buyers on buyers.id = service_orders.buyer_id
        left join owners on owners.id = buyers.owner_id
        left join worker_counts on worker_counts.report_id = service_reports.id
        where {" and ".join(clauses)}
        order by actual_work_date desc, users.name, service_orders.order_number
        """,
        params,
    ).fetchall()
    entries = []
    for row in rows:
        hours = split_report_labor_hours(row, row["worker_count"])
        entry = dict(row)
        entry.update(hours)
        entry["work_hours"] = hours["standard_hours"] + hours["overtime_hours"] + hours["holiday_hours"]
        entries.append(entry)
    return entries


@app.route("/reports/labor-hours")
@login_required
def labor_hours_report():
    if not can_view_labor_payroll_reports():
        abort(403)
    date_from = request.args.get("date_from", "")
    date_to = request.args.get("date_to", "")
    worker_id = request.args.get("worker_id", "")
    can_filter_workers = normalized_role() in {"admin", "manager", "finance"}
    effective_worker_id = worker_id if can_filter_workers else str(g.user["id"])
    workers = db().execute(
        """
        select id, name
        from users
        where role in ('manager', 'finance', 'employee', 'external_employee')
        order by name
        """
    ).fetchall() if can_filter_workers else []
    rows = labor_report_entries(date_from, date_to, effective_worker_id)
    totals = {
        "work_hours": sum(row["work_hours"] for row in rows),
        "standard_hours": sum(row["standard_hours"] for row in rows),
        "transport_hours": sum(row["transport_hours"] for row in rows),
        "overtime_hours": sum(row["overtime_hours"] for row in rows),
        "holiday_hours": sum(row["holiday_hours"] for row in rows),
    }
    return render_template(
        "labor_hours_report.html",
        rows=rows,
        totals=totals,
        workers=workers,
        worker_id=int(effective_worker_id) if str(effective_worker_id).isdigit() else "",
        can_filter_workers=can_filter_workers,
        date_from=date_from,
        date_to=date_to,
    )


def current_payroll_period_start(today=None):
    today = today or date.today()
    try:
        cycle_start = date.fromisoformat(get_setting("payroll_cycle_start", "2026-07-06"))
    except ValueError:
        cycle_start = date(2026, 7, 6)
    if today < cycle_start:
        return cycle_start
    elapsed_periods = (today - cycle_start).days // 14
    return cycle_start + timedelta(days=elapsed_periods * 14)


def payroll_period_dates(period_start):
    period_end = period_start + timedelta(days=13)
    pay_date = period_end + timedelta(days=14)
    return period_end, pay_date


def effective_payroll_worker_id():
    return "" if normalized_role() in {"admin", "manager", "finance"} else str(g.user["id"])


def payroll_rows_for_range(period_start, period_end, pay_date, worker_id="", date_mode="attendance"):
    subsidy_settings = payroll_subsidy_settings()
    show_meal_allowance = subsidy_settings["meal_allowance_method"] == "daily"
    rows = labor_report_entries(period_start.isoformat(), period_end.isoformat(), worker_id, date_mode=date_mode)
    payroll = {}
    for row in rows:
        worker = payroll.setdefault(
            row["worker_id"],
            {
                "worker_id": row["worker_id"],
                "worker_name": row["worker_name"],
                "grade_name": row["grade_name"] or "未指定",
                "base_salary": float(row["base_salary"] or 0),
                "standard_rate": float(row["standard_hourly_rate"] or 0),
                "transport_rate": float(row["transport_hourly_rate"] or 0),
                "overtime_rate": float(row["overtime_hourly_rate"] or 0),
                "holiday_rate": float(row["holiday_hourly_rate"] or 0),
                "standard_hours": 0,
                "transport_hours": 0,
                "overtime_hours": 0,
                "holiday_hours": 0,
                "attendance_dates": set(),
                "driving_miles": 0,
            },
        )
        for key in ("standard_hours", "transport_hours", "overtime_hours", "holiday_hours"):
            worker[key] += row[key]
        if row["attendance_date"]:
            worker["attendance_dates"].add(row["attendance_date"])
        worker["driving_miles"] += float(row["worker_driving_miles"] or 0)
    payroll_rows = []
    for worker in payroll.values():
        attendance_days = len(worker["attendance_dates"])
        standard_pay = worker["standard_hours"] * worker["standard_rate"]
        transport_pay = worker["transport_hours"] * worker["transport_rate"]
        overtime_pay = worker["overtime_hours"] * worker["overtime_rate"]
        holiday_pay = worker["holiday_hours"] * worker["holiday_rate"]
        if subsidy_settings["car_allowance_method"] == "mileage":
            car_allowance = worker["driving_miles"] * subsidy_settings["car_mileage_rate"]
        else:
            car_allowance = attendance_days * subsidy_settings["car_daily_amount"]
        meal_allowance = attendance_days * subsidy_settings["meal_daily_amount"] if show_meal_allowance else 0
        total_pay = (
            worker["base_salary"]
            + standard_pay
            + transport_pay
            + overtime_pay
            + holiday_pay
            + car_allowance
            + meal_allowance
        )
        payroll_rows.append(
            {
                **worker,
                "attendance_days": attendance_days,
                "standard_pay": standard_pay,
                "transport_pay": transport_pay,
                "overtime_pay": overtime_pay,
                "holiday_pay": holiday_pay,
                "car_allowance": car_allowance,
                "meal_allowance": meal_allowance,
                "total_pay": total_pay,
            }
        )
    payroll_rows.sort(key=lambda row: row["worker_name"])
    totals = {
        "base_salary": sum(row["base_salary"] for row in payroll_rows),
        "standard_pay": sum(row["standard_pay"] for row in payroll_rows),
        "transport_pay": sum(row["transport_pay"] for row in payroll_rows),
        "overtime_pay": sum(row["overtime_pay"] for row in payroll_rows),
        "holiday_pay": sum(row["holiday_pay"] for row in payroll_rows),
        "car_allowance": sum(row["car_allowance"] for row in payroll_rows),
        "meal_allowance": sum(row["meal_allowance"] for row in payroll_rows),
        "total_pay": sum(row["total_pay"] for row in payroll_rows),
    }
    return {
        "rows": payroll_rows,
        "totals": totals,
        "period_start": period_start,
        "period_end": period_end,
        "pay_date": pay_date,
        "subsidy_settings": subsidy_settings,
        "show_meal_allowance": show_meal_allowance,
    }


def payroll_rows_for_period(period_start, worker_id=""):
    period_end, pay_date = payroll_period_dates(period_start)
    return payroll_rows_for_range(period_start, period_end, pay_date, worker_id)


def payroll_row_export(row, show_meal_allowance):
    payload = {
        "员工": row["worker_name"],
        "员工等级": row["grade_name"],
        "基本工资": round(row["base_salary"], 2),
        "出勤天数": row["attendance_days"],
        "里程": round(row["driving_miles"], 1),
        "标准工时": round(row["standard_hours"], 2),
        "标准工资": round(row["standard_pay"], 2),
        "交通工时": round(row["transport_hours"], 2),
        "交通工资": round(row["transport_pay"], 2),
        "加班工时": round(row["overtime_hours"], 2),
        "加班工资": round(row["overtime_pay"], 2),
        "假期工时": round(row["holiday_hours"], 2),
        "假期工资": round(row["holiday_pay"], 2),
        "车补": round(row["car_allowance"], 2),
    }
    if show_meal_allowance:
        payload["餐补"] = round(row["meal_allowance"], 2)
    payload["合计工资"] = round(row["total_pay"], 2)
    return payload


def payroll_payslip_payload(row, show_meal_allowance):
    lines = [
        {"label": "基本工资", "amount": round(row["base_salary"], 2)},
        {"label": "标准工资", "hours": round(row["standard_hours"], 2), "rate": round(row["standard_rate"], 2), "amount": round(row["standard_pay"], 2)},
        {"label": "交通工资", "hours": round(row["transport_hours"], 2), "rate": round(row["transport_rate"], 2), "amount": round(row["transport_pay"], 2)},
        {"label": "加班工资", "hours": round(row["overtime_hours"], 2), "rate": round(row["overtime_rate"], 2), "amount": round(row["overtime_pay"], 2)},
        {"label": "假期工资", "hours": round(row["holiday_hours"], 2), "rate": round(row["holiday_rate"], 2), "amount": round(row["holiday_pay"], 2)},
        {"label": "车补", "days": row["attendance_days"], "miles": round(row["driving_miles"], 1), "amount": round(row["car_allowance"], 2)},
    ]
    if show_meal_allowance:
        lines.append({"label": "餐补", "days": row["attendance_days"], "amount": round(row["meal_allowance"], 2)})
    return {
        "employee": row["worker_name"],
        "grade": row["grade_name"],
        "attendance_days": row["attendance_days"],
        "driving_miles": round(row["driving_miles"], 1),
        "lines": lines,
        "total_pay": round(row["total_pay"], 2),
    }


def payroll_historical_paid_date():
    try:
        return date.fromisoformat(get_setting("payroll_historical_paid_date", date.today().isoformat()))
    except ValueError:
        return date.today()


def historical_payroll_period(worker_id=""):
    cycle_start = payroll_cycle_start_date()
    period_end = min(cycle_start - timedelta(days=1), payroll_historical_paid_date())
    clauses = ["date(service_reports.report_date) <= ?"]
    params = [period_end.isoformat()]
    if worker_id and str(worker_id).isdigit():
        clauses.append("users.id = ?")
        params.append(int(worker_id))
    row = db().execute(
        f"""
        select min(date(service_reports.report_date)) as period_start
        from service_reports
        join service_report_workers on service_report_workers.report_id = service_reports.id
        join users on users.id = service_report_workers.user_id
        where {" and ".join(clauses)}
        """,
        params,
    ).fetchone()
    if not row or not row["period_start"]:
        return None
    period_start = date.fromisoformat(row["period_start"])
    if period_start > period_end:
        return None
    return {
        "batch_type": "historical",
        "label": "历史工资补发",
        "period_start": period_start,
        "period_end": period_end,
        "pay_date": payroll_historical_paid_date(),
        "status": "paid",
    }


def payroll_batch_payload(period_start, batch_type="regular"):
    if batch_type == "historical":
        period = historical_payroll_period(effective_payroll_worker_id())
        if not period or period["period_start"] != period_start:
            abort(404)
        batch = payroll_rows_for_range(
            period["period_start"],
            period["period_end"],
            period["pay_date"],
            effective_payroll_worker_id(),
            date_mode="report",
        )
        label = period["label"]
        status = "paid"
    else:
        batch = payroll_rows_for_period(period_start, effective_payroll_worker_id())
        label = "工资发放"
        status = "paid" if batch["pay_date"] <= date.today() else "scheduled"
    return {
        "batch_type": batch_type,
        "label": label,
        "period_start": batch["period_start"].isoformat(),
        "period_end": batch["period_end"].isoformat(),
        "pay_date": batch["pay_date"].isoformat(),
        "status": status,
        "show_meal_allowance": batch["show_meal_allowance"],
        "rows": [payroll_row_export(row, batch["show_meal_allowance"]) for row in batch["rows"]],
        "payslips": [payroll_payslip_payload(row, batch["show_meal_allowance"]) for row in batch["rows"]],
        "totals": {key: round(value, 2) for key, value in batch["totals"].items()},
    }


def payroll_cycle_start_date():
    try:
        return date.fromisoformat(get_setting("payroll_cycle_start", "2026-07-06"))
    except ValueError:
        return date(2026, 7, 6)


def payroll_periods_for_month(year, month):
    month_start = date(year, month, 1)
    month_end = (date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)) - timedelta(days=1)
    cycle_start = payroll_cycle_start_date()
    start = cycle_start
    if month_start > cycle_start + timedelta(days=27):
        periods_before = max(((month_start - cycle_start).days - 27) // 14 - 1, 0)
        start = cycle_start + timedelta(days=periods_before * 14)
    periods = []
    while start <= month_end:
        period_end, pay_date = payroll_period_dates(start)
        if month_start <= pay_date <= month_end:
            periods.append(
                {
                    "batch_type": "regular",
                    "label": "工资发放",
                    "period_start": start,
                    "period_end": period_end,
                    "pay_date": pay_date,
                    "status": "paid" if pay_date <= date.today() else "scheduled",
                }
            )
        start += timedelta(days=14)
    historical_period = historical_payroll_period(effective_payroll_worker_id())
    if historical_period and month_start <= historical_period["pay_date"] <= month_end:
        periods.append(historical_period)
    periods.sort(key=lambda period: (period["pay_date"], period["period_start"], period["batch_type"]))
    return periods


def payroll_calendar_weeks(year, month, periods):
    month_start = date(year, month, 1)
    month_end = (date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)) - timedelta(days=1)
    pay_days = {}
    for period in periods:
        pay_days.setdefault(period["pay_date"], []).append(period)
    current = month_start - timedelta(days=month_start.weekday())
    weeks = []
    while current <= month_end or current.weekday() != 0:
        week = []
        for _ in range(7):
            day_periods = pay_days.get(current, [])
            week.append(
                {
                    "date": current,
                    "in_month": current.month == month,
                    "periods": day_periods,
                    "status": "paid" if any(period["status"] == "paid" for period in day_periods) else ("scheduled" if day_periods else ""),
                    "is_today": current == date.today(),
                }
            )
            current += timedelta(days=1)
        weeks.append(week)
    return weeks


def xlsx_column_name(index):
    name = ""
    index += 1
    while index:
        index, remainder = divmod(index - 1, 26)
        name = chr(65 + remainder) + name
    return name


def build_simple_xlsx(headers, rows, sheet_name="Sheet1"):
    def cell_xml(row_index, column_index, value):
        ref = f"{xlsx_column_name(column_index)}{row_index}"
        if isinstance(value, (int, float, Decimal)) and not isinstance(value, bool):
            return f'<c r="{ref}"><v>{float(value):.2f}</v></c>'
        text = html.escape("" if value is None else str(value))
        return f'<c r="{ref}" t="inlineStr"><is><t>{text}</t></is></c>'

    worksheet_rows = []
    all_rows = [headers] + rows
    for row_index, row in enumerate(all_rows, start=1):
        cells = "".join(cell_xml(row_index, column_index, value) for column_index, value in enumerate(row))
        worksheet_rows.append(f'<row r="{row_index}">{cells}</row>')
    dimension = f"A1:{xlsx_column_name(max(len(headers) - 1, 0))}{max(len(all_rows), 1)}"
    sheet_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        f'<dimension ref="{dimension}"/><sheetData>{"".join(worksheet_rows)}</sheetData></worksheet>'
    )
    workbook_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<sheets><sheet name="{html.escape(sheet_name)}" sheetId="1" r:id="rId1"/></sheets></workbook>'
    )
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        '</Types>'
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
        '</Relationships>'
    )
    workbook_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>'
        '</Relationships>'
    )
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as workbook:
        workbook.writestr("[Content_Types].xml", content_types)
        workbook.writestr("_rels/.rels", root_rels)
        workbook.writestr("xl/workbook.xml", workbook_xml)
        workbook.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        workbook.writestr("xl/worksheets/sheet1.xml", sheet_xml)
    buffer.seek(0)
    return buffer


@app.route("/reports/payroll")
@login_required
def payroll_report():
    if not can_view_labor_payroll_reports():
        abort(403)
    requested_start = request.args.get("period_start", "")
    try:
        period_start = date.fromisoformat(requested_start) if requested_start else current_payroll_period_start()
    except ValueError:
        period_start = current_payroll_period_start()
    payroll_batch = payroll_rows_for_period(period_start, effective_payroll_worker_id())
    return render_template(
        "payroll_report.html",
        rows=payroll_batch["rows"],
        totals=payroll_batch["totals"],
        period_start=payroll_batch["period_start"].isoformat(),
        period_end=payroll_batch["period_end"].isoformat(),
        pay_date=payroll_batch["pay_date"].isoformat(),
        subsidy_settings=payroll_batch["subsidy_settings"],
        show_meal_allowance=payroll_batch["show_meal_allowance"],
    )


@app.route("/payroll/calendar")
@login_required
def payroll_calendar():
    if not can_view_labor_payroll_reports():
        abort(403)
    today = date.today()
    try:
        year = int(request.args.get("year", today.year))
        month = int(request.args.get("month", today.month))
        visible_month = date(year, month, 1)
    except ValueError:
        visible_month = date(today.year, today.month, 1)
    periods = payroll_periods_for_month(visible_month.year, visible_month.month)
    previous_month = date(visible_month.year - 1, 12, 1) if visible_month.month == 1 else date(visible_month.year, visible_month.month - 1, 1)
    next_month = date(visible_month.year + 1, 1, 1) if visible_month.month == 12 else date(visible_month.year, visible_month.month + 1, 1)
    return render_template(
        "payroll_calendar.html",
        visible_month=visible_month,
        previous_month=previous_month,
        next_month=next_month,
        weeks=payroll_calendar_weeks(visible_month.year, visible_month.month, periods),
    )


@app.get("/payroll/calendar/batch")
@login_required
def payroll_calendar_batch():
    if not can_view_labor_payroll_reports():
        abort(403)
    try:
        period_start = date.fromisoformat(request.args.get("period_start", ""))
    except ValueError:
        abort(400)
    return jsonify(payroll_batch_payload(period_start, request.args.get("batch_type", "regular")))


@app.get("/payroll/calendar/export.xlsx")
@login_required
def payroll_calendar_export():
    if not can_view_labor_payroll_reports():
        abort(403)
    try:
        period_start = date.fromisoformat(request.args.get("period_start", ""))
    except ValueError:
        abort(400)
    payload = payroll_batch_payload(period_start, request.args.get("batch_type", "regular"))
    headers = list(payload["rows"][0].keys()) if payload["rows"] else [
        "员工", "员工等级", "基本工资", "出勤天数", "里程", "标准工时", "标准工资",
        "交通工时", "交通工资", "加班工时", "加班工资", "假期工时", "假期工资",
        "车补", "餐补", "合计工资",
    ]
    if not payload["show_meal_allowance"] and "餐补" in headers:
        headers.remove("餐补")
    rows = [[row.get(header, "") for header in headers] for row in payload["rows"]]
    rows.append([])
    rows.append(["周期", f"{payload['period_start']} 至 {payload['period_end']}"])
    rows.append(["发薪日期", payload["pay_date"]])
    rows.append(["合计工资", payload["totals"]["total_pay"]])
    rows.append([])
    rows.append(["工资条"])
    for payslip in payload["payslips"]:
        rows.append([])
        rows.append(["员工", payslip["employee"], "员工等级", payslip["grade"], "合计工资", payslip["total_pay"]])
        rows.append(["项目", "工时", "单价", "天数", "里程", "金额"])
        for line in payslip["lines"]:
            rows.append([
                line.get("label", ""),
                line.get("hours", ""),
                line.get("rate", ""),
                line.get("days", ""),
                line.get("miles", ""),
                line.get("amount", ""),
            ])
    workbook = build_simple_xlsx(headers, rows, sheet_name="工资批次")
    filename = f"payroll-{payload['batch_type']}-{payload['pay_date']}.xlsx"
    return send_file(
        workbook,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/reports/customer-reimbursements")
@login_required
def customer_reimbursement_query():
    if not can_view_customer_reimbursement():
        abort(403)
    q = request.args.get("q", "").strip()
    status = request.args.get("status", "")
    date_from = request.args.get("date_from", "")
    date_to = request.args.get("date_to", "")
    region_code, country_code, countries, regions, location_clauses, location_params = report_location_filters()
    clauses = ["1 = 1"]
    params = []
    if not is_external_manager():
        clauses.extend(location_clauses)
        params.extend(location_params)
    else:
        clauses.append("service_orders.client_id = ?")
        params.append(g.user["client_id"])
    if q:
        clauses.append(
            """
            (
                customer_reimbursements.file_name like ?
                or service_orders.order_number like ?
                or service_orders.client_name like ?
                or service_orders.client_order_number like ?
                or coalesce(owners.name, buyers.owner) like ?
                or invoices.invoice_number like ?
            )
            """
        )
        params.extend([f"%{q}%"] * 6)
    if status in CUSTOMER_REIMBURSEMENT_STATUS_LABELS:
        clauses.append("customer_reimbursements.status = ?")
        params.append(status)
    else:
        status = ""
    if date_from:
        clauses.append("date(customer_reimbursements.created_at) >= ?")
        params.append(date_from)
    if date_to:
        clauses.append("date(customer_reimbursements.created_at) <= ?")
        params.append(date_to)
    rows = db().execute(
        f"""
        select customer_reimbursements.*,
               service_orders.order_number,
               service_orders.client_name,
               service_orders.client_order_number,
               service_orders.country_code,
               service_orders.region_code,
               coalesce(owners.name, buyers.owner) as buyer_owner,
               coalesce(country_local.name, country_zh.name, service_orders.country_code) as country_name,
               coalesce(country_local.region_name, country_zh.region_name, service_orders.region_code) as region_name,
               creators.name as creator_name,
               reviewers.name as reviewer_name,
               invoices.invoice_number
        from customer_reimbursements
        join service_orders on service_orders.id = customer_reimbursements.service_order_id
        left join buyers on buyers.id = service_orders.buyer_id
        left join owners on owners.id = buyers.owner_id
        left join users creators on creators.id = customer_reimbursements.created_by
        left join users reviewers on reviewers.id = customer_reimbursements.reviewed_by
        left join invoices on invoices.id = customer_reimbursements.invoice_id
        left join country_translations country_local
          on country_local.country_code = service_orders.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = service_orders.country_code and country_zh.language_code = 'zh-CN'
        where {" and ".join(clauses)}
        order by customer_reimbursements.created_at desc, customer_reimbursements.id desc
        """,
        [current_language(), *params],
    ).fetchall()
    totals = {
        "labor_total": sum(float(row["labor_total"] or 0) for row in rows),
        "travel_total": sum(float(row["travel_total"] or 0) for row in rows),
        "mileage_total": sum(float(row["mileage_total"] or 0) for row in rows),
        "total_amount": sum(float(row["total_amount"] or 0) for row in rows),
    }
    return render_template(
        "customer_reimbursement_query.html",
        rows=rows,
        totals=totals,
        q=q,
        status=status,
        date_from=date_from,
        date_to=date_to,
        countries=countries,
        regions=regions,
        region_code=region_code,
        country_code=country_code,
        labels=CUSTOMER_REIMBURSEMENT_STATUS_LABELS,
        show_location_filters=not is_external_manager(),
    )


@app.route("/reports/expenses")
@login_required
def expense_query():
    q = request.args.get("q", "").strip()
    status = request.args.get("status", "")
    payout_status = request.args.get("payout_status", "")
    date_from = request.args.get("date_from", "")
    date_to = request.args.get("date_to", "")
    region_code, country_code, countries, regions, location_clauses, location_params = report_location_filters()
    clauses = ["1 = 1"]
    params = []
    clauses.extend(location_clauses)
    params.extend(location_params)
    if normalized_role() == "employee":
        clauses.append("expenses.created_by = ?")
        params.append(g.user["id"])
    if q:
        clauses.append("(expenses.expense_number like ? or expenses.project like ? or service_orders.order_number like ? or service_orders.client_name like ?)")
        params.extend([f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"])
    if status:
        clauses.append("expenses.status = ?")
        params.append(status)
    if payout_status in EXPENSE_PAYOUT_LABELS:
        clauses.append("expenses.payout_status = ?")
        params.append(payout_status)
    else:
        payout_status = ""
    if date_from:
        clauses.append("expenses.expense_date >= ?")
        params.append(date_from)
    if date_to:
        clauses.append("expenses.expense_date <= ?")
        params.append(date_to)
    rows = db().execute(
        f"""
        select expenses.*, service_orders.order_number, service_orders.client_name, users.name as creator_name,
               coalesce(country_local.name, country_zh.name, service_orders.country_code) as country_name,
               coalesce(country_local.region_name, country_zh.region_name, service_orders.region_code) as region_name
        from expenses
        join service_orders on service_orders.id = expenses.service_order_id
        left join users on users.id = expenses.created_by
        left join country_translations country_local
          on country_local.country_code = service_orders.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = service_orders.country_code and country_zh.language_code = 'zh-CN'
        where {" and ".join(clauses)}
        order by expenses.expense_date desc, expenses.id desc
        """,
        [current_language(), *params],
    ).fetchall()
    total = sum(
        float(row["amount"] or 0)
        for row in rows
        if row["status"] == "approved" and row["payout_status"] != "paid"
    )
    return render_template(
        "expense_query.html",
        rows=rows,
        q=q,
        status=status,
        payout_status=payout_status,
        date_from=date_from,
        date_to=date_to,
        total=total,
        labels=EXPENSE_STATUS_LABELS,
        payout_labels=EXPENSE_PAYOUT_LABELS,
        countries=countries,
        regions=regions,
        region_code=region_code,
        country_code=country_code,
    )


@app.route("/expense-processing")
@login_required
def expense_processing():
    if not is_internal_user():
        abort(403)
    q = request.args.get("q", "").strip()
    payout_status = request.args.get("payout_status", "").strip()
    clauses = ["expenses.status = 'approved'"]
    params = []
    if q:
        clauses.append(
            """
            (expenses.expense_number like ? or service_orders.order_number like ?
             or service_orders.client_name like ? or users.name like ?)
            """
        )
        params.extend([f"%{q}%"] * 4)
    if payout_status in EXPENSE_PAYOUT_LABELS:
        clauses.append("expenses.payout_status = ?")
        params.append(payout_status)
    else:
        payout_status = ""
    rows = db().execute(
        f"""
        select expenses.*, service_orders.order_number, service_orders.client_name,
               users.name as creator_name, reimbursers.name as reimbursed_by_name
        from expenses
        join service_orders on service_orders.id = expenses.service_order_id
        left join users on users.id = expenses.created_by
        left join users as reimbursers on reimbursers.id = expenses.reimbursed_by
        where {" and ".join(clauses)}
        order by
            case when expenses.payout_status = 'pending' then 0 else 1 end,
            expenses.expense_date desc,
            expenses.id desc
        """,
        params,
    ).fetchall()
    totals = db().execute(
        """
        select
            coalesce(sum(case when payout_status != 'paid' then amount else 0 end), 0) as pending_total,
            coalesce(sum(case when payout_status = 'paid' then amount else 0 end), 0) as paid_total
        from expenses
        where status = 'approved'
        """
    ).fetchone()
    return render_template(
        "expense_processing.html",
        rows=rows,
        q=q,
        payout_status=payout_status,
        payout_labels=EXPENSE_PAYOUT_LABELS,
        pending_total=totals["pending_total"],
        paid_total=totals["paid_total"],
    )


@app.post("/expense-processing/action")
@login_required
def process_expense_action():
    if not is_internal_user():
        abort(403)
    try:
        expense_id = int(request.form.get("expense_id", ""))
    except (TypeError, ValueError):
        flash("请选择一张报销单据。", "error")
        return redirect(url_for("expense_processing"))
    expense = db().execute("select * from expenses where id = ?", (expense_id,)).fetchone()
    if not expense:
        abort(404)
    action = request.form.get("action", "")
    if action == "reimburse":
        if not has_action_permission("expenses", "approve"):
            abort(403)
        if expense["status"] != "approved":
            flash("只有已审核通过的报销可以发放。", "error")
            return redirect(url_for("expense_processing"))
        if expense["payout_status"] == "paid":
            flash("这张报销单已经完成报销。", "error")
            return redirect(url_for("expense_processing"))
        cursor = db().execute(
            """
            update expenses
            set payout_status = 'paid', reimbursed_by = ?, reimbursed_at = ?, updated_at = ?
            where id = ? and status = 'approved' and payout_status != 'paid'
            """,
            (g.user["id"], now(), now(), expense_id),
        )
        if cursor.rowcount != 1:
            db().rollback()
            flash("这张报销单已经完成报销或流程状态已变化。", "error")
            return redirect(url_for("expense_processing"))
        message = f"报销 {expense['expense_number']} 已完成发放，金额 {money(expense['amount'], expense['currency'])}。"
        create_message(
            expense["created_by"],
            "报销已发放",
            message,
            url_for("expense_detail", expense_id=expense_id),
        )
        log_action("reimburse", "expense", expense_id, expense["expense_number"], message)
        db().commit()
        flash("报销已标记为已报销。", "success")
    elif action == "reset_payout":
        if normalized_role() != "admin":
            abort(403)
        db().execute(
            """
            update expenses
            set payout_status = 'pending', reimbursed_by = null, reimbursed_at = null, updated_at = ?
            where id = ?
            """,
            (now(), expense_id),
        )
        log_action("reset", "expense", expense_id, expense["expense_number"], "发放状态重置为待报销")
        db().commit()
        flash("发放状态已重置为待报销。", "success")
    elif action == "reset_workflow":
        if normalized_role() != "admin":
            abort(403)
        db().execute(
            """
            update expenses
            set status = 'submitted', return_reason = null,
                reviewed_by = null, reviewed_at = null,
                payout_status = 'pending', reimbursed_by = null, reimbursed_at = null,
                updated_at = ?
            where id = ?
            """,
            (now(), expense_id),
        )
        log_action("reset", "expense", expense_id, expense["expense_number"], "流程状态重置为待经理审核，发放状态重置为待报销")
        db().commit()
        flash("流程状态和发放状态已重置。", "success")
    else:
        flash("请选择要执行的操作。", "error")
    return redirect(url_for("expense_processing"))


@app.route("/reports/audit-logs")
@login_required
def audit_log_report():
    if not can_view_audit_logs():
        abort(403)
    entity_labels = {
        "service_order": "工单",
        "invoice": "发票",
        "service_report": "工作日报",
        "expense": "报销",
        "customer_reimbursement": "工单结算",
        "user": "用户",
        "database": "数据库",
    }
    action_labels = {
        "create": "创建",
        "update": "修改",
        "delete": "删除",
        "submit": "提交审核",
        "approve": "审核通过",
        "return": "退回",
        "void": "作废",
        "mark_paid": "核销",
        "unmark_paid": "取消核销",
        "status_change": "调整状态",
        "reimburse": "报销发放",
        "reset": "重置状态",
        "send": "发送邮件",
        "login": "登录",
        "sql": "执行 SQL",
    }
    q = request.args.get("q", "").strip()
    entity_type = request.args.get("entity_type", "").strip()
    action = request.args.get("action", "").strip()
    date_from = request.args.get("date_from", "")
    date_to = request.args.get("date_to", "")
    clauses = ["1 = 1"]
    params = []
    if q:
        clauses.append("(user_name like ? or entity_label like ? or summary like ?)")
        params.extend([f"%{q}%", f"%{q}%", f"%{q}%"])
    if entity_type in entity_labels:
        clauses.append("entity_type = ?")
        params.append(entity_type)
    else:
        entity_type = ""
    if action in action_labels:
        clauses.append("action = ?")
        params.append(action)
    else:
        action = ""
    if date_from:
        clauses.append("date(created_at) >= ?")
        params.append(date_from)
    if date_to:
        clauses.append("date(created_at) <= ?")
        params.append(date_to)
    rows = db().execute(
        f"""
        select * from audit_logs
        where {" and ".join(clauses)}
        order by datetime(created_at) desc, id desc
        limit 1000
        """,
        params,
    ).fetchall()
    return render_template(
        "audit_logs.html",
        rows=rows,
        q=q,
        entity_type=entity_type,
        action=action,
        date_from=date_from,
        date_to=date_to,
        entity_labels=entity_labels,
        action_labels=action_labels,
    )


@app.route("/service-orders")
@login_required
def service_orders():
    q = request.args.get("q", "").strip()
    buyer_id = request.args.get("buyer_id", "").strip()
    clauses = ["1 = 1"]
    params = []
    if is_external_manager():
        if not g.user["client_id"]:
            clauses.append("1 = 0")
        else:
            clauses.append("service_orders.client_id = ?")
            params.append(g.user["client_id"])
    elif is_external_employee():
        clauses.append(
            "service_orders.id in (select service_order_id from user_service_orders where user_id = ?)"
        )
        params.append(g.user["id"])
    if q:
        clauses.append(
            "(service_orders.order_number like ? or service_orders.client_name like ? or coalesce(owners.name, buyers.owner) like ? or service_orders.client_order_number like ?)"
        )
        params.extend([f"%{q}%", f"%{q}%", f"%{q}%", f"%{q}%"])
    if buyer_id.isdigit():
        clauses.append("service_orders.buyer_id = ?")
        params.append(int(buyer_id))
    rows = db().execute(
        f"""
        select service_orders.*,
               work_order_types.name as work_order_type_name,
               coalesce(owners.name, buyers.owner) as buyer_owner,
               buyers.equipment_manufacturer as buyer_equipment_manufacturer,
               contracts.contract_number,
               count(distinct service_reports.id) as report_count,
               count(distinct invoices.id) as invoice_count
        from service_orders
        left join work_order_types on work_order_types.id = service_orders.work_order_type_id
        left join buyers on buyers.id = service_orders.buyer_id
        left join owners on owners.id = buyers.owner_id
        left join contracts on contracts.id = service_orders.contract_id
        left join service_reports on service_reports.service_order_id = service_orders.id
        left join invoices on invoices.service_order_id = service_orders.id and invoices.status != 'void'
        where {" and ".join(clauses)}
        group by service_orders.id
        order by service_orders.created_at desc, service_orders.id desc
        """,
        params,
    ).fetchall()
    return render_template("service_orders.html", orders=rows, q=q, buyer_id=buyer_id)


def buyer_map_payload(buyer):
    payload = {
        "id": buyer["id"],
        "buyer_number": buyer["buyer_number"],
        "name": buyer["name"],
        "owner": buyer["owner_name"],
        "contact_name": buyer["contact_name"],
        "contact_details": buyer["contact_details"],
        "email": buyer["email"],
        "country": buyer["country_name"] or buyer["country_code"] or buyer["country"],
        "country_code": buyer["country_code"],
        "detailed_address": buyer["detailed_address"],
        "equipment_manufacturer": buyer["equipment_manufacturer"],
        "latitude": buyer["latitude"],
        "longitude": buyer["longitude"],
        "geocode_status": buyer["geocode_status"] or "pending",
        "work_order_total": buyer["work_order_total"],
        "work_order_completed": buyer["work_order_completed"],
        "status": (
            "completed"
            if buyer["work_order_total"] and buyer["work_order_total"] == buyer["work_order_completed"]
            else "active"
        ),
        "detail_url": url_for("service_orders", buyer_id=buyer["id"]),
    }
    if can_manage_buyers():
        payload["edit_url"] = url_for("edit_buyer", buyer_id=buyer["id"])
    if can_view_invoices():
        payload["paid_invoice_amount"] = buyer["paid_invoice_amount"]
        payload["completed_invoice_amount"] = buyer["completed_invoice_amount"]
    return payload


def buyer_map_rows(where_clause="1 = 1", params=()):
    return db().execute(
        f"""
        with order_stats as (
            select buyer_id,
                   count(*) as work_order_total,
                   sum(case when status = 'closed' then 1 else 0 end) as work_order_completed
            from service_orders
            where buyer_id is not null
            group by buyer_id
        ),
        invoice_amounts as (
            select invoices.id, invoices.service_order_id, invoices.status, invoices.paid_at,
                   invoices.payment_amount,
                   coalesce(sum(invoice_items.amount * (1 + invoice_items.tax_rate / 100.0)), 0) as invoice_total
            from invoices
            left join invoice_items on invoice_items.invoice_id = invoices.id
            where invoices.status != 'void'
            group by invoices.id
        ),
        invoice_stats as (
            select service_orders.buyer_id,
                   sum(
                       case when invoice_amounts.paid_at is not null
                       then coalesce(invoice_amounts.payment_amount, invoice_amounts.invoice_total)
                       else 0 end
                   ) as paid_invoice_amount,
                   sum(
                       case when invoice_amounts.status = 'completed'
                       then invoice_amounts.invoice_total
                       else 0 end
                   ) as completed_invoice_amount
            from service_orders
            join invoice_amounts on invoice_amounts.service_order_id = service_orders.id
            where service_orders.buyer_id is not null
            group by service_orders.buyer_id
        )
        select buyers.*,
               coalesce(owners.name, buyers.owner) as owner_name,
               owners.owner_number as owner_number,
               coalesce(country_local.name, country_zh.name, country_en.name, buyers.country, buyers.country_code) as country_name,
               coalesce(order_stats.work_order_total, 0) as work_order_total,
               coalesce(order_stats.work_order_completed, 0) as work_order_completed,
               coalesce(invoice_stats.paid_invoice_amount, 0) as paid_invoice_amount,
               coalesce(invoice_stats.completed_invoice_amount, 0) as completed_invoice_amount
        from buyers
        left join owners on owners.id = buyers.owner_id
        left join country_translations country_local
          on country_local.country_code = buyers.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = buyers.country_code and country_zh.language_code = 'zh-CN'
        left join country_translations country_en
          on country_en.country_code = buyers.country_code and country_en.language_code = 'en'
        left join order_stats on order_stats.buyer_id = buyers.id
        left join invoice_stats on invoice_stats.buyer_id = buyers.id
        where {where_clause}
        order by owner_name, buyers.name, buyers.id
        """,
        (current_language(), *params),
    ).fetchall()


@app.route("/service-orders/map")
@login_required
def service_order_map():
    if not is_internal_user() and not is_external_manager():
        abort(403)
    rows = (
        buyer_map_rows("buyers.client_id = ?", (g.user["client_id"],))
        if is_external_manager()
        else buyer_map_rows()
    )
    buyers_payload = [buyer_map_payload(row) for row in rows]
    google_maps_browser_api_key = get_google_maps_browser_api_key()
    company = get_company_profile()
    route_origin_address = (g.user["address"] or "").strip() or company["address"]
    return render_template(
        "service_order_map.html",
        map_buyers=buyers_payload,
        show_invoice_amounts=can_view_invoices(),
        headquarters={
            "name": "Prasinos Power LLC",
            "address": "518 Anacacho Dr, Spring, TX 77386",
            "latitude": 30.11295,
            "longitude": -95.41663,
        },
        company_address=company["address"],
        route_origin_address=route_origin_address,
        geocoding_enabled=GEOCODING_ENABLED,
        google_maps_enabled=bool(google_maps_browser_api_key),
        google_maps_browser_api_key=google_maps_browser_api_key,
    )


@app.post("/service-orders/map/geocode-next")
@login_required
def geocode_next_service_order():
    if not is_internal_user() and not is_external_manager():
        abort(403)
    if not GEOCODING_ENABLED:
        return jsonify({"available": False, "remaining": 0}), 503
    client_clause = "and client_id = ?" if is_external_manager() else ""
    geocode_params = [GEOCODER_VERSION]
    if is_external_manager():
        geocode_params.append(g.user["client_id"])
    buyer = db().execute(
        f"""
        select id from buyers
        where (
           geocode_status is null or geocode_status = 'pending'
           or geocode_address is null
           or lower(trim(geocode_address)) != lower(trim(detailed_address))
           or coalesce(geocode_version, '') != ?
        )
        {client_clause}
        order by created_at desc, id desc
        limit 1
        """,
        geocode_params,
    ).fetchone()
    if not buyer:
        return jsonify({"available": True, "remaining": 0, "buyer": None})
    try:
        geocode_buyer(buyer["id"])
    except TemporaryGeocodingError as error:
        db().rollback()
        app.logger.warning("Temporary geocoding failure for buyer %s: %s", buyer["id"], error)
        return jsonify({"available": True, "temporary_error": True}), 503
    db().commit()
    refreshed = buyer_map_rows("buyers.id = ?", (buyer["id"],))[0]
    remaining = db().execute(
        f"""
        select count(*) as count from buyers
        where (
           geocode_status is null or geocode_status = 'pending'
           or geocode_address is null
           or lower(trim(geocode_address)) != lower(trim(detailed_address))
           or coalesce(geocode_version, '') != ?
        )
        {client_clause}
        """,
        geocode_params,
    ).fetchone()["count"]
    return jsonify(
        {
            "available": True,
            "remaining": remaining,
            "buyer": buyer_map_payload(refreshed),
        }
    )


@app.post("/service-orders/map/retry-failed")
@login_required
def retry_failed_service_order_geocodes():
    if not is_internal_user() and not is_external_manager():
        abort(403)
    client_clause = "and client_id = ?" if is_external_manager() else ""
    params = [g.user["client_id"]] if is_external_manager() else []
    db().execute(
        f"""
        update buyers
        set geocode_status = 'pending', geocode_attempted_at = null, geocode_version = null
        where geocode_status = 'failed'
        {client_clause}
        """
        ,
        params,
    )
    db().commit()
    return jsonify({"ok": True})


@app.route("/service-orders/new", methods=["GET", "POST"])
@login_required
def new_service_order():
    if not can_create_service_order():
        abort(403)
    buyers_rows = db().execute(
        """
        select buyers.*, coalesce(owners.name, buyers.owner) as owner_name,
               coalesce(country_local.name, country_zh.name, country_en.name, buyers.country, buyers.country_code) as country_name
        from buyers
        left join owners on owners.id = buyers.owner_id
        left join country_translations country_local
          on country_local.country_code = buyers.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = buyers.country_code and country_zh.language_code = 'zh-CN'
        left join country_translations country_en
          on country_en.country_code = buyers.country_code and country_en.language_code = 'en'
        order by owner_name, buyers.name, buyers.buyer_number
        """,
        (current_language(),),
    ).fetchall()
    clients_rows = db().execute("select * from clients order by client_number, name").fetchall()
    work_order_types_rows = db().execute(
        "select * from work_order_types where is_active = 1 order by code, name"
    ).fetchall()
    contracts_rows = db().execute(
        """
        select contracts.*, clients.name as client_name
        from contracts join clients on clients.id = contracts.client_id
        where contracts.status in ('signed', 'active')
        order by clients.name, contracts.contract_number
        """
    ).fetchall()
    if not buyers_rows:
        flash("请先由会计或经理维护站点资料。", "error")
        return redirect(url_for("buyers") if can_manage_buyers() else url_for("service_orders"))
    if not clients_rows:
        flash("请先由管理员或经理维护客户资料。", "error")
        return redirect(url_for("clients") if is_manager() else url_for("service_orders"))
    if not work_order_types_rows:
        flash("请先由管理员或经理维护工单类型。", "error")
        return redirect(url_for("work_order_types") if is_manager() else url_for("service_orders"))
    if request.method == "POST":
        country = country_from_form()
        buyer = db().execute(
            "select * from buyers where id = ?",
            (request.form.get("buyer_id"),),
        ).fetchone()
        work_order_type = db().execute(
            "select * from work_order_types where id = ? and is_active = 1",
            (request.form.get("work_order_type_id"),),
        ).fetchone()
        client = db().execute(
            "select * from clients where id = ?",
            (request.form.get("client_id"),),
        ).fetchone()
        contract_id = request.form.get("contract_id") or None
        contract = None
        if contract_id:
            contract = db().execute(
                "select * from contracts where id = ?",
                (contract_id,),
            ).fetchone()
        site_address = request.form.get("site_address", "").strip()
        client_order_number = request.form.get("client_order_number", "").strip()
        if not buyer or not client or not work_order_type or not site_address or not client_order_number:
            flash("请选择客户、站点和工单类型，并填写服务现场地址、服务订单号码。", "error")
            return redirect(url_for("new_service_order"))
        if contract_id and (not contract or contract["client_id"] != client["id"]):
            flash("关联合同必须属于所选客户。", "error")
            return redirect(url_for("new_service_order"))
        order_number = next_service_order_number()
        cursor = db().execute(
            """
            insert into service_orders (
                order_number, client_id, contract_id, buyer_id, client_name, buyer_contact_name, buyer_contact_details,
                site_address, client_order_number, start_date, work_order_type_id,
                region_code, country_code, created_by, created_at
            ) values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                order_number,
                client["id"],
                contract["id"] if contract else None,
                buyer["id"],
                buyer["name"],
                buyer["contact_name"],
                buyer["contact_details"],
                site_address,
                client_order_number,
                request.form.get("start_date") or None,
                work_order_type["id"],
                country["region_code"],
                country["code"],
                g.user["id"],
                now(),
            ),
        )
        log_action("create", "service_order", cursor.lastrowid, order_number, f"站点：{buyer['name']}")
        folder_created = True
        try:
            ensure_service_order_picture_folder(order_number)
        except OSError:
            folder_created = False
            app.logger.exception("Failed to create NAS pictures folder for service order %s", order_number)
            flash("工单已创建，但 NAS 照片文件夹创建失败，请检查共享照片目录挂载或权限。", "error")
        db().commit()
        if folder_created:
            flash("工单已创建，NAS 照片文件夹已自动创建。", "success")
        return redirect(url_for("service_orders"))
    return render_template(
        "service_order_form.html",
        order=None,
        clients=clients_rows,
        buyers=buyers_rows,
        work_order_types=work_order_types_rows,
        contracts=contracts_rows,
        form_title="新建工单",
        start_date_only=False,
        countries=country_rows(),
    )


@app.route("/service-orders/<int:order_id>/edit", methods=["GET", "POST"])
@login_required
def edit_service_order(order_id):
    order = require_service_order(order_id)
    if normalized_role() in {"employee", "external_employee", "external_manager"} and order["status"] == "closed":
        flash("已关闭的工单不能编辑。", "error")
        return redirect(url_for("service_order_detail", order_id=order_id))
    if is_external_user():
        if request.method == "POST":
            db().execute(
                "update service_orders set start_date = ? where id = ?",
                (request.form.get("start_date") or None, order_id),
            )
            log_action(
                "update",
                "service_order",
                order_id,
                order["order_number"],
                "修改工单开始日期",
            )
            db().commit()
            flash("工单开始日期已保存。", "success")
            return redirect(url_for("service_order_detail", order_id=order_id))
        return render_template(
            "service_order_form.html",
            order=order,
            clients=[],
            buyers=[],
            work_order_types=[],
            form_title="修改工单开始日期",
            start_date_only=True,
            countries=[],
        )
    buyers_rows = db().execute(
        """
        select buyers.*, coalesce(owners.name, buyers.owner) as owner_name,
               coalesce(country_local.name, country_zh.name, country_en.name, buyers.country, buyers.country_code) as country_name
        from buyers
        left join owners on owners.id = buyers.owner_id
        left join country_translations country_local
          on country_local.country_code = buyers.country_code and country_local.language_code = ?
        left join country_translations country_zh
          on country_zh.country_code = buyers.country_code and country_zh.language_code = 'zh-CN'
        left join country_translations country_en
          on country_en.country_code = buyers.country_code and country_en.language_code = 'en'
        order by owner_name, buyers.name, buyers.buyer_number
        """,
        (current_language(),),
    ).fetchall()
    clients_rows = db().execute("select * from clients order by client_number, name").fetchall()
    work_order_types_rows = db().execute(
        """
        select * from work_order_types
        where is_active = 1 or id = ?
        order by is_active desc, code, name
        """,
        (order["work_order_type_id"],),
    ).fetchall()
    contracts_rows = db().execute(
        """
        select contracts.*, clients.name as client_name
        from contracts join clients on clients.id = contracts.client_id
        where contracts.status in ('signed', 'active') or contracts.id = ?
        order by clients.name, contracts.contract_number
        """,
        (order["contract_id"] or 0,),
    ).fetchall()
    if request.method == "POST":
        country = country_from_form(order["country_code"])
        buyer = db().execute(
            "select * from buyers where id = ?",
            (request.form.get("buyer_id"),),
        ).fetchone()
        work_order_type = db().execute(
            "select * from work_order_types where id = ?",
            (request.form.get("work_order_type_id"),),
        ).fetchone()
        client = db().execute(
            "select * from clients where id = ?",
            (request.form.get("client_id"),),
        ).fetchone()
        contract_id = request.form.get("contract_id") or None
        contract = None
        if contract_id:
            contract = db().execute("select * from contracts where id = ?", (contract_id,)).fetchone()
        site_address = request.form.get("site_address", "").strip()
        if not buyer or not client or not work_order_type or not site_address:
            flash("请选择有效的客户、站点和工单类型，并填写服务现场地址。", "error")
            return redirect(url_for("edit_service_order", order_id=order_id))
        if contract_id and (not contract or contract["client_id"] != client["id"]):
            flash("关联合同必须属于所选客户。", "error")
            return redirect(url_for("edit_service_order", order_id=order_id))
        db().execute(
            """
            update service_orders
            set client_id = ?, contract_id = ?, buyer_id = ?, client_name = ?, buyer_contact_name = ?, buyer_contact_details = ?,
                site_address = ?, client_order_number = ?, start_date = ?,
                work_order_type_id = ?, status = ?, region_code = ?, country_code = ?
            where id = ?
            """,
            (
                client["id"],
                contract["id"] if contract else None,
                buyer["id"],
                buyer["name"],
                buyer["contact_name"],
                buyer["contact_details"],
                site_address,
                request.form.get("client_order_number", "").strip(),
                request.form.get("start_date") or None,
                work_order_type["id"],
                request.form.get("status", "open"),
                country["region_code"],
                country["code"],
                order_id,
            ),
        )
        log_action("update", "service_order", order_id, order["order_number"], "修改工单信息")
        db().commit()
        flash("工单已保存。", "success")
        return redirect(url_for("service_order_detail", order_id=order_id))
    return render_template(
        "service_order_form.html",
        order=order,
        clients=clients_rows,
        buyers=buyers_rows,
        work_order_types=work_order_types_rows,
        contracts=contracts_rows,
        form_title="编辑工单",
        start_date_only=False,
        countries=country_rows(),
    )


@app.route("/service-orders/<int:order_id>")
@login_required
def service_order_detail(order_id):
    order = require_service_order(order_id)
    reports_rows = db().execute(
        """
        select service_reports.*, group_concat(worker_users.name, ', ') as worker_names,
               creator_users.name as creator_name
        from service_reports
        left join service_report_workers on service_report_workers.report_id = service_reports.id
        left join users as worker_users on worker_users.id = service_report_workers.user_id
        left join users as creator_users on creator_users.id = service_reports.created_by
        where service_reports.service_order_id = ?
        group by service_reports.id
        order by service_reports.report_date desc, service_reports.id desc
        """,
        (order_id,),
    ).fetchall()
    invoices_rows = db().execute(
        """
        select invoices.*, clients.name as client_name
        from invoices left join clients on clients.id = invoices.client_id
        where invoices.service_order_id = ? and invoices.status != 'void'
        order by invoices.issue_date desc, invoices.id desc
        """,
        (order_id,),
    ).fetchall()
    if is_external_user():
        expenses_rows = []
    else:
        expense_access = "" if normalized_role() in {"admin", "manager", "finance"} else "and expenses.created_by = ?"
        expense_params = [order_id] if not expense_access else [order_id, g.user["id"]]
        expenses_rows = db().execute(
            f"""
            select expenses.*, users.name as creator_name
            from expenses left join users on users.id = expenses.created_by
            where expenses.service_order_id = ? {expense_access}
            order by expenses.created_at desc, expenses.id desc
            """,
            expense_params,
        ).fetchall()
    customer_reimbursement = latest_customer_reimbursement(order_id) if can_view_customer_reimbursement() else None
    delete_blockers = service_order_delete_blockers(order_id)
    return render_template(
        "service_order_detail.html",
        order=order,
        reports=reports_rows,
        invoices=invoices_rows,
        expenses=expenses_rows,
        customer_reimbursement=customer_reimbursement,
        can_delete_order=can_delete_service_order() and not delete_blockers,
        delete_blockers=delete_blockers,
        labels=STATUS_LABELS,
        expense_labels=EXPENSE_STATUS_LABELS,
    )


@app.post("/service-orders/<int:order_id>/delete")
@login_required
def delete_service_order(order_id):
    order = require_service_order(order_id)
    if not can_delete_service_order():
        abort(403)
    blockers = service_order_delete_blockers(order_id)
    if blockers:
        flash(f"这个工单已有{', '.join(blockers)}，不能删除。", "error")
        return redirect(url_for("service_order_detail", order_id=order_id))
    db().execute("delete from user_service_orders where service_order_id = ?", (order_id,))
    db().execute("delete from service_orders where id = ?", (order_id,))
    log_action("delete", "service_order", order_id, order["order_number"], f"站点：{order['client_name']}")
    db().commit()
    flash("工单已删除。", "success")
    return redirect(url_for("service_orders"))


@app.post("/service-orders/<int:order_id>/import-google-photos")
@login_required
def import_google_photos_for_service_order(order_id):
    order = require_service_order(order_id)
    if not can_create_service_report(order):
        abort(403)
    share_url = request.form.get("google_photos_url", "").strip()
    if not is_google_photos_share_url(share_url):
        flash("请输入有效的 Google Photos 分享链接。", "error")
        return redirect(url_for("service_order_detail", order_id=order_id))
    try:
        ensure_service_order_picture_folder(order["order_number"])
        service_order_incoming_folder(order["order_number"], "google-photos")
    except OSError:
        app.logger.exception("Failed to prepare Google Photos import folder for service order %s", order["order_number"])
        flash("无法写入 NAS 照片目录，请检查共享照片目录挂载或权限。", "error")
        return redirect(url_for("service_order_detail", order_id=order_id))
    start_google_photos_import_job(order_id, g.user["id"], share_url)
    flash("Google Photos 分享链接导入任务已开始。完成后会通过邮件和系统消息通知你。", "success")
    return redirect(url_for("service_order_detail", order_id=order_id))


@app.route("/service-orders/<int:order_id>/customer-reimbursement", methods=["GET", "POST"])
@login_required
def customer_reimbursement_form(order_id):
    if not can_view_customer_reimbursement():
        abort(403)
    order = require_service_order(order_id)
    if request.method == "POST" and not can_manage_customer_reimbursement():
        abort(403)
    start_date_redirect = require_service_order_start_date(order)
    if start_date_redirect:
        return start_date_redirect
    reimbursement = latest_customer_reimbursement(order_id)
    if not reimbursement:
        if not can_manage_customer_reimbursement():
            flash("该工单还没有工单结算。", "error")
            return redirect(url_for("service_order_detail", order_id=order_id))
        try:
            reimbursement = create_customer_reimbursement(order)
            build_customer_reimbursement_pdf(reimbursement, order, customer_reimbursement_items(reimbursement["id"]))
            log_action("create", "customer_reimbursement", reimbursement["id"], reimbursement["file_name"], f"工单：{order['order_number']}")
            db().commit()
            flash("工单结算已根据工作日报生成，请补充费用和附件。", "success")
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("service_order_detail", order_id=order_id))
    else:
        reimbursement = ensure_customer_reimbursement_pdf_record(reimbursement, order)
        db().commit()
    if request.method == "POST":
        try:
            action = request.form.get("action", "save")
            if reimbursement["status"] not in {"draft", "returned"}:
                if action == "generate_pdf":
                    path = build_customer_reimbursement_pdf(
                        reimbursement,
                        order,
                        customer_reimbursement_items(reimbursement["id"]),
                    )
                    return send_file(path, as_attachment=True, download_name=reimbursement["file_name"])
                if action == "send_email":
                    if reimbursement["status"] != "approved":
                        raise ValueError("工单结算审核通过后才能发送邮件。")
                    recipient = deliver_customer_reimbursement_email(reimbursement, order)
                    db().commit()
                    flash(f"工单结算已发送至 {recipient}。", "success")
                    return redirect(url_for("customer_reimbursement_form", order_id=order_id))
                if action == "generate_invoice":
                    if reimbursement["status"] != "approved":
                        raise ValueError("工单结算审核通过后才能生成发票。")
                    if not can_create_invoice():
                        abort(403)
                    linked_invoice = service_order_active_invoice(order_id)
                    if linked_invoice:
                        if not reimbursement["invoice_id"]:
                            db().execute(
                                "update customer_reimbursements set invoice_id = ? where id = ?",
                                (linked_invoice["id"], reimbursement["id"]),
                            )
                            db().commit()
                        return redirect(url_for("invoice_detail", invoice_id=linked_invoice["id"]))
                    return redirect(
                        url_for(
                            "new_invoice",
                            service_order_id=order["id"],
                            customer_reimbursement_id=reimbursement["id"],
                        )
                    )
                raise ValueError("只有保存未提交或已退回的工单结算可以修改。")
            rows = customer_reimbursement_items_from_form()
            save_customer_reimbursement_items(reimbursement["id"], rows)
            totals = update_customer_reimbursement_totals(reimbursement["id"], rows)
            save_customer_reimbursement_uploads(reimbursement["id"])
            next_status = "submitted" if action == "submit" else reimbursement["status"]
            db().execute(
                """
                update customer_reimbursements
                set status = ?, return_reason = null
                where id = ?
                """,
                (next_status, reimbursement["id"]),
            )
            reimbursement = db().execute("select * from customer_reimbursements where id = ?", (reimbursement["id"],)).fetchone()
            remove_customer_reimbursement_pdf(reimbursement)
            if action == "generate_pdf":
                build_customer_reimbursement_pdf(reimbursement, order, customer_reimbursement_items(reimbursement["id"]))
            log_action(
                "update",
                "customer_reimbursement",
                reimbursement["id"],
                reimbursement["file_name"],
                f"金额：{money(totals['total_amount'])}",
            )
            db().commit()
            if action == "submit":
                notify_role(
                    ["admin", "manager"],
                    "新工单结算待审核",
                    (
                        f"{g.user['name']}提交了工单 {order['order_number']} 的工单结算，"
                        f"金额 {money(totals['total_amount'])}。"
                    ),
                    url_for("customer_reimbursement_form", order_id=order_id),
                    exclude_user_ids={g.user["id"]},
                )
                log_action(
                    "submit",
                    "customer_reimbursement",
                    reimbursement["id"],
                    reimbursement["file_name"],
                    f"金额：{money(totals['total_amount'])}",
                )
                db().commit()
                flash("工单结算已提交经理审核。", "success")
                return redirect(url_for("customer_reimbursement_form", order_id=order_id))
            if action == "generate_pdf":
                return send_file(
                    customer_reimbursement_file_path(reimbursement),
                    as_attachment=True,
                    download_name=reimbursement["file_name"],
                )
            if action == "send_email":
                if reimbursement["status"] != "approved":
                    raise ValueError("工单结算审核通过后才能发送邮件。")
                try:
                    recipient = deliver_customer_reimbursement_email(reimbursement, order)
                except (ValueError, RuntimeError) as error:
                    flash(str(error), "error")
                    return redirect(url_for("customer_reimbursement_form", order_id=order_id))
                db().commit()
                flash(f"工单结算已发送至 {recipient}。", "success")
                return redirect(url_for("customer_reimbursement_form", order_id=order_id))
            if action == "generate_invoice":
                if reimbursement["status"] != "approved":
                    raise ValueError("工单结算审核通过后才能生成发票。")
                if not can_create_invoice():
                    abort(403)
                linked_invoice = service_order_active_invoice(order_id)
                if linked_invoice:
                    if not reimbursement["invoice_id"]:
                        db().execute(
                            "update customer_reimbursements set invoice_id = ? where id = ?",
                            (linked_invoice["id"], reimbursement["id"]),
                        )
                        db().commit()
                    return redirect(url_for("invoice_detail", invoice_id=linked_invoice["id"]))
                return redirect(
                    url_for(
                        "new_invoice",
                        service_order_id=order["id"],
                        customer_reimbursement_id=reimbursement["id"],
                    )
                )
            flash("工单结算已保存。", "success")
            return redirect(url_for("customer_reimbursement_form", order_id=order_id))
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("customer_reimbursement_form", order_id=order_id))
    items = customer_reimbursement_items(reimbursement["id"])
    linked_invoice = customer_reimbursement_linked_invoice(reimbursement, order_id)
    if linked_invoice and reimbursement["invoice_id"] != linked_invoice["id"]:
        db().execute(
            "update customer_reimbursements set invoice_id = ? where id = ?",
            (linked_invoice["id"], reimbursement["id"]),
        )
        db().commit()
        reimbursement = db().execute("select * from customer_reimbursements where id = ?", (reimbursement["id"],)).fetchone()
    reviewer = db().execute(
        "select name from users where id = ?",
        (reimbursement["reviewed_by"],),
    ).fetchone() if reimbursement["reviewed_by"] else None
    return render_template(
        "customer_reimbursement_form.html",
        order=order,
        reimbursement=reimbursement,
        items=items,
        totals=customer_reimbursement_totals(items),
        rates=customer_reimbursement_rates(),
        attachments=get_customer_reimbursement_attachments(reimbursement["id"]),
        linked_invoice=linked_invoice,
        reviewer=reviewer,
        can_edit=can_manage_customer_reimbursement() and reimbursement["status"] in {"draft", "returned"},
    )


@app.get("/customer-reimbursements/<int:reimbursement_id>/download")
@login_required
def download_customer_reimbursement(reimbursement_id):
    reimbursement, order = require_customer_reimbursement(reimbursement_id)
    reimbursement = ensure_customer_reimbursement_pdf_record(reimbursement, order)
    path = customer_reimbursement_file_path(reimbursement)
    if not os.path.exists(path):
        build_customer_reimbursement_pdf(reimbursement, order, customer_reimbursement_items(reimbursement_id))
    db().commit()
    return send_file(path, as_attachment=True, download_name=reimbursement["file_name"])


@app.get("/customer-reimbursements/<int:reimbursement_id>/preview")
@login_required
def preview_customer_reimbursement(reimbursement_id):
    reimbursement, order = require_customer_reimbursement(reimbursement_id)
    reimbursement = ensure_customer_reimbursement_pdf_record(reimbursement, order)
    path = build_customer_reimbursement_pdf(
        reimbursement,
        order,
        customer_reimbursement_items(reimbursement_id),
    )
    db().commit()
    return send_file(
        path,
        mimetype="application/pdf",
        as_attachment=False,
        download_name=reimbursement["file_name"],
        conditional=True,
    )


@app.post("/customer-reimbursements/<int:reimbursement_id>/approve")
@login_required
def approve_customer_reimbursement(reimbursement_id):
    reimbursement, order = require_customer_reimbursement(reimbursement_id)
    if reimbursement["status"] != "submitted":
        flash("只有待经理审核的工单结算可以审核通过。", "error")
        return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))
    db().execute(
        """
        update customer_reimbursements
        set status = 'approved', return_reason = null, reviewed_by = ?, reviewed_at = ?
        where id = ?
        """,
        (g.user["id"], now(), reimbursement_id),
    )
    message = (
        f"{g.user['name']}已审核通过工单 {order['order_number']} 的工单结算，"
        f"金额 {money(reimbursement['total_amount'])}。"
    )
    create_message(
        reimbursement["created_by"],
        "工单结算已审核通过",
        message,
        url_for("customer_reimbursement_form", order_id=order["id"]),
    )
    log_action(
        "approve",
        "customer_reimbursement",
        reimbursement_id,
        reimbursement["file_name"],
        f"金额：{money(reimbursement['total_amount'])}",
    )
    db().commit()
    flash("工单结算已审核通过。", "success")
    return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))


@app.post("/customer-reimbursements/<int:reimbursement_id>/return")
@login_required
def return_customer_reimbursement(reimbursement_id):
    reimbursement, order = require_customer_reimbursement(reimbursement_id)
    if reimbursement["status"] != "submitted":
        flash("只有待经理审核的工单结算可以退回。", "error")
        return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))
    reason = request.form.get("return_reason", "").strip()
    if not reason:
        flash("请填写退回原因。", "error")
        return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))
    db().execute(
        """
        update customer_reimbursements
        set status = 'returned', return_reason = ?, reviewed_by = ?, reviewed_at = ?
        where id = ?
        """,
        (reason, g.user["id"], now(), reimbursement_id),
    )
    create_message(
        reimbursement["created_by"],
        "工单结算已被退回",
        f"工单 {order['order_number']} 的工单结算已被退回。原因：{reason}",
        url_for("customer_reimbursement_form", order_id=order["id"]),
    )
    log_action(
        "return",
        "customer_reimbursement",
        reimbursement_id,
        reimbursement["file_name"],
        f"原因：{reason}",
    )
    db().commit()
    flash("工单结算已退回。", "success")
    return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))


@app.post("/customer-reimbursements/<int:reimbursement_id>/reset")
@login_required
def reset_customer_reimbursement(reimbursement_id):
    reimbursement, order = require_customer_reimbursement(reimbursement_id)
    if not can_manage_customer_reimbursement():
        abort(403)
    if reimbursement["status"] != "approved":
        flash("只有已审核通过的工单结算可以重置。", "error")
        return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))
    linked_invoice = customer_reimbursement_linked_invoice(reimbursement, order["id"])
    if linked_invoice:
        flash("这份工单结算已经生成发票。请先删除对应发票，再重置工单结算后进行修改。", "error")
        return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))
    db().execute(
        """
        update customer_reimbursements
        set status = 'draft', return_reason = null, reviewed_by = null, reviewed_at = null,
            invoice_id = null
        where id = ?
        """,
        (reimbursement_id,),
    )
    remove_customer_reimbursement_pdf(reimbursement)
    log_action(
        "reset",
        "customer_reimbursement",
        reimbursement_id,
        reimbursement["file_name"],
        "重置为保存未提交，允许重新编辑",
    )
    db().commit()
    flash("工单结算已重置为保存未提交，可以重新编辑。修改后请重新提交经理审核。", "success")
    return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))


@app.post("/customer-reimbursements/<int:reimbursement_id>/delete")
@login_required
def delete_customer_reimbursement(reimbursement_id):
    reimbursement, order = require_customer_reimbursement(reimbursement_id)
    if not can_manage_customer_reimbursement():
        abort(403)
    if reimbursement["status"] not in {"draft", "returned"}:
        flash("只有保存未提交或已退回的工单结算草稿可以删除。", "error")
        return redirect(url_for("customer_reimbursement_form", order_id=order["id"]))
    reimbursement_dir = os.path.join(CUSTOMER_REIMBURSEMENT_DIR, str(reimbursement_id))
    db().execute("delete from customer_reimbursement_attachments where customer_reimbursement_id = ?", (reimbursement_id,))
    db().execute("delete from customer_reimbursement_items where customer_reimbursement_id = ?", (reimbursement_id,))
    db().execute("delete from customer_reimbursements where id = ?", (reimbursement_id,))
    shutil.rmtree(reimbursement_dir, ignore_errors=True)
    log_action(
        "delete",
        "customer_reimbursement",
        reimbursement_id,
        reimbursement["file_name"],
        f"工单：{order['order_number']}，状态：{CUSTOMER_REIMBURSEMENT_STATUS_LABELS.get(reimbursement['status'], reimbursement['status'])}",
    )
    db().commit()
    flash("工单结算草稿已删除。", "success")
    return redirect(url_for("service_order_detail", order_id=order["id"]))


@app.get("/customer-reimbursement-attachments/<int:attachment_id>/download")
@login_required
def download_customer_reimbursement_attachment(attachment_id):
    attachment = db().execute("select * from customer_reimbursement_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    require_customer_reimbursement(attachment["customer_reimbursement_id"])
    return send_file(customer_reimbursement_attachment_path(attachment), as_attachment=True, download_name=attachment["original_filename"])


@app.get("/customer-reimbursement-attachments/<int:attachment_id>/preview")
@login_required
def preview_customer_reimbursement_attachment(attachment_id):
    attachment = db().execute("select * from customer_reimbursement_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    require_customer_reimbursement(attachment["customer_reimbursement_id"])
    return send_file(customer_reimbursement_attachment_path(attachment), mimetype=attachment["content_type"] or None)


@app.post("/customer-reimbursement-attachments/<int:attachment_id>/delete")
@login_required
def delete_customer_reimbursement_attachment(attachment_id):
    attachment = db().execute("select * from customer_reimbursement_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    reimbursement, order = require_customer_reimbursement(attachment["customer_reimbursement_id"])
    if not can_manage_customer_reimbursement():
        abort(403)
    if reimbursement["status"] not in {"draft", "returned"}:
        flash("只有保存未提交或已退回的工单结算可以删除附件。", "error")
        return redirect(url_for("customer_reimbursement_form", order_id=order["id"], _anchor="customerReimbursementAttachmentsSection"))
    try:
        os.remove(customer_reimbursement_attachment_path(attachment))
    except FileNotFoundError:
        pass
    db().execute("delete from customer_reimbursement_attachments where id = ?", (attachment_id,))
    log_action("delete", "customer_reimbursement_attachment", attachment_id, attachment["original_filename"], f"工单：{order['order_number']}")
    db().commit()
    flash("附件已删除。", "success")
    return redirect(url_for("customer_reimbursement_form", order_id=order["id"], _anchor="customerReimbursementAttachmentsSection"))


@app.route("/service-orders/<int:order_id>/reports/new", methods=["GET", "POST"])
@login_required
def new_service_report(order_id):
    order = require_service_order(order_id)
    if not can_create_service_report(order):
        abort(403)
    if not is_external_employee():
        start_date_redirect = require_service_order_start_date(order)
        if start_date_redirect:
            return start_date_redirect
    users_rows = service_report_worker_options(order)
    if request.method == "POST":
        save_token = request.form.get("save_token", "")
        if not claim_report_save_token(save_token):
            existing = db().execute(
                "select report_id from service_report_save_tokens where token = ?",
                (save_token,),
            ).fetchone()
            if existing and existing["report_id"]:
                return redirect(url_for("edit_service_report", report_id=existing["report_id"]))
            flash("该日报正在保存，请稍候。", "error")
            return redirect(url_for("new_service_report", order_id=order_id))
        try:
            arrival_time = posted_report_time("arrival_time")
            departure_time = posted_report_time("departure_time")
            total_service_hours = calculated_report_total_service_hours(arrival_time, departure_time)
            total_time = calculated_report_total_time()
            cursor = db().execute(
                """
                insert into service_reports (
                    service_order_id, report_date, actual_work_date, total_service_hours, travel_hours, public_transport_hours,
                    driving_miles, departure_address, site_address, total_time, cabinet_number,
                    arrival_time, departure_time, service_description, created_by, created_at, updated_at
                ) values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    order_id,
                    request.form.get("report_date"),
                    request.form.get("actual_work_date") or request.form.get("report_date"),
                    total_service_hours,
                    to_float(request.form.get("travel_hours")),
                    to_float(request.form.get("public_transport_hours")),
                    to_float(request.form.get("driving_miles")),
                    request.form.get("departure_address", "").strip(),
                    request.form.get("site_address", "").strip(),
                    total_time,
                    request.form.get("cabinet_number", "").strip(),
                    arrival_time,
                    departure_time,
                    request.form.get("service_description", "").strip(),
                    g.user["id"],
                    now(),
                    now(),
                ),
            )
            report_id = cursor.lastrowid
            finish_report_save_token(save_token, report_id)
            save_report_detail_rows(report_id)
            save_report_uploads(report_id)
            log_action(
                "create",
                "service_report",
                report_id,
                f"{order['order_number']} / {request.form.get('report_date')}",
                "创建工作日报",
            )
            db().commit()
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("new_service_report", order_id=order_id))
        flash("工作日报已保存。", "success")
        return redirect(url_for("edit_service_report", report_id=report_id))
    return render_template(
        "service_report_form.html",
        order=order,
        report=report_form_defaults(order=order),
        users=users_rows,
        selected_workers=[],
        selected_worker_miles={},
        saved_parts=[{} for _ in range(4)],
        replaced_parts=[{} for _ in range(4)],
        attachments=get_report_attachments(0),
        save_token=secrets.token_urlsafe(24),
        is_edit=False,
        can_edit_report=True,
    )


@app.route("/service-reports/<int:report_id>/edit", methods=["GET", "POST"])
@login_required
def edit_service_report(report_id):
    report, order = require_service_report(report_id)
    if is_external_employee() and report["created_by"] != g.user["id"]:
        abort(403)
    users_rows = service_report_worker_options(order, report_id=report_id)
    if request.method == "POST":
        if is_external_manager():
            abort(403)
        save_token = request.form.get("save_token", "")
        if not claim_report_save_token(save_token, report_id):
            flash("该日报已经保存，请勿重复提交。", "success")
            return redirect(url_for("edit_service_report", report_id=report_id))
        try:
            arrival_time = posted_report_time("arrival_time")
            departure_time = posted_report_time("departure_time")
            total_service_hours = calculated_report_total_service_hours(arrival_time, departure_time)
            total_time = calculated_report_total_time()
            db().execute(
                """
                update service_reports
                set report_date = ?, actual_work_date = ?, total_service_hours = ?, travel_hours = ?, public_transport_hours = ?,
                    driving_miles = ?, departure_address = ?, site_address = ?, total_time = ?, cabinet_number = ?,
                    arrival_time = ?, departure_time = ?, service_description = ?, updated_at = ?
                where id = ?
                """,
                (
                    request.form.get("report_date"),
                    request.form.get("actual_work_date") or request.form.get("report_date"),
                    total_service_hours,
                    to_float(request.form.get("travel_hours")),
                    to_float(request.form.get("public_transport_hours")),
                    to_float(request.form.get("driving_miles")),
                    request.form.get("departure_address", "").strip(),
                    request.form.get("site_address", "").strip(),
                    total_time,
                    request.form.get("cabinet_number", "").strip(),
                    arrival_time,
                    departure_time,
                    request.form.get("service_description", "").strip(),
                    now(),
                    report_id,
                ),
            )
            save_report_detail_rows(report_id)
            relocate_report_attachments(report_id)
            save_report_uploads(report_id)
            log_action(
                "update",
                "service_report",
                report_id,
                f"{order['order_number']} / {request.form.get('report_date')}",
                "修改工作日报",
            )
            db().commit()
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("edit_service_report", report_id=report_id))
        flash("工作日报已保存。", "success")
        return redirect(url_for("edit_service_report", report_id=report_id))
    worker_rows = service_report_workers(report_id)
    selected_workers = [row["id"] for row in worker_rows]
    selected_worker_miles = {row["id"]: row["driving_miles"] for row in worker_rows}
    saved_parts = list(report_parts("service_report_saved_parts", report_id)) or [{} for _ in range(4)]
    replaced_parts = list(report_parts("service_report_replaced_parts", report_id)) or [{} for _ in range(4)]
    return render_template(
        "service_report_form.html",
        order=order,
        report=report_form_defaults(report=report),
        users=users_rows,
        selected_workers=selected_workers,
        selected_worker_miles=selected_worker_miles,
        saved_parts=saved_parts,
        replaced_parts=replaced_parts,
        attachments=get_report_attachments(report_id),
        save_token=secrets.token_urlsafe(24),
        is_edit=True,
        can_edit_report=not is_external_manager(),
    )


@app.post("/service-reports/<int:report_id>/delete")
@login_required
def delete_service_report(report_id):
    report, order = require_service_report(report_id)
    if report["created_by"] != g.user["id"] and not is_manager():
        abort(403)
    attachments = db().execute(
        "select * from service_report_attachments where report_id = ?",
        (report_id,),
    ).fetchall()
    for attachment in attachments:
        attachment_path = report_attachment_path(attachment)
        try:
            os.remove(attachment_path)
            prune_empty_report_folders(attachment_path)
        except FileNotFoundError:
            pass
    shutil.rmtree(os.path.join(REPORT_ATTACHMENTS_DIR, str(report_id)), ignore_errors=True)
    db().execute("delete from service_report_attachments where report_id = ?", (report_id,))
    db().execute("delete from service_report_workers where report_id = ?", (report_id,))
    db().execute("delete from service_report_saved_parts where report_id = ?", (report_id,))
    db().execute("delete from service_report_replaced_parts where report_id = ?", (report_id,))
    db().execute("delete from service_reports where id = ?", (report_id,))
    log_action(
        "delete",
        "service_report",
        report_id,
        f"{order['order_number']} / {report['report_date']}",
        "删除工作日报",
    )
    db().commit()
    flash("工作日报已删除。", "success")
    return redirect(url_for("service_order_detail", order_id=order["id"]))


@app.route("/shared-photos/browse")
@login_required
def browse_shared_photos():
    if not is_internal_user():
        abort(403)
    if not shared_photos_root().is_dir():
        return jsonify(
            {
                "available": False,
                "current": "",
                "parent": None,
                "folders": [],
                "images": [],
                "status": {"waiting": 0, "processing": 0, "completed": 0, "failed": 0},
            }
        )
    requested_path = request.args.get("path", "")
    order_dir = resolve_shared_photo(requested_path, allow_missing=True)
    if not order_dir.is_dir():
        return jsonify(
            {
                "available": True,
                "folder_exists": False,
                "current": requested_path,
                "parent": None,
                "folders": [],
                "images": [],
                "status": {"waiting": 0, "processing": 0, "completed": 0, "failed": 0},
            }
        )
    current = order_dir / "pictures"
    images = []
    if current.is_dir():
        try:
            entries = sorted(
                (
                    entry
                    for entry in current.rglob("*")
                    if entry.is_file()
                    and not entry.name.startswith(".")
                    and entry.suffix.lower().lstrip(".") in ALLOWED_IMAGE_EXTENSIONS
                    and "@eadir" not in {part.casefold() for part in entry.relative_to(current).parts}
                ),
                key=lambda item: item.relative_to(current).as_posix().casefold(),
            )
        except OSError:
            abort(403)
        for entry in entries:
            try:
                entry.relative_to(current)
            except ValueError:
                continue
            if not valid_image_file(entry):
                continue
            relative = shared_photo_relative(entry)
            display_name = entry.relative_to(current).as_posix()
            images.append(
                {
                    "name": display_name,
                    "path": relative,
                    "thumbnail": url_for("shared_photo_thumbnail", path=relative),
                    "preview": url_for("shared_photo_preview", path=relative),
                }
            )
    return jsonify(
        {
            "available": True,
            "folder_exists": True,
            "current": requested_path,
            "parent": None,
            "folders": [],
            "images": images,
            "status": order_photo_status(order_dir),
        }
    )


@app.post("/shared-photos/download")
@login_required
def download_shared_photos():
    if not is_internal_user():
        abort(403)
    selected_paths = list(dict.fromkeys(path for path in request.form.getlist("path") if path))
    if not selected_paths:
        abort(400)

    files = []
    for relative_path in selected_paths:
        source_path, relative_parts = resolve_shared_picture(relative_path)
        files.append((source_path, Path(*relative_parts[2:]).as_posix()))
    if not files:
        abort(400)

    archive_root = secure_filename(Path(selected_paths[0]).parts[0]) or "nas-photos"
    used_names = set()
    archive_file = tempfile.NamedTemporaryFile(prefix=f"{archive_root}-", suffix=".zip", delete=False)
    archive_path = archive_file.name
    archive_file.close()
    with zipfile.ZipFile(archive_path, "w", compression=zipfile.ZIP_STORED) as archive:
        for source_path, display_name in files:
            arcname = f"{archive_root}/{display_name}"
            if arcname in used_names:
                stem, extension = os.path.splitext(display_name)
                suffix = 2
                while f"{archive_root}/{stem}-{suffix}{extension}" in used_names:
                    suffix += 1
                arcname = f"{archive_root}/{stem}-{suffix}{extension}"
            used_names.add(arcname)
            archive.write(source_path, arcname=arcname)

    response = send_file(
        archive_path,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"{archive_root}-photos.zip",
    )

    def remove_archive():
        try:
            os.remove(archive_path)
        except FileNotFoundError:
            pass

    response.call_on_close(remove_archive)
    return response


@app.route("/shared-photos/preview")
@login_required
def shared_photo_preview():
    if not is_internal_user():
        abort(403)
    source_path, _ = resolve_shared_picture(request.args.get("path", ""))
    return send_file(source_path, as_attachment=False, conditional=True)


@app.route("/shared-photos/thumbnail")
@login_required
def shared_photo_thumbnail():
    if not is_internal_user():
        abort(403)
    source_path = resolve_shared_photo(request.args.get("path", ""), require_file=True)
    relative_parts = source_path.relative_to(shared_photos_root()).parts
    if len(relative_parts) >= 3 and relative_parts[1].casefold() == "pictures":
        thumbnail_path = shared_photos_root() / relative_parts[0] / "thumbnails" / Path(*relative_parts[2:])
        if thumbnail_path.is_file() and valid_image_file(thumbnail_path):
            return send_file(thumbnail_path, mimetype="image/jpeg", max_age=3600)
    buffer = BytesIO()
    try:
        with Image.open(source_path) as source:
            source.seek(0)
            image = ImageOps.exif_transpose(source)
            if image.mode not in {"RGB", "L"}:
                background = Image.new("RGB", image.size, "white")
                if "A" in image.getbands():
                    background.paste(image, mask=image.getchannel("A"))
                else:
                    background.paste(image.convert("RGB"))
                image = background
            else:
                image = image.convert("RGB")
            image.thumbnail((420, 320), Image.Resampling.LANCZOS)
            image.save(buffer, format="JPEG", quality=70, optimize=True)
    except (OSError, ValueError):
        abort(404)
    buffer.seek(0)
    return send_file(buffer, mimetype="image/jpeg", max_age=3600)


@app.route("/service-report-attachments/<int:attachment_id>")
@login_required
def preview_report_attachment(attachment_id):
    attachment = db().execute("select * from service_report_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    require_service_report(attachment["report_id"])
    return send_file(
        report_attachment_path(attachment),
        as_attachment=False,
        download_name=attachment["original_filename"],
        mimetype=attachment["content_type"] or None,
        conditional=True,
    )


@app.post("/service-report-attachments/<int:attachment_id>/delete")
@login_required
def delete_report_attachment(attachment_id):
    attachment = db().execute("select * from service_report_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    report, order = require_service_report(attachment["report_id"])
    try:
        attachment_path = report_attachment_path(attachment)
        os.remove(attachment_path)
        prune_empty_report_folders(attachment_path)
    except FileNotFoundError:
        pass
    db().execute("delete from service_report_attachments where id = ?", (attachment_id,))
    db().commit()
    flash("附件已删除。", "success")
    redirect_anchor = request.form.get("redirect_anchor", "").strip()
    if not re.fullmatch(r"#[A-Za-z0-9_-]+", redirect_anchor):
        redirect_anchor = ""
    return redirect(url_for("edit_service_report", report_id=report["id"]) + redirect_anchor)


@app.post("/service-reports/<int:report_id>/export")
@login_required
def export_service_report(report_id):
    report, order = require_service_report(report_id)
    try:
        document_bytes = build_service_report_docx(report, order)
    except Exception:
        app.logger.exception("Failed to export service report %s", report_id)
        flash("工作日报导出失败，请检查日报内容或照片后重试。", "error")
        return redirect(url_for("edit_service_report", report_id=report_id))
    filename = secure_filename(f"{order['client_order_number']}-{report['report_date']}-report.docx") or "service-report.docx"
    return send_file(
        BytesIO(document_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


def expense_defaults(expense=None):
    if expense:
        return dict(expense)
    return {
        "expense_number": next_expense_number(),
        "project_id": None,
        "project": "",
        "expense_date": date.today().isoformat(),
        "amount": "",
        "currency": "USD",
        "description": "",
        "status": "draft",
    }


def expense_items(expense_id):
    return db().execute(
        """
        select expense_items.*, projects.is_active
        from expense_items
        left join projects on projects.id = expense_items.project_id
        where expense_items.expense_id = ?
        order by expense_items.sort_order, expense_items.id
        """,
        (expense_id,),
    ).fetchall()


def expense_items_from_form(expense_id=None):
    project_ids = request.form.getlist("project_id")
    amounts = request.form.getlist("item_amount")
    descriptions = request.form.getlist("item_description")
    rows = []
    for index, project_id in enumerate(project_ids):
        if not project_id:
            continue
        project = db().execute(
            """
            select * from projects
            where id = ? and project_type = 'expense'
              and (is_active = 1 or id in (
                  select project_id from expense_items where expense_id = ?
              ))
            """,
            (project_id, expense_id or 0),
        ).fetchone()
        if not project:
            raise ValueError("选择的员工报销项目不存在或已停用，请重新选择。")
        amount = to_float(amounts[index] if index < len(amounts) else 0)
        if amount <= 0:
            raise ValueError("每个员工报销项目的金额必须大于 0。")
        if is_lodging_project_name(project["name"]):
            validate_lodging_reimbursement(amount, "员工住宿报销")
        if is_fuel_project_name(project["name"]):
            validate_fuel_reimbursement_allowed(amount, "员工油费")
        description = descriptions[index].strip() if index < len(descriptions) else ""
        rows.append(
            {
                "project": project,
                "amount": amount,
                "description": description,
                "sort_order": len(rows),
            }
        )
    if not rows:
        raise ValueError("请至少添加一个员工报销项目。")
    return rows


def save_expense_items(expense_id, item_rows):
    db().execute("delete from expense_items where expense_id = ?", (expense_id,))
    for item in item_rows:
        project = item["project"]
        db().execute(
            """
            insert into expense_items (
                expense_id, project_id, project, amount, description, sort_order
            ) values (?, ?, ?, ?, ?, ?)
            """,
            (
                expense_id,
                project["id"],
                project["name"],
                item["amount"],
                item["description"],
                item["sort_order"],
            ),
        )


@app.route("/service-orders/<int:order_id>/expenses/new", methods=["GET", "POST"])
@login_required
def new_expense(order_id):
    if not can_create_expense():
        abort(403)
    order = require_service_order(order_id)
    start_date_redirect = require_service_order_start_date(order)
    if start_date_redirect:
        return start_date_redirect
    expense_projects = db().execute(
        "select * from projects where project_type = 'expense' and is_active = 1 order by name"
    ).fetchall()
    if not expense_projects:
        flash("请先由经理或管理员创建员工报销项目。", "error")
        return redirect(url_for("projects") if is_manager() else url_for("service_order_detail", order_id=order_id))
    if request.method == "POST":
        save_token = request.form.get("save_token", "")
        if not claim_expense_save_token(save_token):
            existing = db().execute(
                "select expense_id from expense_save_tokens where token = ?",
                (save_token,),
            ).fetchone()
            if existing and existing["expense_id"]:
                return redirect(url_for("edit_expense", expense_id=existing["expense_id"]))
            flash("该报销正在保存，请稍候。", "error")
            return redirect(url_for("new_expense", order_id=order_id))
        submit_for_review = request.form.get("action") == "submit"
        try:
            item_rows = expense_items_from_form()
            total_amount = sum(item["amount"] for item in item_rows)
            project_names = ", ".join(dict.fromkeys(item["project"]["name"] for item in item_rows))
            first_project = item_rows[0]["project"]
            expense_number = next_expense_number()
            cursor = db().execute(
                """
                insert into expenses (
                    service_order_id, expense_number, project_id, project, expense_date, amount, currency,
                    description, status, created_by, created_at, updated_at
                ) values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    order_id,
                    expense_number,
                    first_project["id"],
                    project_names,
                    request.form.get("expense_date"),
                    total_amount,
                    "USD",
                    request.form.get("description", "").strip(),
                    "submitted" if submit_for_review else "draft",
                    g.user["id"],
                    now(),
                    now(),
                ),
            )
            expense_id = cursor.lastrowid
            finish_expense_save_token(save_token, expense_id)
            save_expense_items(expense_id, item_rows)
            save_expense_uploads(expense_id)
            expense_summary = f"工单：{order['order_number']}；金额：{money(total_amount)}"
            log_action("create", "expense", expense_id, expense_number, expense_summary)
            if submit_for_review:
                log_action("submit", "expense", expense_id, expense_number, expense_summary)
            db().commit()
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("new_expense", order_id=order_id))
        if submit_for_review:
            notify_role(
                ["admin", "manager"],
                "新报销待审核",
                f"{g.user['name']}提交了报销 {expense_number}，工单 {order['order_number']}，金额 {money(total_amount)}。",
                url_for("expense_detail", expense_id=expense_id),
            )
            db().commit()
            flash("报销已提交经理审核。", "success")
            return redirect(url_for("expense_detail", expense_id=expense_id))
        flash("报销已保存。", "success")
        return redirect(url_for("edit_expense", expense_id=expense_id))
    return render_template(
        "expense_form.html",
        order=order,
        expense=expense_defaults(),
        form_items=[],
        expense_projects=expense_projects,
        is_edit=False,
        attachments=[],
        save_token=secrets.token_urlsafe(24),
    )


@app.route("/expenses/<int:expense_id>/edit", methods=["GET", "POST"])
@login_required
def edit_expense(expense_id):
    expense, order = require_expense(expense_id)
    if expense["status"] not in {"draft", "returned"}:
        flash("只有保存未提交或被退回的报销可以编辑。", "error")
        return redirect(url_for("expense_detail", expense_id=expense_id))
    if expense["created_by"] != g.user["id"] and not is_manager():
        abort(403)
    expense_projects = db().execute(
        """
        select * from projects
        where project_type = 'expense' and (
            is_active = 1 or id in (
                select project_id from expense_items where expense_id = ?
            )
        )
        order by is_active desc, name
        """,
        (expense_id,),
    ).fetchall()
    if request.method == "POST":
        save_token = request.form.get("save_token", "")
        if not claim_expense_save_token(save_token, expense_id):
            flash("该报销已经保存，请勿重复提交。", "success")
            return redirect(url_for("edit_expense", expense_id=expense_id))
        submit_for_review = request.form.get("action") == "submit"
        try:
            item_rows = expense_items_from_form(expense_id)
            total_amount = sum(item["amount"] for item in item_rows)
            project_names = ", ".join(dict.fromkeys(item["project"]["name"] for item in item_rows))
            first_project = item_rows[0]["project"]
            db().execute(
                """
                update expenses
                set project_id = ?, project = ?, expense_date = ?, amount = ?, currency = ?, description = ?,
                    status = ?, return_reason = null, updated_at = ?
                where id = ?
                """,
                (
                    first_project["id"],
                    project_names,
                    request.form.get("expense_date"),
                    total_amount,
                    "USD",
                    request.form.get("description", "").strip(),
                    "submitted" if submit_for_review else "draft",
                    now(),
                    expense_id,
                ),
            )
            save_expense_items(expense_id, item_rows)
            save_expense_uploads(expense_id)
            expense_summary = f"工单：{order['order_number']}；金额：{money(total_amount)}"
            log_action("update", "expense", expense_id, expense["expense_number"], expense_summary)
            if submit_for_review:
                log_action("submit", "expense", expense_id, expense["expense_number"], expense_summary)
            db().commit()
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("edit_expense", expense_id=expense_id))
        if submit_for_review:
            notify_role(
                ["admin", "manager"],
                "报销已提交审核",
                f"{g.user['name']}提交了报销 {expense['expense_number']}，工单 {order['order_number']}，金额 {money(total_amount)}。",
                url_for("expense_detail", expense_id=expense_id),
            )
            db().commit()
            flash("报销已提交经理审核。", "success")
            return redirect(url_for("expense_detail", expense_id=expense_id))
        flash("报销已保存。", "success")
        return redirect(url_for("edit_expense", expense_id=expense_id))
    return render_template(
        "expense_form.html",
        order=order,
        expense=expense_defaults(expense),
        form_items=expense_items(expense_id),
        expense_projects=expense_projects,
        is_edit=True,
        attachments=get_expense_attachments(expense_id),
        save_token=secrets.token_urlsafe(24),
    )


@app.route("/expenses/<int:expense_id>")
@login_required
def expense_detail(expense_id):
    expense, order = require_expense(expense_id)
    creator = db().execute("select name, email from users where id = ?", (expense["created_by"],)).fetchone()
    reviewer = db().execute("select name, email from users where id = ?", (expense["reviewed_by"],)).fetchone() if expense["reviewed_by"] else None
    reimburser = db().execute(
        "select name, email from users where id = ?",
        (expense["reimbursed_by"],),
    ).fetchone() if expense["reimbursed_by"] else None
    return render_template(
        "expense_detail.html",
        order=order,
        expense=expense,
        items=expense_items(expense_id),
        creator=creator,
        reviewer=reviewer,
        reimburser=reimburser,
        attachments=get_expense_attachments(expense_id),
        labels=EXPENSE_STATUS_LABELS,
    )


@app.post("/expenses/<int:expense_id>/approve")
@login_required
def approve_expense(expense_id):
    expense, order = require_expense(expense_id)
    if expense["status"] != "submitted":
        flash("只有待经理审核的报销可以通过。", "error")
        return redirect(url_for("expense_detail", expense_id=expense_id))
    db().execute(
        """
        update expenses
        set status = 'approved', return_reason = null, reviewed_by = ?, reviewed_at = ?,
            payout_status = 'pending', reimbursed_by = null, reimbursed_at = null, updated_at = ?
        where id = ?
        """,
        (g.user["id"], now(), now(), expense_id),
    )
    message_link = url_for("expense_detail", expense_id=expense_id)
    message_body = (
        f"{g.user['name']}已审核通过报销 {expense['expense_number']}，"
        f"工单 {order['order_number']}，金额 {money(expense['amount'], expense['currency'])}。"
    )
    create_message(expense["created_by"], "报销已审核通过", message_body, message_link)
    notify_role(
        ["admin"],
        "报销已审核通过",
        message_body,
        message_link,
        exclude_user_ids={expense["created_by"], g.user["id"]},
    )
    log_action("approve", "expense", expense_id, expense["expense_number"], f"工单：{order['order_number']}")
    db().commit()
    flash("报销已审核通过。", "success")
    return redirect(url_for("expense_detail", expense_id=expense_id))


@app.post("/expenses/<int:expense_id>/return")
@login_required
def return_expense(expense_id):
    expense, order = require_expense(expense_id)
    if expense["status"] != "submitted":
        flash("只有待经理审核的报销可以退回。", "error")
        return redirect(url_for("expense_detail", expense_id=expense_id))
    reason = request.form.get("return_reason", "").strip()
    if not reason:
        flash("请填写退回原因。", "error")
        return redirect(url_for("expense_detail", expense_id=expense_id))
    db().execute(
        "update expenses set status = 'returned', return_reason = ?, reviewed_by = ?, reviewed_at = ?, updated_at = ? where id = ?",
        (reason, g.user["id"], now(), now(), expense_id),
    )
    create_message(expense["created_by"], "报销已被退回", f"报销 {expense['expense_number']} 已被退回。原因：{reason}", url_for("expense_detail", expense_id=expense_id))
    log_action("return", "expense", expense_id, expense["expense_number"], f"原因：{reason}")
    db().commit()
    flash("报销已退回。", "success")
    return redirect(url_for("expense_detail", expense_id=expense_id))


@app.post("/expenses/<int:expense_id>/delete")
@login_required
def delete_expense(expense_id):
    expense, order = require_expense(expense_id)
    is_admin = normalized_role() == "admin"
    if not is_admin and expense["created_by"] != g.user["id"] and not is_manager():
        abort(403)
    if not is_admin and expense["status"] not in {"draft", "returned"}:
        flash("只有保存未提交或被退回的报销可以删除。", "error")
        return redirect(url_for("expense_detail", expense_id=expense_id))
    shutil.rmtree(expense_attachment_dir(expense_id), ignore_errors=True)
    db().execute("delete from expense_attachments where expense_id = ?", (expense_id,))
    db().execute("delete from expense_items where expense_id = ?", (expense_id,))
    db().execute("delete from expenses where id = ?", (expense_id,))
    log_action("delete", "expense", expense_id, expense["expense_number"], f"工单：{order['order_number']}")
    db().commit()
    flash("报销已删除。", "success")
    return redirect(url_for("service_order_detail", order_id=order["id"]))


@app.route("/expense-attachments/<int:attachment_id>")
@login_required
def preview_expense_attachment(attachment_id):
    attachment = db().execute("select * from expense_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    require_expense(attachment["expense_id"])
    return send_file(
        expense_attachment_path(attachment),
        as_attachment=False,
        download_name=attachment["original_filename"],
        mimetype=attachment["content_type"] or None,
        conditional=True,
    )


@app.route("/expense-attachments/<int:attachment_id>/download")
@login_required
def download_expense_attachment(attachment_id):
    attachment = db().execute("select * from expense_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    require_expense(attachment["expense_id"])
    return send_file(expense_attachment_path(attachment), as_attachment=True, download_name=attachment["original_filename"])


@app.post("/expense-attachments/<int:attachment_id>/delete")
@login_required
def delete_expense_attachment(attachment_id):
    attachment = db().execute("select * from expense_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    expense, order = require_expense(attachment["expense_id"])
    if expense["status"] not in {"draft", "returned"}:
        abort(403)
    try:
        os.remove(expense_attachment_path(attachment))
    except FileNotFoundError:
        pass
    db().execute("delete from expense_attachments where id = ?", (attachment_id,))
    db().commit()
    flash("附件已删除。", "success")
    return redirect(url_for("edit_expense", expense_id=expense["id"]))


@app.route("/invoices/new", methods=["GET", "POST"])
@login_required
def new_invoice():
    if not can_create_invoice():
        abort(403)
    requested_order_id = request.args.get("service_order_id", "")
    requested_reimbursement_id = request.args.get("customer_reimbursement_id", "")
    source_reimbursement = None
    prefilled_items = []
    if request.method == "GET" and not requested_reimbursement_id.isdigit():
        flash("请从审核通过的工单结算生成发票。", "error")
        if requested_order_id.isdigit():
            return redirect(url_for("customer_reimbursement_form", order_id=int(requested_order_id)))
        return redirect(url_for("service_orders"))
    if request.method == "GET" and requested_order_id.isdigit():
        requested_order = require_service_order(int(requested_order_id))
        start_date_redirect = require_service_order_start_date(requested_order)
        if start_date_redirect:
            return start_date_redirect
    if request.method == "GET" and requested_reimbursement_id.isdigit():
        source_reimbursement, source_order = require_customer_reimbursement(int(requested_reimbursement_id))
        start_date_redirect = require_service_order_start_date(source_order)
        if start_date_redirect:
            return start_date_redirect
        if source_reimbursement["invoice_id"]:
            flash("这份工单结算已经生成过发票。", "success")
            return redirect(url_for("invoice_detail", invoice_id=source_reimbursement["invoice_id"]))
        linked_invoice = service_order_active_invoice(source_order["id"])
        if linked_invoice:
            db().execute(
                "update customer_reimbursements set invoice_id = ? where id = ?",
                (linked_invoice["id"], source_reimbursement["id"]),
            )
            db().commit()
            flash("这个工单已经有关联发票。", "success")
            return redirect(url_for("invoice_detail", invoice_id=linked_invoice["id"]))
        if source_reimbursement["status"] != "approved":
            flash("工单结算审核通过后才能生成发票。", "error")
            return redirect(url_for("customer_reimbursement_form", order_id=source_order["id"]))
        requested_order_id = str(source_order["id"])
        try:
            prefilled_items = customer_reimbursement_invoice_items(source_reimbursement)
        except ValueError as error:
            flash(str(error), "error")
            return redirect(url_for("customer_reimbursement_form", order_id=source_order["id"]))
    if is_external_user():
        if not g.user["client_id"]:
            flash("请联系管理员绑定客户。", "error")
            return redirect(url_for("dashboard"))
        clients_rows = db().execute("select * from clients where id = ?", (g.user["client_id"],)).fetchall()
    else:
        clients_rows = db().execute("select * from clients order by client_number").fetchall()
    if not clients_rows:
        flash("请先创建一个客户。", "error")
        return redirect(url_for("clients"))
    projects_rows = db().execute(
        "select * from projects where project_type = 'invoice' and is_active = 1 order by name"
    ).fetchall()
    service_orders_rows = db().execute("select * from service_orders where status != 'closed' order by created_at desc, id desc").fetchall()
    if not projects_rows:
        flash("请先由经理或管理员维护项目。", "error")
        return redirect(url_for("projects") if is_manager() else url_for("dashboard"))
    if request.method == "POST":
        save_token = request.form.get("save_token", "")
        if not claim_invoice_save_token(save_token):
            existing = db().execute(
                "select invoice_id from invoice_save_tokens where token = ?",
                (save_token,),
            ).fetchone()
            if existing and existing["invoice_id"]:
                return redirect(url_for("edit_invoice", invoice_id=existing["invoice_id"]))
            flash("该发票正在保存，请稍候。", "error")
            return redirect(url_for("new_invoice"))
        uploads = uploaded_attachments_from_request()
        if not validate_attachment_uploads(uploads):
            db().rollback()
            return redirect(url_for("new_invoice"))
        try:
            invoice_id = create_invoice_from_form(save_token=save_token)
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("new_invoice"))
        except sqlite3.IntegrityError:
            db().rollback()
            flash("保存发票时发生数据库冲突，请重新提交。", "error")
            return redirect(url_for("new_invoice"))
        flash("发票已生成。", "success")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    today = date.today()
    defaults = {
        "invoice_number": next_invoice_number(),
        "issue_date": today.isoformat(),
        "due_date": (today + timedelta(days=30)).isoformat(),
        "service_order_id": int(requested_order_id) if requested_order_id.isdigit() else None,
        "customer_reimbursement_id": source_reimbursement["id"] if source_reimbursement else None,
        "client_id": source_order["client_id"] if source_reimbursement else None,
    }
    return render_template(
        "invoice_form.html",
        clients=clients_rows,
        projects=projects_rows,
        service_orders=service_orders_rows,
        defaults=defaults,
        form_title="根据工单结算生成发票" if source_reimbursement else "新建发票",
        form_items=prefilled_items,
        is_edit=False,
        attachments=[],
        save_token=secrets.token_urlsafe(24),
    )


def invoice_items_from_form():
    rows = []
    for project_id, amount in zip(request.form.getlist("project_id"), request.form.getlist("amount")):
        if not project_id:
            continue
        project = db().execute(
            "select * from projects where id = ? and project_type = 'invoice' and is_active = 1",
            (project_id,),
        ).fetchone()
        if not project:
            raise ValueError("选择的项目不存在或已停用，请重新选择项目。")
        rows.append((project, to_float(amount)))
    if not rows:
        raise ValueError("请至少选择一个项目明细。")
    if sum(amount for _, amount in rows) <= 0:
        raise ValueError("发票金额必须大于 0。")
    return rows


def posted_client_id():
    try:
        return int(request.form.get("client_id"))
    except (TypeError, ValueError):
        raise ValueError("请选择客户。")


def posted_service_order_id(require_start_date=False, required=False):
    value = request.form.get("service_order_id")
    if not value:
        if required:
            raise ValueError("请选择关联工单。")
        return None
    try:
        order_id = int(value)
    except (TypeError, ValueError):
        raise ValueError("请选择有效工单。")
    order = require_service_order(order_id)
    if require_start_date and not order["start_date"]:
        raise ValueError("所选工单还没有开始日期，请先编辑工单并维护开始日期。")
    return order_id


def create_invoice_from_form(save_token=None):
    client_id = posted_client_id()
    if not can_access_client(client_id):
        abort(403)
    service_order_id = posted_service_order_id(require_start_date=True, required=True)
    source_reimbursement = None
    reimbursement_id = request.form.get("customer_reimbursement_id", "")
    if reimbursement_id:
        if not reimbursement_id.isdigit():
            raise ValueError("工单结算来源无效。")
        source_reimbursement = db().execute(
            "select * from customer_reimbursements where id = ?",
            (int(reimbursement_id),),
        ).fetchone()
        if not source_reimbursement or source_reimbursement["service_order_id"] != service_order_id:
            raise ValueError("工单结算与所选工单不匹配。")
        if source_reimbursement["status"] != "approved":
            raise ValueError("工单结算审核通过后才能生成发票。")
        if source_reimbursement["invoice_id"]:
            raise ValueError("这份工单结算已经生成过发票。")
        if service_order_active_invoice(service_order_id):
            raise ValueError("这个工单已经有关联发票，不能重复生成。")
        source_order = db().execute(
            "select client_id from service_orders where id = ?",
            (service_order_id,),
        ).fetchone()
        if source_order and source_order["client_id"] and client_id != source_order["client_id"]:
            raise ValueError("发票客户必须与工单关联客户一致。")
    invoice_number = request.form.get("invoice_number", "").strip()
    if not invoice_number:
        invoice_number = next_invoice_number()
    if invoice_number_exists(invoice_number):
        invoice_number = next_invoice_number()
    item_rows = invoice_items_from_form()
    status = "completed"
    cursor = None
    for _ in range(30):
        try:
            cursor = db().execute(
                """
                insert into invoices (
                    invoice_number, client_id, service_order_id, issue_date, due_date, currency, notes, status, created_by, created_at
                ) values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    invoice_number,
                    client_id,
                    service_order_id,
                    request.form.get("issue_date"),
                    request.form.get("due_date"),
                    request.form.get("currency", "USD"),
                    request.form.get("notes", "").strip(),
                    status,
                    g.user["id"],
                    now(),
                ),
            )
            break
        except sqlite3.IntegrityError:
            invoice_number = next_invoice_number()
    if cursor is None:
        raise RuntimeError("Unable to generate a unique invoice number.")
    invoice_id = cursor.lastrowid
    if save_token:
        finish_invoice_save_token(save_token, invoice_id)
    for project, amount in item_rows:
        db().execute(
            """
            insert into invoice_items (invoice_id, project_id, description, amount, tax_rate)
            values (?, ?, ?, ?, ?)
            """,
            (invoice_id, project["id"], project["name"], to_float(amount), project["tax_rate"]),
        )
    for uploaded in uploaded_attachments_from_request():
        if uploaded and uploaded.filename:
            save_uploaded_attachment(invoice_id, uploaded)
    if source_reimbursement:
        copy_mileage_proofs_to_invoice(invoice_id, service_order_id)
    if source_reimbursement:
        db().execute(
            "update customer_reimbursements set invoice_id = ? where id = ? and invoice_id is null",
            (invoice_id, source_reimbursement["id"]),
        )
    invoice_summary = f"金额：{money(invoice_totals(invoice_id)['total'], request.form.get('currency', 'USD'))}"
    log_action("create", "invoice", invoice_id, invoice_number, invoice_summary)
    db().commit()
    return invoice_id


@app.route("/invoices/<int:invoice_id>/edit", methods=["GET", "POST"])
@login_required
def edit_invoice(invoice_id):
    invoice, client, items = load_invoice(invoice_id)
    if invoice["status"] != "draft":
        flash("只有管理员调整为草稿的发票可以编辑。", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    if invoice["created_by"] != g.user["id"] and not is_manager():
        abort(403)
    if is_external_user():
        clients_rows = db().execute("select * from clients where id = ?", (g.user["client_id"],)).fetchall()
    else:
        clients_rows = db().execute("select * from clients order by client_number").fetchall()
    projects_rows = db().execute(
        """
        select * from projects
        where project_type = 'invoice' and (is_active = 1 or id in (
            select project_id from invoice_items where invoice_id = ?
        ))
        order by is_active desc, name
        """,
        (invoice_id,),
    ).fetchall()
    service_orders_rows = db().execute("select * from service_orders where status != 'closed' or id = ? order by created_at desc, id desc", (invoice["service_order_id"] or 0,)).fetchall()
    if request.method == "POST":
        save_token = request.form.get("save_token", "")
        if not claim_invoice_save_token(save_token, invoice_id):
            flash("该发票已经保存，请勿重复提交。", "success")
            return redirect(url_for("edit_invoice", invoice_id=invoice_id))
        uploads = uploaded_attachments_from_request()
        if not validate_attachment_uploads(uploads, invoice_id):
            db().rollback()
            return redirect(url_for("edit_invoice", invoice_id=invoice_id))
        try:
            update_invoice_from_form(invoice_id)
            log_action("update", "invoice", invoice_id, invoice["invoice_number"], "修改发票内容")
        except ValueError as error:
            db().rollback()
            flash(str(error), "error")
            return redirect(url_for("edit_invoice", invoice_id=invoice_id))
        except sqlite3.IntegrityError:
            db().rollback()
            flash("发票编号已经存在，请更换一个发票编号后再保存。", "error")
            return redirect(url_for("edit_invoice", invoice_id=invoice_id))
        db().commit()
        flash("发票已保存。", "success")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    defaults = {
        "id": invoice["id"],
        "invoice_number": invoice["invoice_number"],
        "client_id": invoice["client_id"],
        "created_by": invoice["created_by"],
        "status": invoice["status"],
        "issue_date": invoice["issue_date"],
        "due_date": invoice["due_date"],
        "currency": invoice["currency"],
        "notes": invoice["notes"],
        "service_order_id": invoice["service_order_id"],
    }
    return render_template(
        "invoice_form.html",
        clients=clients_rows,
        projects=projects_rows,
        service_orders=service_orders_rows,
        defaults=defaults,
        form_title="编辑发票",
        form_items=items,
        is_edit=True,
        attachments=get_invoice_attachments(invoice_id),
        save_token=secrets.token_urlsafe(24),
    )


def update_invoice_from_form(invoice_id):
    client_id = posted_client_id()
    if not can_access_client(client_id):
        abort(403)
    service_order_id = posted_service_order_id()
    item_rows = invoice_items_from_form()
    db().execute(
        """
        update invoices
        set client_id = ?, service_order_id = ?, issue_date = ?, due_date = ?, currency = ?,
            notes = ?
        where id = ?
        """,
        (
            client_id,
            service_order_id,
            request.form.get("issue_date"),
            request.form.get("due_date"),
            request.form.get("currency", "USD"),
            request.form.get("notes", "").strip(),
            invoice_id,
        ),
    )
    db().execute("delete from invoice_items where invoice_id = ?", (invoice_id,))
    for project, amount in item_rows:
        db().execute(
            """
            insert into invoice_items (invoice_id, project_id, description, amount, tax_rate)
            values (?, ?, ?, ?, ?)
            """,
            (invoice_id, project["id"], project["name"], amount, project["tax_rate"]),
        )
    for uploaded in uploaded_attachments_from_request():
        if uploaded and uploaded.filename:
            save_uploaded_attachment(invoice_id, uploaded)
    db().execute("update invoices set status = 'completed', return_reason = null where id = ?", (invoice_id,))


@app.route("/invoices/<int:invoice_id>")
@login_required
def invoice_detail(invoice_id):
    invoice, client, items = load_invoice(invoice_id)
    if invoice["status"] == "draft" and invoice["created_by"] == g.user["id"]:
        return redirect(url_for("edit_invoice", invoice_id=invoice_id))
    creator = db().execute("select name, email from users where id = ?", (invoice["created_by"],)).fetchone()
    service_order = None
    if invoice["service_order_id"]:
        service_order = db().execute(
            """
            select service_orders.*, contracts.contract_number, contracts.title as contract_title
            from service_orders
            left join contracts on contracts.id = service_orders.contract_id
            where service_orders.id = ?
            """,
            (invoice["service_order_id"],),
        ).fetchone()
    return render_template(
        "invoice_detail.html",
        invoice=invoice,
        client=client,
        items=items,
        totals=invoice_totals(invoice_id),
        creator=creator,
        service_order=service_order,
        attachments=get_invoice_attachments(invoice_id),
        company=get_company_profile(),
        terms=get_invoice_terms(),
        payment=get_payment_instructions(),
        labels=STATUS_LABELS,
        today=date.today().isoformat(),
    )


@app.post("/invoices/<int:invoice_id>/admin-status")
@login_required
def admin_update_invoice_status(invoice_id):
    invoice = require_invoice_access(invoice_id)
    if g.user["role"] != "admin":
        abort(403)
    if invoice["paid_at"]:
        flash("这张发票已经核销，不能再修改为保存未提交状态。", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    target_status = request.form.get("status", "")
    if target_status != "draft":
        flash("当前只允许管理员将未核销发票改为保存未提交状态。", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    db().execute(
        "update invoices set status = 'draft', return_reason = null where id = ?",
        (invoice_id,),
    )
    create_message(
        invoice["created_by"],
        "发票状态已调整",
        f"管理员已将发票 {invoice['invoice_number']} 调整为保存未提交状态。",
        url_for("edit_invoice", invoice_id=invoice_id),
    )
    log_action("status_change", "invoice", invoice_id, invoice["invoice_number"], "调整为保存未提交")
    db().commit()
    flash("发票状态已修改为保存未提交。", "success")
    return redirect(url_for("invoice_detail", invoice_id=invoice_id))


@app.post("/invoices/<int:invoice_id>/void")
@login_required
def void_invoice(invoice_id):
    invoice = require_invoice_access(invoice_id)
    if invoice["status"] == "completed":
        flash("审核完成后的发票不能作废。", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    db().execute("update invoices set status = 'void' where id = ?", (invoice_id,))
    log_action("void", "invoice", invoice_id, invoice["invoice_number"], "作废发票")
    db().commit()
    flash("发票已作废。", "success")
    return redirect(url_for("invoice_detail", invoice_id=invoice_id))


@app.post("/invoices/<int:invoice_id>/delete")
@login_required
def delete_invoice(invoice_id):
    invoice = require_invoice_access(invoice_id)
    has_delete_permission = has_action_permission("invoices", "delete")
    can_delete_returned = invoice["status"] in {"draft", "returned"} and (
        has_delete_permission or invoice["created_by"] == g.user["id"]
    )
    can_delete_void = invoice["status"] == "void" and has_delete_permission
    if not (has_delete_permission or can_delete_returned or can_delete_void):
        flash("只有有删除权限，或草稿/退回状态下由发起人删除的发票可以删除。", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    shutil.rmtree(invoice_attachment_path(invoice_id), ignore_errors=True)
    db().execute(
        "update messages set link = null where link = ? or link = ?",
        (url_for("invoice_detail", invoice_id=invoice_id), url_for("edit_invoice", invoice_id=invoice_id)),
    )
    db().execute("update customer_reimbursements set invoice_id = null where invoice_id = ?", (invoice_id,))
    db().execute("delete from invoices where id = ?", (invoice_id,))
    log_action("delete", "invoice", invoice_id, invoice["invoice_number"], f"删除状态为 {invoice['status']} 的发票")
    db().commit()
    flash("发票已删除。", "success")
    return redirect(url_for("invoices"))


@app.post("/invoices/<int:invoice_id>/mark-paid")
@login_required
def mark_invoice_paid(invoice_id):
    invoice = require_invoice_access(invoice_id)
    if not is_internal_user():
        abort(403)
    if invoice["status"] != "completed":
        flash("只有已完成的发票才能核销。", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    totals = invoice_totals(invoice_id)
    db().execute(
        """
        update invoices set paid_at = ?, payment_amount = ?, payment_note = ?
        where id = ?
        """,
        (
            request.form.get("paid_at") or date.today().isoformat(),
            to_float(request.form.get("payment_amount"), totals["total"]),
            request.form.get("payment_note", "").strip(),
            invoice_id,
        ),
    )
    log_action("mark_paid", "invoice", invoice_id, invoice["invoice_number"], "记录发票核销")
    db().commit()
    flash("发票已核销。", "success")
    return redirect(url_for("invoice_detail", invoice_id=invoice_id))


@app.post("/invoices/<int:invoice_id>/unmark-paid")
@login_required
def unmark_invoice_paid(invoice_id):
    invoice = require_invoice_access(invoice_id)
    db().execute("update invoices set paid_at = null, payment_amount = null, payment_note = null where id = ?", (invoice_id,))
    log_action("unmark_paid", "invoice", invoice_id, invoice["invoice_number"], "取消发票核销")
    db().commit()
    flash("发票核销记录已取消。", "success")
    return redirect(url_for("invoice_detail", invoice_id=invoice_id))


@app.post("/invoices/<int:invoice_id>/send")
@login_required
def send_invoice(invoice_id):
    invoice = require_invoice_access(invoice_id)
    if is_external_user():
        abort(403)
    if invoice["status"] != "completed":
        flash("只有经理审核完成后的发票才能发送邮件。", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    try:
        send_invoice_email(invoice_id)
    except Exception as error:
        app.logger.exception("Unable to send invoice %s", invoice_id)
        flash(f"发票邮件发送失败：{error}", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    flash("发票邮件已发送。", "success")
    return redirect(url_for("invoice_detail", invoice_id=invoice_id))


@app.post("/invoices/<int:invoice_id>/attachments")
@login_required
def upload_attachment(invoice_id):
    require_invoice_access(invoice_id)
    uploads = uploaded_attachments_from_request()
    if not uploads:
        flash("请选择要上传的附件。", "error")
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    if not validate_attachment_uploads(uploads, invoice_id):
        return redirect(url_for("invoice_detail", invoice_id=invoice_id))
    for uploaded in uploads:
        save_uploaded_attachment(invoice_id, uploaded)
    db().commit()
    flash(f"已上传 {len(uploads)} 个附件。", "success")
    return redirect(url_for("invoice_detail", invoice_id=invoice_id))


@app.route("/attachments/<int:attachment_id>")
@login_required
def download_attachment(attachment_id):
    attachment = db().execute("select * from invoice_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    require_invoice_access(attachment["invoice_id"])
    return send_file(attachment_file_path(attachment), as_attachment=True, download_name=attachment["original_filename"])


@app.route("/attachments/<int:attachment_id>/preview")
@login_required
def preview_attachment(attachment_id):
    attachment = db().execute("select * from invoice_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    require_invoice_access(attachment["invoice_id"])
    return send_file(
        attachment_file_path(attachment),
        as_attachment=False,
        download_name=attachment["original_filename"],
        mimetype=attachment["content_type"] or None,
        conditional=True,
    )


@app.post("/attachments/<int:attachment_id>/delete")
@login_required
def delete_attachment(attachment_id):
    attachment = db().execute("select * from invoice_attachments where id = ?", (attachment_id,)).fetchone()
    if not attachment:
        abort(404)
    require_invoice_access(attachment["invoice_id"])
    try:
        os.remove(attachment_file_path(attachment))
    except FileNotFoundError:
        pass
    db().execute("delete from invoice_attachments where id = ?", (attachment_id,))
    db().commit()
    flash("附件已删除。", "success")
    return redirect(url_for("invoice_detail", invoice_id=attachment["invoice_id"]))


@app.post("/invoices/<int:invoice_id>/export")
@login_required
def export_invoice(invoice_id):
    invoice, client, items = load_invoice(invoice_id)
    archive_name, archive_bytes = build_invoice_zip(invoice, client, items)
    return send_file(
        BytesIO(archive_bytes),
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"{archive_name}.zip",
    )


@app.get("/invoices/<int:invoice_id>/export-pdf")
@login_required
def export_invoice_pdf(invoice_id):
    invoice, client, items = load_invoice(invoice_id)
    invoice_name = secure_filename(invoice["invoice_number"]) or f"invoice-{invoice['id']}"
    return send_file(
        BytesIO(render_invoice_pdf(invoice, client, items)),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{invoice_name}.pdf",
    )


def build_invoice_zip(invoice, client, items):
    archive_name = secure_filename(invoice["invoice_number"])
    if not archive_name:
        archive_name = f"invoice-{invoice['id']}"
    pdf_bytes = render_invoice_pdf(invoice, client, items)
    attachments = get_invoice_attachments(invoice["id"])
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(f"{archive_name}/{archive_name}.pdf", pdf_bytes)
        for attachment in attachments:
            archive.write(attachment_file_path(attachment), arcname=f"{archive_name}/{attachment['original_filename']}")
    return archive_name, buffer.getvalue()


def customer_reimbursement_email_attachments(invoice):
    if not invoice["service_order_id"]:
        return []
    reimbursement = latest_customer_reimbursement(invoice["service_order_id"])
    if not reimbursement:
        return []
    order = require_service_order(invoice["service_order_id"])
    _, attachments = customer_reimbursement_outgoing_attachments(reimbursement, order)
    return attachments


def unique_email_attachment_filename(filename, used_names):
    filename = os.path.basename(filename or "").strip() or "attachment"
    stem, extension = os.path.splitext(filename)
    candidate = filename
    counter = 2
    while candidate.casefold() in used_names:
        candidate = f"{stem or 'attachment'}-{counter}{extension}"
        counter += 1
    used_names.add(candidate.casefold())
    return candidate


def invoice_email_attachments(invoice, client, items):
    invoice_name = secure_filename(invoice["invoice_number"]) or f"invoice-{invoice['id']}"
    attachments = [
        {
            "filename": f"{invoice_name}.pdf",
            "content": render_invoice_pdf(invoice, client, items),
            "maintype": "application",
            "subtype": "pdf",
        }
    ]
    for attachment in get_invoice_attachments(invoice["id"]):
        path = attachment_file_path(attachment)
        if not os.path.exists(path):
            continue
        maintype, _, subtype = (attachment["content_type"] or "application/octet-stream").partition("/")
        with open(path, "rb") as file:
            attachments.append(
                {
                    "filename": attachment["original_filename"],
                    "content": file.read(),
                    "maintype": maintype or "application",
                    "subtype": subtype or "octet-stream",
                }
            )
    attachments.extend(customer_reimbursement_email_attachments(invoice))

    used_names = set()
    for attachment in attachments:
        attachment["filename"] = unique_email_attachment_filename(attachment["filename"], used_names)
    return attachments


def render_invoice_export_html(invoice, client, items):
    return render_template(
        "invoice_export.html",
        invoice=invoice,
        client=client,
        items=items,
        totals=invoice_totals(invoice["id"]),
        company=get_company_profile(),
        terms=get_invoice_terms(),
        payment=get_payment_instructions(),
        labels=STATUS_LABELS,
    )


def render_invoice_pdf(invoice, client, items):
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    buffer = BytesIO()
    page = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    left = 42
    top = height - 42
    line = 15
    company = get_company_profile()
    payment = get_payment_instructions()
    terms = get_invoice_terms()
    totals = invoice_totals(invoice["id"])

    def ensure_space(required):
        nonlocal y
        if y - required < 42:
            page.showPage()
            y = height - 42

    def wrapped_lines(text, max_width, font_name="STSong-Light", font_size=8):
        lines = []
        for source_line in pdf_text(text).splitlines() or [""]:
            words = source_line.split(" ")
            current = ""
            for word in words:
                candidate = word if not current else f"{current} {word}"
                if pdfmetrics.stringWidth(candidate, font_name, font_size) <= max_width:
                    current = candidate
                    continue
                if current:
                    lines.append(current)
                if pdfmetrics.stringWidth(word, font_name, font_size) <= max_width:
                    current = word
                    continue
                current = ""
                chunk = ""
                for char in word:
                    candidate = chunk + char
                    if pdfmetrics.stringWidth(candidate, font_name, font_size) <= max_width:
                        chunk = candidate
                    else:
                        if chunk:
                            lines.append(chunk)
                        chunk = char
                current = chunk
            lines.append(current)
        return [line for line in lines if line]

    def draw_section(title, body):
        nonlocal y
        if not pdf_text(body).strip():
            return
        body_lines = wrapped_lines(body, width - left * 2, font_size=8)
        ensure_space(16 + len(body_lines) * 11)
        page.setFont("STSong-Light", 9)
        page.drawString(left, y, title)
        y -= 12
        page.setFont("STSong-Light", 8)
        page.setFillColor(colors.HexColor("#344054"))
        for line_text in body_lines:
            ensure_space(11)
            page.drawString(left, y, line_text)
            y -= 11
        page.setFillColor(colors.black)
        y -= 6

    page.setFont("STSong-Light", 18)
    page.drawString(left, top, pdf_text(company["name"]))
    page.setFont("STSong-Light", 9)
    y = top - 22
    for part in pdf_text(company["address"]).splitlines():
        page.drawString(left, y, part)
        y -= line

    page.setFont("STSong-Light", 26)
    page.drawRightString(width - left, top, "INVOICE")
    page.setFont("STSong-Light", 10)
    meta_y = top - 40
    for label, value in (
        ("Invoice No.", invoice["invoice_number"]),
        ("Issue Date", invoice["issue_date"]),
        ("Due Date", invoice["due_date"]),
        ("Currency", invoice["currency"]),
    ):
        page.drawRightString(width - left - 95, meta_y, label)
        page.drawRightString(width - left, meta_y, pdf_text(value))
        meta_y -= line

    y -= 18
    page.setFont("STSong-Light", 11)
    page.drawString(left, y, "Bill To")
    y -= line
    page.setFont("STSong-Light", 10)
    for value in (client["name"], client["address"], client["country"]):
        for part in pdf_text(value).splitlines():
            if part:
                page.drawString(left, y, part)
                y -= line

    y -= 10
    page.setStrokeColor(colors.HexColor("#d9dee7"))
    page.line(left, y, width - left, y)
    y -= 18
    page.setFont("STSong-Light", 9)
    headers = [("Description", left), ("Amount", 285), ("Tax Rate", 360), ("Tax", 430), ("Line Total", 500)]
    for text, x in headers:
        page.drawString(x, y, text)
    y -= 8
    page.line(left, y, width - left, y)
    y -= 16
    for item in items:
        amount = float(item["amount"] or 0)
        tax = amount * float(item["tax_rate"] or 0) / 100
        page.drawString(left, y, pdf_text(item["description"])[:46])
        page.drawRightString(340, y, money(amount, invoice["currency"]))
        page.drawRightString(410, y, f"{float(item['tax_rate']):.2f}%")
        page.drawRightString(480, y, money(tax, invoice["currency"]))
        page.drawRightString(width - left, y, money(amount + tax, invoice["currency"]))
        y -= line

    y -= 12
    page.line(360, y, width - left, y)
    y -= 16
    for label, value in (("Subtotal", totals["subtotal"]), ("Tax", totals["tax"]), ("Amount Due", totals["total"])):
        page.drawString(365, y, label)
        page.drawRightString(width - left, y, money(value, invoice["currency"]))
        y -= line

    y -= 10
    draw_section("Notes", invoice["notes"])
    draw_section("Terms", terms)
    draw_section(
        "Payment Instructions",
        "\n".join(
            (
                f"Method: {payment['method']}",
                f"Beneficiary Name: {payment['beneficiary']}",
                f"Bank Name: {payment['bank_name']}",
                f"Account Number: {payment['account_number']}",
                f"Routing Number: {payment['routing_number']}",
                f"SWIFT/BIC: {payment['swift_bic']}",
            )
        ),
    )
    draw_section("Tax Note", company["tax_note"])
    page.showPage()
    page.save()
    buffer.seek(0)
    return buffer.getvalue()


def build_service_report_docx(report, order):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.right_margin = Inches(0.65)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(9)
    normal_style.paragraph_format.space_after = Pt(2)

    def docx_text(value):
        text = "" if value is None else str(value)
        return "".join(
            char
            for char in text
            if char in "\t\n\r" or ord(char) >= 0x20
        )

    def format_run(run, size=9, bold=False, color=None):
        run.font.name = "Microsoft YaHei"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    def set_cell_shading(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)

    def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=8.5):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run(docx_text(text))
        format_run(run, size=size, bold=bold)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def style_table(table, header_rows=0, column_widths=None):
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        if column_widths:
            for row in table.rows:
                for index, width in enumerate(column_widths):
                    row.cells[index].width = Inches(width)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index < header_rows:
                    set_cell_shading(cell, "D9EAD3")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    def add_section_title(title):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(title)
        format_run(run, size=10, bold=True, color=(15, 118, 110))

    def docx_photo_stream(path):
        buffer = BytesIO()
        with Image.open(path) as source:
            source.seek(0)
            image = ImageOps.exif_transpose(source)
            if image.mode not in {"RGB", "L"}:
                background = Image.new("RGB", image.size, "white")
                if "A" in image.getbands():
                    background.paste(image, mask=image.getchannel("A"))
                else:
                    background.paste(image.convert("RGB"))
                image = background
            else:
                image = image.convert("RGB")
            image.thumbnail((960, 960), Image.Resampling.LANCZOS)
            image.save(buffer, format="JPEG", quality=68, optimize=True)
        buffer.seek(0)
        return buffer

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    format_run(title.add_run("现场服务日报"), size=18, bold=True)

    for label, value in (
        ("站点名称", order["client_name"]),
        ("业主", order["buyer_owner"]),
        ("服务现场地址", order["site_address"]),
        ("服务订单号码", order["client_order_number"]),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        format_run(paragraph.add_run(f"{label}："), bold=True)
        format_run(paragraph.add_run(docx_text(value)))

    workers = service_report_workers(report["id"])
    add_section_title("现场服务人员")
    worker_rows = max(len(workers), 2)
    worker_table = document.add_table(rows=worker_rows + 1, cols=2)
    set_cell_text(worker_table.rows[0].cells[0], "姓名", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(worker_table.rows[0].cells[1], "公司", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for index in range(worker_rows):
        worker = workers[index] if index < len(workers) else None
        set_cell_text(worker_table.rows[index + 1].cells[0], worker["name"] if worker else "")
        set_cell_text(worker_table.rows[index + 1].cells[1], "Prasinos Power LLC" if worker else "")
    style_table(worker_table, header_rows=1, column_widths=[3.5, 3.5])

    summary = document.add_table(rows=2, cols=3)
    labels = ["报告日期", "总计服务工时（小时）", "交通时长（小时）"]
    values = [report["report_date"], report["total_service_hours"], report["travel_hours"]]
    for index, label in enumerate(labels):
        set_cell_text(summary.rows[0].cells[index], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(summary.rows[1].cells[index], values[index], align=WD_ALIGN_PARAGRAPH.CENTER)
    style_table(summary, header_rows=1, column_widths=[2.35, 2.35, 2.3])

    add_section_title("现场服务描述")
    description_table = document.add_table(rows=2, cols=1)
    set_cell_text(description_table.rows[0].cells[0], report["service_description"] or "")
    description_table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(20)
    set_cell_text(description_table.rows[1].cells[0], f"机柜编号：{report['cabinet_number'] or ''}", bold=True)
    style_table(description_table, column_widths=[7.0])

    add_section_title("公共交通时长及自驾车里程细节")
    travel_table = document.add_table(rows=3, cols=2)
    travel_table.rows[0].cells[0].merge(travel_table.rows[0].cells[1])
    set_cell_text(
        travel_table.rows[0].cells[0],
        f"公共交通时长：{report['public_transport_hours'] or 0} 小时    自驾里程总计：{report['driving_miles'] or 0} 英里",
        bold=True,
    )
    set_cell_text(travel_table.rows[1].cells[0], f"出发地址：{report['departure_address'] or ''}")
    set_cell_text(travel_table.rows[1].cells[1], f"场地地址：{report['site_address'] or order['site_address'] or ''}")
    travel_table.rows[2].cells[0].merge(travel_table.rows[2].cells[1])
    set_cell_text(travel_table.rows[2].cells[0], f"合计用时：{report['total_time'] or ''}")
    style_table(travel_table, column_widths=[3.5, 3.5])

    def add_parts_table(title_text, headers, keys, rows, minimum_rows=4):
        add_section_title(title_text)
        table_rows = list(rows)
        while len(table_rows) < minimum_rows:
            table_rows.append(None)
        table = document.add_table(rows=len(table_rows) + 2, cols=len(headers))
        title_cell = table.rows[0].cells[0]
        for index in range(1, len(headers)):
            title_cell = title_cell.merge(table.rows[0].cells[index])
        set_cell_text(table.rows[0].cells[0], title_text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[0], "B6D7A8")
        for index, header in enumerate(headers):
            set_cell_text(table.rows[1].cells[index], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)
        for row_index, part in enumerate(table_rows, start=2):
            for column_index, key in enumerate(keys):
                set_cell_text(table.rows[row_index].cells[column_index], part[key] if part else "", align=WD_ALIGN_PARAGRAPH.CENTER, size=8)
        widths = [7.0 / len(headers)] * len(headers)
        style_table(table, header_rows=2, column_widths=widths)

    add_parts_table(
        "保存的配件",
        ["零件号", "零件名称", "数量", "状态（新/报废）"],
        ["part_number", "part_name", "quantity", "status"],
        report_parts("service_report_saved_parts", report["id"]),
    )
    add_parts_table(
        "现场更换的配件",
        ["零件号", "零件名称", "旧配件序列号", "新配件序列号", "数量"],
        ["part_number", "part_name", "old_serial_number", "new_serial_number", "quantity"],
        report_parts("service_report_replaced_parts", report["id"]),
    )

    add_section_title("现场时间")
    time_table = document.add_table(rows=2, cols=2)
    set_cell_text(time_table.rows[0].cells[0], "到达现场时间", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(time_table.rows[0].cells[1], "离开现场时间", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(time_table.rows[1].cells[0], report["arrival_time"] or "", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(time_table.rows[1].cells[1], report["departure_time"] or "", align=WD_ALIGN_PARAGRAPH.CENTER)
    style_table(time_table, header_rows=1, column_widths=[3.5, 3.5])

    attachment_groups = get_report_attachments(report["id"])
    photo_labels = {
        "arrival": "到达现场时间照片",
        "departure": "离开现场时间照片",
        "self_check": "自检照片",
        "site": "现场服务照片",
    }
    for category, section_title in photo_labels.items():
        photos = attachment_groups.get(category, [])
        if not photos:
            continue
        add_section_title(section_title)
        photo_table = document.add_table(rows=(len(photos) + 1) // 2, cols=2)
        photo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        photo_table.autofit = False
        for index, attachment in enumerate(photos):
            cell = photo_table.rows[index // 2].cells[index % 2]
            path = report_attachment_path(attachment)
            if not os.path.exists(path):
                continue
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = paragraph.add_run()
                run.add_picture(docx_photo_stream(path), width=Inches(3.15))
            except Exception:
                format_run(paragraph.add_run("图片无法嵌入"), size=8)

    document.add_paragraph()
    sign_table = document.add_table(rows=2, cols=2)
    set_cell_text(sign_table.rows[0].cells[0], "报告人签字：")
    set_cell_text(sign_table.rows[0].cells[1], "现场服务人员：")
    set_cell_text(sign_table.rows[1].cells[0], f"报告时间：{report['report_date'] or ''}")
    set_cell_text(sign_table.rows[1].cells[1], "公司：Prasinos Power LLC")
    style_table(sign_table, column_widths=[3.5, 3.5])

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def send_invoice_email(invoice_id):
    invoice, client, items = load_invoice(invoice_id)
    if not client["email"]:
        raise RuntimeError("客户没有邮箱，无法发送。")
    mail_attachments = invoice_email_attachments(invoice, client, items)
    html = render_template(
        "email_invoice.html",
        invoice=invoice,
        client=client,
        items=items,
        totals=invoice_totals(invoice_id),
        company=get_company_profile(),
        terms=get_invoice_terms(),
        payment=get_payment_instructions(),
    )
    send_email(
        to=client["email"],
        subject=f"Invoice {invoice['invoice_number']} from {get_company_profile()['name']}",
        html=html,
        attachments=mail_attachments,
    )
    db().execute("update invoices set sent_at = ? where id = ?", (now(), invoice_id))
    db().commit()


def send_email(to, subject, html, attachments=None):
    smtp_settings = get_smtp_settings()
    host = smtp_settings["host"].strip()
    if not host:
        raise RuntimeError("公司设置中的 SMTP Host 未配置。")
    user = smtp_settings["user"].strip()
    password = smtp_settings["password"]
    if user and not password:
        raise RuntimeError("公司设置中的 SMTP Password 未配置。Gmail 需要使用应用专用密码，不是普通登录密码。")
    message = EmailMessage()
    message["From"] = smtp_settings["from"] or user or "billing@example.com"
    message["To"] = to
    message["Subject"] = subject
    message.set_content("Please see the attached documents.")
    message.add_alternative(html, subtype="html")
    for attachment in attachments or []:
        message.add_attachment(
            attachment["content"],
            maintype=attachment["maintype"],
            subtype=attachment["subtype"],
            filename=attachment["filename"],
        )
    port = int(smtp_settings["port"] or "587")
    use_tls = smtp_settings["tls"].lower() == "true"
    try:
        with smtplib.SMTP(host, port, timeout=30) as smtp:
            if use_tls:
                smtp.starttls()
            if user:
                smtp.login(user, password)
            smtp.send_message(message)
    except smtplib.SMTPAuthenticationError as error:
        raise RuntimeError("Gmail 登录失败。请确认公司设置中的 SMTP Password 是 Gmail 应用专用密码，并且账号已开启两步验证。") from error
    except (smtplib.SMTPException, OSError, ValueError) as error:
        raise RuntimeError(f"邮件发送失败：{error}") from error


@app.errorhandler(RuntimeError)
def runtime_error(error):
    flash(str(error), "error")
    return redirect(request.referrer or url_for("dashboard"))


@app.errorhandler(404)
def not_found(error):
    return (
        render_template(
            "error.html",
            status_code="404",
            title="没有找到这个页面或记录",
            message="你输入的地址可能不正确，或者这张发票、附件、客户资料已经被删除。",
        ),
        404,
    )


def get_metrics():
    access_clause, access_params = client_filter_clause("invoices")
    rows = db().execute(
        f"select id, currency, paid_at, payment_amount from invoices where status = 'completed' and {access_clause}",
        access_params,
    ).fetchall()
    completed = paid = unpaid = invoice_count = 0
    for row in rows:
        total = invoice_totals(row["id"])["total"]
        invoice_count += 1
        completed += total
        if row["paid_at"]:
            paid += float(row["payment_amount"] or total)
        else:
            unpaid += total
    expense_rows = db().execute(
        """
        select status, amount
        from expenses
        where status in ('approved', 'submitted', 'returned')
        """
    ).fetchall()
    pending_reimbursement = db().execute(
        """
        select coalesce(sum(amount), 0) as total
        from expenses
        where status = 'approved' and payout_status != 'paid'
        """
    ).fetchone()["total"]
    reimbursed_expenses = db().execute(
        """
        select coalesce(sum(amount), 0) as total
        from expenses
        where status = 'approved' and payout_status = 'paid'
        """
    ).fetchone()["total"]
    pending_expenses = sum(float(row["amount"] or 0) for row in expense_rows if row["status"] != "approved")
    if is_external_user():
        pending_reimbursement = 0
        reimbursed_expenses = 0
        pending_expenses = 0
    return {
        "invoice_count": invoice_count,
        "completed": completed,
        "paid": paid,
        "unpaid": unpaid,
        "total": completed,
        "pending_reimbursement": float(pending_reimbursement or 0),
        "reimbursed_expenses": float(reimbursed_expenses or 0),
        "pending_expenses": pending_expenses,
    }


def monthly_chart():
    access_clause, access_params = client_filter_clause("invoices")
    rows = db().execute(
        f"select id, issue_date from invoices where status = 'completed' and {access_clause} order by issue_date asc",
        access_params,
    ).fetchall()
    buckets = {}
    for row in rows:
        month = row["issue_date"][:7]
        buckets[month] = buckets.get(month, 0) + invoice_totals(row["id"])["total"]
    if not buckets:
        return []
    max_value = max(buckets.values()) or 1
    return [{"month": month, "value": value, "height": round(value / max_value * 100)} for month, value in buckets.items()]


def available_projects_for_access():
    access_clause, access_params = client_filter_clause("invoices")
    rows = db().execute(
        f"""
        select distinct projects.id, projects.name
        from projects
        join invoice_items on invoice_items.project_id = projects.id
        join invoices on invoices.id = invoice_items.invoice_id
        where invoices.status = 'completed' and {access_clause}
        order by projects.name
        """,
        access_params,
    ).fetchall()
    if rows:
        return rows
    return db().execute(
        "select id, name from projects where project_type = 'invoice' and is_active = 1 order by name"
    ).fetchall()


def dashboard_project_options():
    return [
        {"id": row["id"], "name": row["name"], "color": project_color(row["id"])}
        for row in available_projects_for_access()
    ]


def report_project_options():
    if is_external_user():
        access_clause, access_params = client_filter_clause("invoices")
        return db().execute(
            f"""
            select distinct projects.id, projects.name
            from projects
            join invoice_items on invoice_items.project_id = projects.id
            join invoices on invoices.id = invoice_items.invoice_id
            where invoices.status != 'void' and {access_clause}
            order by projects.name
            """,
            access_params,
        ).fetchall()
    return db().execute(
        """
        select id, name from projects where project_type = 'invoice' and is_active = 1
        union
        select distinct projects.id, projects.name
        from projects
        join invoice_items on invoice_items.project_id = projects.id
        join invoices on invoices.id = invoice_items.invoice_id
        where invoices.status != 'void'
        order by name
        """
    ).fetchall()


def monthly_project_chart(selected_project_ids):
    access_clause, access_params = client_filter_clause("invoices")
    clauses = ["invoices.status = 'completed'", access_clause]
    params = list(access_params)
    if selected_project_ids:
        placeholders = ",".join("?" for _ in selected_project_ids)
        clauses.append(f"invoice_items.project_id in ({placeholders})")
        params.extend(selected_project_ids)
    rows = db().execute(
        f"""
        select substr(invoices.issue_date, 1, 7) as month, projects.id as project_id,
               projects.name as project_name, sum(invoice_items.amount + invoice_items.amount * invoice_items.tax_rate / 100) as total
        from invoice_items
        join invoices on invoices.id = invoice_items.invoice_id
        join projects on projects.id = invoice_items.project_id
        where {" and ".join(clauses)}
        group by month, projects.id, projects.name
        order by month asc, projects.name asc
        """,
        params,
    ).fetchall()
    buckets = {}
    project_names = {}
    for row in rows:
        month = row["month"]
        amount = float(row["total"] or 0)
        project_id = row["project_id"]
        project_names[project_id] = row["project_name"]
        buckets.setdefault(month, {"month": month, "total": 0, "segments": []})
        buckets[month]["total"] += amount
        buckets[month]["segments"].append(
            {"project_id": project_id, "name": row["project_name"], "value": amount, "color": project_color(project_id)}
        )
    if not buckets:
        return []
    max_total = max(bucket["total"] for bucket in buckets.values()) or 1
    for bucket in buckets.values():
        bucket["height"] = round(bucket["total"] / max_total * 100)
        for segment in bucket["segments"]:
            segment["height"] = round(segment["value"] / bucket["total"] * 100) if bucket["total"] else 0
    return list(buckets.values())


def monthly_paid_chart():
    access_clause, access_params = client_filter_clause("invoices")
    rows = db().execute(
        f"select id, paid_at, payment_amount from invoices where status = 'completed' and paid_at is not null and {access_clause} order by paid_at asc",
        access_params,
    ).fetchall()
    buckets = {}
    for row in rows:
        month = row["paid_at"][:7]
        amount = float(row["payment_amount"] or invoice_totals(row["id"])["total"])
        buckets[month] = buckets.get(month, 0) + amount
    if not buckets:
        return []
    max_value = max(buckets.values()) or 1
    return [{"month": month, "value": value, "height": round(value / max_value * 100)} for month, value in buckets.items()]


if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=8000, debug=True)
else:
    init_db()

